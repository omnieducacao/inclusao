# pei_functions.py
from datetime import date, datetime
import re
import json
import base64
import os
from io import BytesIO
from docx import Document
from openai import OpenAI
from pypdf import PdfReader
from fpdf import FPDF
import requests

# ==============================================================================
# LISTAS DE DADOS
# ==============================================================================
LISTA_SERIES = [
    "Educação Infantil (Creche)", "Educação Infantil (Pré-Escola)",
    "1º Ano (Fund. I)", "2º Ano (Fund. I)", "3º Ano (Fund. I)", "4º Ano (Fund. I)", "5º Ano (Fund. I)",
    "6º Ano (Fund. II)", "7º Ano (Fund. II)", "8º Ano (Fund. II)", "9º Ano (Fund. II)",
    "1ª Série (EM)", "2ª Série (EM)", "3ª Série (EM)", "EJA (Educação de Jovens e Adultos)"
]

LISTA_ALFABETIZACAO = [
    "Não se aplica (Educação Infantil)",
    "Pré-Silábico (Garatuja/Desenho sem letras)",
    "Pré-Silábico (Letras aleatórias sem valor sonoro)",
    "Silábico (Sem valor sonoro convencional)",
    "Silábico (Com valor sonoro vogais/consoantes)",
    "Silábico-Alfabético (Transição)",
    "Alfabético (Escrita fonética, com erros ortográficos)",
    "Ortográfico (Escrita convencional consolidada)"
]

LISTAS_BARREIRAS = {
    "Funções Cognitivas": ["Atenção Sustentada/Focada", "Memória de Trabalho (Operacional)", "Flexibilidade Mental", "Planejamento e Organização", "Velocidade de Processamento", "Abstração e Generalização"],
    "Comunicação e Linguagem": ["Linguagem Expressiva (Fala)", "Linguagem Receptiva (Compreensão)", "Pragmática (Uso social da língua)", "Processamento Auditivo", "Intenção Comunicativa"],
    "Socioemocional": ["Regulação Emocional (Autocontrole)", "Tolerância à Frustração", "Interação Social com Pares", "Autoestima e Autoimagem", "Reconhecimento de Emoções"],
    "Sensorial e Motor": ["Praxias Globais (Coordenação Grossa)", "Praxias Finas (Coordenação Fina)", "Hipersensibilidade Sensorial", "Hipossensibilidade (Busca Sensorial)", "Planejamento Motor"],
    "Acadêmico": ["Decodificação Leitora", "Compreensão Textual", "Raciocínio Lógico-Matemático", "Grafomotricidade (Escrita manual)", "Produção Textual"]
}

LISTA_POTENCIAS = ["Memória Visual", "Musicalidade/Ritmo", "Interesse em Tecnologia", "Hiperfoco Construtivo", "Liderança Natural", "Habilidades Cinestésicas (Esportes)", "Expressão Artística (Desenho)", "Cálculo Mental Rápido", "Oralidade/Vocabulário", "Criatividade/Imaginação", "Empatia/Cuidado com o outro", "Resolução de Problemas", "Curiosidade Investigativa"]

LISTA_PROFISSIONAIS = ["Psicólogo Clínico", "Neuropsicólogo", "Fonoaudiólogo", "Terapeuta Ocupacional", "Neuropediatra", "Psiquiatra Infantil", "Psicopedagogo Clínico", "Professor de Apoio (Mediador)", "Acompanhante Terapêutico (AT)", "Musicoterapeuta", "Equoterapeuta", "Oftalmologista"]

LISTA_FAMILIA = ["Mãe", "Mãe 2", "Pai", "Pai 2", "Madrasta", "Padrasto", "Avó Materna", "Avó Paterna", "Avô Materno", "Avô Paterno", "Irmãos", "Tios", "Primos", "Tutor Legal", "Abrigo Institucional"]

# ==============================================================================
# FUNÇÕES AUXILIARES
# ==============================================================================
def calcular_idade(data_nasc):
    if not data_nasc:
        return ""
    hoje = date.today()
    idade = hoje.year - data_nasc.year - ((hoje.month, hoje.day) < (data_nasc.month, data_nasc.day))
    return f"{idade} anos"

def detectar_nivel_ensino(serie_str: str | None):
    if not serie_str:
        return "INDEFINIDO"
    s = serie_str.lower()
    if "infantil" in s:
        return "EI"
    if "1º ano" in s or "2º ano" in s or "3º ano" in s or "4º ano" in s or "5º ano" in s:
        return "FI"
    if "6º ano" in s or "7º ano" in s or "8º ano" in s or "9º ano" in s:
        return "FII"
    if "série" in s or "médio" in s or "eja" in s:
        return "EM"
    return "INDEFINIDO"

def get_segmento_info_visual(serie: str | None):
    nivel = detectar_nivel_ensino(serie or "")
    if nivel == "EI":
        return "Educação Infantil", "#4299e1", "Foco: Campos de Experiência (BNCC)."
    if nivel == "FI":
        return "Anos Iniciais (Fund. I)", "#48bb78", "Foco: Alfabetização e BNCC."
    if nivel == "FII":
        return "Anos Finais (Fund. II)", "#ed8936", "Foco: Autonomia e Identidade."
    if nivel == "EM":
        return "Ensino Médio / EJA", "#9f7aea", "Foco: Projeto de Vida."
    return "Selecione a Série", "grey", "Aguardando seleção..."

def get_hiperfoco_emoji(texto: str | None):
    if not texto:
        return "🚀"
    t = texto.lower()
    if "jogo" in t or "game" in t or "minecraft" in t or "roblox" in t:
        return "🎮"
    if "dino" in t:
        return "🦖"
    if "fute" in t or "bola" in t:
        return "⚽"
    if "desenho" in t or "arte" in t:
        return "🎨"
    if "músic" in t or "music" in t:
        return "🎵"
    if "anim" in t or "gato" in t or "cachorro" in t:
        return "🐾"
    if "carro" in t:
        return "🏎️"
    if "espaço" in t or "espaco" in t:
        return "🪐"
    return "🚀"

def calcular_complexidade_pei(dados: dict):
    n_bar = sum(len(v) for v in (dados.get("barreiras_selecionadas") or {}).values())
    n_suporte_alto = sum(
        1 for v in (dados.get("niveis_suporte") or {}).values()
        if v in ["Substancial", "Muito Substancial"]
    )
    recursos = 0
    if dados.get("rede_apoio"):
        recursos += 3
    if dados.get("lista_medicamentos"):
        recursos += 2
    saldo = (n_bar + n_suporte_alto) - recursos
    if saldo <= 2:
        return "FLUIDA", "#F0FFF4", "#276749"
    if saldo <= 7:
        return "ATENÇÃO", "#FFFFF0", "#D69E2E"
    return "CRÍTICA", "#FFF5F5", "#C53030"

def extrair_tag_ia(texto: str, tag: str):
    if not texto:
        return ""
    padrao = fr"\[{tag}\](.*?)(\[|$)"
    match = re.search(padrao, texto, re.DOTALL)
    return match.group(1).strip() if match else ""

def extrair_metas_estruturadas(texto: str):
    bloco = extrair_tag_ia(texto or "", "METAS_SMART")
    metas = {"Curto": "Definir...", "Medio": "Definir...", "Longo": "Definir..."}
    if bloco:
        linhas = bloco.split("\n")
        for l in linhas:
            l_clean = re.sub(r"^[\-\*]+", "", l).strip()
            if not l_clean:
                continue
            if "Curto" in l or "2 meses" in l:
                metas["Curto"] = l_clean.split(":")[-1].strip()
            elif "Médio" in l or "Semestre" in l or "Medio" in l:
                metas["Medio"] = l_clean.split(":")[-1].strip()
            elif "Longo" in l or "Ano" in l:
                metas["Longo"] = l_clean.split(":")[-1].strip()
    return metas

def get_pro_icon(nome_profissional: str | None):
    p = (nome_profissional or "").lower()
    if "psic" in p:
        return "🧠"
    if "fono" in p:
        return "🗣️"
    if "terapeuta" in p or "equo" in p or "musico" in p:
        return "🧩"
    if "neuro" in p or "psiq" in p or "medico" in p:
        return "🩺"
    return "👨‍⚕️"

def inferir_componentes_impactados(dados: dict):
    barreiras = dados.get("barreiras_selecionadas", {}) or {}
    serie = (dados.get("serie") or "")
    nivel = detectar_nivel_ensino(serie)
    impactados = set()

    # Leitura
    if barreiras.get("Acadêmico") and any("Leitora" in b for b in barreiras["Acadêmico"]):
        impactados.add("Língua Portuguesa")
        impactados.add("História/Sociologia/Filosofia" if nivel == "EM" else "História/Geografia")

    # Matemática
    if barreiras.get("Acadêmico") and any("Matemático" in b for b in barreiras["Acadêmico"]):
        impactados.add("Matemática")
        if nivel == "EM":
            impactados.add("Física/Química")
        elif nivel == "FII":
            impactados.add("Ciências")

    # Cognitivas (transversal)
    if barreiras.get("Funções Cognitivas"):
        impactados.add("Transversal (Todas as áreas)")

    # Motor fino
    if barreiras.get("Sensorial e Motor") and any("Fina" in b for b in barreiras["Sensorial e Motor"]):
        impactados.add("Arte")
        impactados.add("Geometria")

    if not impactados and dados.get("diagnostico"):
        return ["Análise Geral (Baseada no Diagnóstico)"]

    return list(impactados) if impactados else ["Nenhum componente específico detectado automaticamente"]

# ==============================================================================
# FUNÇÕES DE PDF/DOCX
# ==============================================================================
def ler_pdf(arquivo_pdf):
    """Lê o conteúdo de um arquivo PDF"""
    try:
        reader = PdfReader(arquivo_pdf)
        text = ""
        for page in reader.pages[:6]:  # lê no máximo 6 páginas
            text += page.extract_text()
        return text
    except Exception as e:
        return f"Erro ao ler PDF: {e}"

def limpar_texto_pdf(texto: str):
    if not texto:
        return ""
    t = texto.replace("**", "").replace("__", "").replace("#", "").replace("•", "-")
    t = t.replace("“", '"').replace("”", '"').replace("‘", "'").replace("’", "'")
    t = t.replace("–", "-").replace("—", "-")
    return t.encode("latin-1", "replace").decode("latin-1")

class PDF_Classic(FPDF):
    def header(self):
        self.set_fill_color(248, 248, 248)
        self.rect(0, 0, 210, 40, "F")
        logo = None
        # Tenta encontrar o logo
        for f in ["omni_icone.png", "logo.png", "iconeaba.png"]:
            if os.path.exists(f):
                logo = f
                break
        x_offset = 40 if logo else 12
        if logo:
            self.image(logo, 10, 8, 25)
        self.set_xy(x_offset, 12)
        self.set_font("Arial", "B", 14)
        self.set_text_color(50, 50, 50)
        self.cell(0, 8, "PEI - PLANO DE ENSINO INDIVIDUALIZADO", 0, 1, "L")
        self.set_xy(x_offset, 19)
        self.set_font("Arial", "", 9)
        self.set_text_color(100, 100, 100)
        self.cell(0, 5, "Documento de Planejamento e Flexibilização Curricular", 0, 1, "L")
        self.ln(15)

    def footer(self):
        self.set_y(-15)
        self.set_font("Arial", "I", 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, f"Página {self.page_no()} | Gerado via Omnisfera", 0, 0, "C")

    def section_title(self, label):
        self.ln(6)
        self.set_fill_color(230, 230, 230)
        self.rect(10, self.get_y(), 190, 8, "F")
        self.set_font("ZapfDingbats", "", 10)
        self.set_text_color(80, 80, 80)
        self.set_xy(12, self.get_y() + 1)
        self.cell(5, 6, "o", 0, 0)
        self.set_font("Arial", "B", 11)
        self.set_text_color(50, 50, 50)
        self.cell(0, 6, label.upper(), 0, 1, "L")
        self.ln(4)

    def add_flat_icon_item(self, texto, bullet_type="check"):
        self.set_font("ZapfDingbats", "", 10)
        self.set_text_color(80, 80, 80)
        char = "3" if bullet_type == "check" else "l"
        self.cell(6, 5, char, 0, 0)
        self.set_font("Arial", "", 10)
        self.set_text_color(0)
        self.multi_cell(0, 5, texto)
        self.ln(1)

class PDF_Simple_Text(FPDF):
    def header(self):
        self.set_font("Arial", "B", 16)
        self.set_text_color(50)
        self.cell(0, 10, "ROTEIRO DE MISSÃO", 0, 1, "C")
        self.set_draw_color(150)
        self.line(10, 25, 200, 25)
        self.ln(10)

def gerar_pdf_final(dados: dict, tem_pdf: bool = False):
    pdf = PDF_Classic()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    pdf.section_title("Identificação e Contexto")
    pdf.set_font("Arial", "B", 10)
    pdf.cell(35, 6, "Estudante:", 0, 0)
    pdf.set_font("Arial", "", 10)
    pdf.cell(0, 6, dados.get("nome", ""), 0, 1)

    pdf.set_font("Arial", "B", 10)
    pdf.cell(35, 6, "Série/Turma:", 0, 0)
    pdf.set_font("Arial", "", 10)
    pdf.cell(0, 6, f"{dados.get('serie','')} - {dados.get('turma','')}", 0, 1)

    pdf.set_font("Arial", "B", 10)
    pdf.cell(35, 6, "Diagnóstico:", 0, 0)
    pdf.set_font("Arial", "", 10)
    pdf.multi_cell(0, 6, dados.get("diagnostico", ""))
    pdf.ln(2)

    if any((dados.get("barreiras_selecionadas") or {}).values()):
        pdf.section_title("Plano de Suporte (Barreiras x Nível)")
        for area, itens in (dados.get("barreiras_selecionadas") or {}).items():
            if itens:
                pdf.set_font("Arial", "B", 10)
                pdf.cell(0, 8, limpar_texto_pdf(area), 0, 1)
                for item in itens:
                    nivel = (dados.get("niveis_suporte") or {}).get(f"{area}_{item}", "Monitorado")
                    pdf.add_flat_icon_item(limpar_texto_pdf(f"{item} (Nível: {nivel})"), "check")

    if dados.get("ia_sugestao"):
        pdf.add_page()
        pdf.section_title("Planejamento Pedagógico Detalhado")
        texto_limpo = limpar_texto_pdf(dados["ia_sugestao"])
        texto_limpo = re.sub(r"\[.*?\]", "", texto_limpo)

        for linha in texto_limpo.split("\n"):
            l = linha.strip()
            if not l:
                continue
            if l.startswith("###") or l.startswith("##"):
                pdf.ln(5)
                pdf.set_font("Arial", "B", 12)
                pdf.set_text_color(0, 51, 102)
                pdf.cell(0, 8, l.replace("#", "").strip(), 0, 1, "L")
                pdf.set_font("Arial", "", 10)
                pdf.set_text_color(0, 0, 0)
            elif l.startswith("-") or l.startswith("*"):
                pdf.add_flat_icon_item(l.replace("-", "").replace("*", "").strip(), "dot")
            else:
                pdf.multi_cell(0, 6, l)

    return pdf.output(dest="S").encode("latin-1", "replace")

def gerar_pdf_tabuleiro_simples(texto: str):
    pdf = PDF_Simple_Text()
    pdf.add_page()
    pdf.set_font("Arial", size=11)
    for linha in limpar_texto_pdf(texto).split("\n"):
        l = linha.strip()
        if not l:
            continue
        if l.isupper() or "**" in linha:
            pdf.ln(4)
            pdf.set_font("Arial", "B", 11)
            pdf.set_fill_color(240, 240, 240)
            pdf.cell(0, 8, l.replace("**", ""), 0, 1, "L", fill=True)
            pdf.set_font("Arial", "", 11)
        else:
            pdf.multi_cell(0, 6, l)
    return pdf.output(dest="S").encode("latin-1", "ignore")

def gerar_docx_final(dados: dict):
    doc = Document()
    doc.add_heading("PEI - " + (dados.get("nome") or "Sem Nome"), 0)
    if dados.get("ia_sugestao"):
        doc.add_paragraph(re.sub(r"\[.*?\]", "", dados["ia_sugestao"]))
    b = BytesIO()
    doc.save(b)
    b.seek(0)
    return b

# ==============================================================================
# FUNÇÕES DE IA
# ==============================================================================
def extrair_dados_pdf_ia(api_key: str, texto_pdf: str):
    if not api_key:
        return None, "Configure a Chave API OpenAI."
    try:
        client = OpenAI(api_key=api_key)
        prompt = (
            "Analise este laudo médico/escolar. Extraia: 1) Diagnóstico; 2) Medicamentos. "
            'Responda em JSON no formato: { "diagnostico": "...", "medicamentos": [ {"nome": "...", "posologia": "..."} ] }. '
            f"Texto: {texto_pdf[:4000]}"
        )
        res = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
        )
        return json.loads(res.choices[0].message.content), None
    except Exception as e:
        return None, str(e)

def consultar_gpt_pedagogico(api_key: str, dados: dict, contexto_pdf: str = "", modo_pratico: bool = False, feedback_usuario: str = ""):
    if not api_key:
        return None, "⚠️ Configure a Chave API OpenAI."
    try:
        client = OpenAI(api_key=api_key)

        evid = "\n".join([f"- {k.replace('?', '')}" for k, v in (dados.get("checklist_evidencias") or {}).items() if v])
        meds_info = "\n".join(
            [f"- {m.get('nome','')} ({m.get('posologia','')})." for m in (dados.get("lista_medicamentos") or [])]
        ) if dados.get("lista_medicamentos") else "Nenhuma medicação informada."

        hiperfoco_txt = f"HIPERFOCO DO ALUNO: {dados.get('hiperfoco','')}" if dados.get("hiperfoco") else "Hiperfoco: Não identificado."
        serie = dados.get("serie") or ""
        nivel_ensino = detectar_nivel_ensino(serie)
        alfabetizacao = dados.get("nivel_alfabetizacao", "Não Avaliado")

        prompt_identidade = f"""
[PERFIL_NARRATIVO]
Inicie com "👤 QUEM É O ESTUDANTE?". Crie um parágrafo humanizado. {hiperfoco_txt}.
Use o hiperfoco para conectar com a aprendizagem.
[/PERFIL_NARRATIVO]
""".strip()

        prompt_diagnostico = """
### 1. 🏥 DIAGNÓSTICO E IMPACTO (FUNDAMENTAL):
- Cite o Diagnóstico (e o CID se disponível).
- Descreva os impactos diretos na aprendizagem.
- Liste cuidados/pontos de atenção.
""".strip()

        prompt_literacia = ""
        if "Alfabético" not in alfabetizacao and alfabetizacao != "Não se aplica (Educação Infantil)":
            prompt_literacia = f"""[ATENÇÃO CRÍTICA: ALFABETIZAÇÃO] Fase: {alfabetizacao}. Inclua 2 ações de consciência fonológica.[/ATENÇÃO CRÍTICA]"""

        prompt_hub = """
### 6. 🧩 CHECKLIST DE ADAPTAÇÃO E ACESSIBILIDADE:
**A. Mediação (Triângulo de Ouro):**
1) Instruções passo a passo
2) Fragmentação de tarefas
3) Scaffolding

**B. Acessibilidade:**
4) Inferências/figuras de linguagem
5) Descrição de imagens (alt text)
6) Adaptação visual (fonte/espaçamento)
7) Adequação de desafio
""".strip()

        prompt_componentes = ""
        if nivel_ensino != "EI":
            prompt_componentes = f"""
### 4. ⚠️ COMPONENTES CURRICULARES DE ATENÇÃO:
Com base no diagnóstico ({dados.get('diagnostico','')}) e nas barreiras citadas, identifique componentes que exigirão maior flexibilização.
- Liste componentes
- Para cada um, explique o motivo técnico
""".strip()

        prompt_metas = """
[METAS_SMART]
- Meta de Curto Prazo (2 meses): [Descreva a meta]
- Meta de Médio Prazo (1 semestre): [Descreva a meta]
- Meta de Longo Prazo (1 ano): [Descreva a meta]
[/METAS_SMART]
""".strip()

        if nivel_ensino == "EI":
            perfil_ia = "Especialista em EDUCAÇÃO INFANTIL e BNCC."
            estrutura_req = f"""
{prompt_identidade}
{prompt_diagnostico}

### 2. 🌟 AVALIAÇÃO DE REPERTÓRIO:
[CAMPOS_EXPERIENCIA_PRIORITARIOS] Destaque 2 ou 3 Campos BNCC. [/CAMPOS_EXPERIENCIA_PRIORITARIOS]

### 3. 🚀 ESTRATÉGIAS DE INTERVENÇÃO:
(Estratégias de acolhimento, rotina e adaptação sensorial).

{prompt_metas}

### 5. ⚠️ PONTOS DE ATENÇÃO FARMACOLÓGICA:
[ANALISE_FARMA] Se houver medicação, cite efeitos colaterais para atenção pedagógica. [/ANALISE_FARMA]

{prompt_hub}
""".strip()
        else:
            perfil_ia = "Especialista em Inclusão Escolar e BNCC."
            instrucao_bncc = "[MAPEAMENTO_BNCC] Separe por Componente Curricular. Inclua código alfanumérico (ex: EF01LP02). [/MAPEAMENTO_BNCC]"
            instrucao_bloom = "[TAXONOMIA_BLOOM] Explique a categoria cognitiva escolhida. [/TAXONOMIA_BLOOM]"
            estrutura_req = f"""
{prompt_identidade}
{prompt_diagnostico}

### 2. 🌟 AVALIAÇÃO DE REPERTÓRIO:
- Defasagens (anos anteriores)
- Foco do ano atual
{instrucao_bncc}
{instrucao_bloom}

### 3. 🚀 ESTRATÉGIAS DE INTERVENÇÃO:
(Adaptações curriculares e de acesso).
{prompt_literacia}

{prompt_componentes}

{prompt_metas}

### 5. ⚠️ PONTOS DE ATENÇÃO FARMACOLÓGICA:
[ANALISE_FARMA] Se houver medicação, cite efeitos colaterais para atenção pedagógica. [/ANALISE_FARMA]

{prompt_hub}
""".strip()

        prompt_feedback = f"AJUSTE SOLICITADO: {feedback_usuario}" if feedback_usuario else ""
        prompt_formatacao = "IMPORTANTE: Use Markdown simples. Use títulos H3 (###). Evite tabelas."

        prompt_sys = f"""{perfil_ia}
MISSÃO: Criar PEI Técnico Oficial.
ESTRUTURA OBRIGATÓRIA:
{estrutura_req}

{prompt_feedback}
{prompt_formatacao}
""".strip()

        if modo_pratico:
            prompt_sys = f"""{perfil_ia}
GUIA PRÁTICO PARA SALA DE AULA.
{prompt_feedback}

{prompt_hub}
""".strip()

        prompt_user = (
            f"ALUNO: {dados.get('nome','')} | SÉRIE: {serie} | HISTÓRICO: {dados.get('historico','')} | "
            f"DIAGNÓSTICO: {dados.get('diagnostico','')} | MEDS: {meds_info} | "
            f"EVIDÊNCIAS: {evid} | LAUDO: {(contexto_pdf or '')[:3000]}"
        )

        res = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "system", "content": prompt_sys}, {"role": "user", "content": prompt_user}],
        )
        return res.choices[0].message.content, None

    except Exception as e:
        return None, str(e)

def gerar_roteiro_gamificado(api_key: str, dados: dict, pei_tecnico: str, feedback_game: str = ""):
    if not api_key:
        return None, "Configure a chave OpenAI."
    try:
        client = OpenAI(api_key=api_key)
        serie = dados.get("serie") or ""
        nivel_ensino = detectar_nivel_ensino(serie)
        hiperfoco = dados.get("hiperfoco") or "brincadeiras"
        nome_curto = (dados.get("nome","").split() or ["Estudante"])[0]

        contexto_seguro = (
            f"ALUNO: {nome_curto} | HIPERFOCO: {hiperfoco} | "
            f"PONTOS FORTES: {', '.join(dados.get('potencias',[]))}"
        )

        prompt_feedback = f"AJUSTE: {feedback_game}" if feedback_game else ""

        if nivel_ensino == "EI":
            prompt_sys = "Crie uma história visual (4-5 anos) com emojis. Estrutura: começo, desafio, ajuda, conquista, rotina."
        elif nivel_ensino == "FI":
            prompt_sys = "Crie um quadro de missões RPG (6-10 anos). Estrutura: mapa, missões, recompensas, superpoder."
        else:
            prompt_sys = "Crie uma ficha RPG (adolescente). Estrutura: quest, skills, buffs, checklists e metas."

        full_sys = f"{prompt_sys}\n{prompt_feedback}"
        res = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "system", "content": full_sys}, {"role": "user", "content": contexto_seguro}],
        )
        return res.choices[0].message.content, None
    except Exception as e:
        return None, str(e)
