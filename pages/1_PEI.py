# pages/1_PEI.py
from __future__ import annotations

import streamlit as st
from datetime import date, datetime
from zoneinfo import ZoneInfo
from io import BytesIO
from docx import Document
from openai import OpenAI
from pypdf import PdfReader
from fpdf import FPDF
import requests
import base64
import json
import os
import time
import re
import csv

import omni_utils as ou  # módulo atualizado
from omni_utils import get_icon, icon_title



# ✅ set_page_config UMA VEZ SÓ, SEMPRE no topo
st.set_page_config(
    page_title="Omnisfera | PEI",
    page_icon="omni_icone.png",
    layout="wide",
    initial_sidebar_state="collapsed",
)

APP_VERSION = "v150.0 (SaaS Design)"

# ✅ UI lockdown (não quebra se faltar)
try:
    from ui_lockdown import hide_streamlit_chrome_if_needed, hide_default_sidebar_nav
    hide_streamlit_chrome_if_needed()
    hide_default_sidebar_nav()
except Exception:
    pass

# ✅ Header + Navbar (depois do page_config)
ou.render_omnisfera_header()
ou.render_navbar(active_tab="Estratégias & PEI")
if "pei" not in ou.get_enabled_modules():
    st.warning("Este módulo está desabilitado para sua escola.")
    if st.button("Voltar ao Início"):
        st.switch_page("pages/0_Home.py")
    st.stop()
ou.inject_compact_app_css()
# Overlay de loading (ícone girando) quando a IA está gerando algo
ou.inject_loading_overlay_css()

# Adiciona classe no body para cores específicas das abas
st.markdown("<script>document.body.classList.add('page-blue');</script>", unsafe_allow_html=True)

# ==============================================================================
# AJUSTE FINO DE LAYOUT (ANTES DO HERO - PADRONIZADO)
# ==============================================================================
def forcar_layout_hub():
    st.markdown("""
        <style>
            /* 1. Remove o cabeçalho padrão do Streamlit e a linha colorida */
            header[data-testid="stHeader"] {
                visibility: hidden !important;
                height: 0px !important;
            }

            /* 2. Puxa todo o conteúdo para cima (O SEGREDO ESTÁ AQUI) - ESPECIFICIDADE MÁXIMA */
            body .main .block-container,
            body .block-container,
            .main .block-container,
            .block-container {
                padding-top: 0px !important; /* SEM espaço entre navbar e hero */
                padding-bottom: 1rem !important;
                margin-top: 0px !important;
            }
            
            /* 3. Remove qualquer espaçamento do Streamlit */
            [data-testid="stVerticalBlock"],
            div[data-testid="stVerticalBlock"] > div:first-child,
            .main .block-container > div:first-child,
            .main .block-container > *:first-child {
                padding-top: 0px !important;
                margin-top: 0px !important;
            }
            
            /* 4. Remove espaçamento do stMarkdown que renderiza o hero */
            .main .block-container > div:first-child .stMarkdown {
                margin-top: 0px !important;
                padding-top: 0px !important;
            }
            
            /* 5. Hero card colado no menu - margin negativo (ajustado para não ficar muito colado) */
            .mod-card-wrapper {
                margin-top: -96px !important; /* Puxa o hero para cima, mas não tanto quanto Estudantes */
                position: relative;
                z-index: 1;
            }
            
            /* 6. Esconde o menu hambúrguer e rodapé */
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
        </style>
    """, unsafe_allow_html=True)

# CHAME ESTA FUNÇÃO ANTES DO HERO CARD
forcar_layout_hub()

# Cores dos hero cards (paleta vibrante)
ou.inject_hero_card_colors()
# CSS padronizado: abas (pílulas), botões, selects, etc.
ou.inject_unified_ui_css()

# ==============================================================================
# HERO - PEI (ÚNICO)
# ==============================================================================
hora = datetime.now(ZoneInfo("America/Sao_Paulo")).hour
saudacao = "Bom dia" if 5 <= hora < 12 else "Boa tarde" if 12 <= hora < 18 else "Boa noite"
USUARIO_NOME = st.session_state.get("usuario_nome", "Visitante").split()[0]
WORKSPACE_NAME = st.session_state.get("workspace_name", "Workspace")

st.markdown(f"""
<div class="mod-card-wrapper">
    <div class="mod-card-rect">
        <div class="mod-bar c-blue"></div>
        <div class="mod-icon-area bg-blue-soft">
            <i class="ri-book-3-fill"></i>
        </div>
        <div class="mod-content">
            <div class="mod-title">Plano Educacional Individualizado (PEI)</div>
            <div class="mod-desc">
                {saudacao}, <strong>{USUARIO_NOME}</strong>! Crie e gerencie Planos Educacionais Individualizados 
                para estudantes do workspace <strong>{WORKSPACE_NAME}</strong>. 
                Desenvolva estratégias personalizadas e acompanhe o progresso de cada estudante.
            </div>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)

# Espaçamento após hero card
st.markdown("<div style='height:20px'></div>", unsafe_allow_html=True)

# ==============================================================================
# OPENAI (env, secrets ou session_state — mesma lógica do PAEE)
# ==============================================================================
api_key = os.environ.get("OPENAI_API_KEY") or ou.get_setting("OPENAI_API_KEY", "") or st.session_state.get("OPENAI_API_KEY", "")
api_key = (api_key or "").strip() or None


# ==============================================================================
# 1. GUARDAS (LOGIN + SUPABASE)
# ==============================================================================
def verificar_login_app():
    if "autenticado" not in st.session_state or not st.session_state["autenticado"]:
        ou.render_acesso_bloqueado("Faça login na Página Inicial para acessar o PEI.")

def verificar_login_supabase():
    # Supabase é necessário para SALVAR/CARREGAR, mas o PEI pode abrir como rascunho.
    # Então aqui só garantimos chaves mínimas (não bloqueia).
    if "supabase_jwt" not in st.session_state:
        st.session_state["supabase_jwt"] = ""
    if "supabase_user_id" not in st.session_state:
        st.session_state["supabase_user_id"] = ""

verificar_login_app()
verificar_login_supabase()

try:
    from ui.permissions import can_access
    if not can_access("pei"):
        ou.render_acesso_bloqueado(
            "Você não tem permissão para acessar o PEI.",
            "Entre em contato com o responsável pela escola.",
        )
except Exception:
    pass


# =============================================================================
# 2. SUPABASE (CRUD students) — REST (compatível com omni_utils.py)
#    Remove dependência de: sb / OWNER_ID / supabase-py
# =============================================================================

def _rest_ready(debug: bool = False):
    """
    Checa se a nuvem está pronta para operar via REST:
    - autenticado
    - workspace_id presente
    - SUPABASE_URL e alguma KEY (SERVICE ou ANON) presentes
    """
    details = {}
    details["autenticado"] = bool(st.session_state.get("autenticado", False))
    details["has_workspace_id"] = bool(st.session_state.get("workspace_id"))

    try:
        details["has_supabase_url"] = bool(str(os.environ.get("SUPABASE_URL") or ou.get_setting("SUPABASE_URL", "")).strip())
    except Exception:
        details["has_supabase_url"] = False

    try:
        service = str(os.environ.get("SUPABASE_SERVICE_KEY") or ou.get_setting("SUPABASE_SERVICE_KEY", "")).strip()
        anon = str(os.environ.get("SUPABASE_ANON_KEY") or ou.get_setting("SUPABASE_ANON_KEY", "")).strip()
        details["has_supabase_key"] = bool(service or anon)
    except Exception:
        details["has_supabase_key"] = False

    ok = all(details.values())
    if debug:
        details["missing"] = [k for k, v in details.items() if not v]
    return ok, details


def _sb_url() -> str:
    url = str(os.environ.get("SUPABASE_URL") or ou.get_setting("SUPABASE_URL", "")).strip()
    if not url:
        raise RuntimeError("SUPABASE_URL não encontrado nos secrets.")
    return url.rstrip("/")


def _sb_key() -> str:
    # Preferência: SERVICE_KEY (server-side), fallback: ANON_KEY
    key = str(os.environ.get("SUPABASE_SERVICE_KEY") or ou.get_setting("SUPABASE_SERVICE_KEY", "")).strip()
    if not key:
        key = str(os.environ.get("SUPABASE_ANON_KEY") or ou.get_setting("SUPABASE_ANON_KEY", "")).strip()
    if not key:
        raise RuntimeError("SUPABASE_SERVICE_KEY/ANON_KEY não encontrado nos secrets.")
    return key


def _headers() -> dict:
    key = _sb_key()
    return {
        "apikey": key,
        "Authorization": f"Bearer {key}",
        "Content-Type": "application/json",
    }


def _http_error(prefix: str, r):
    raise RuntimeError(f"{prefix}: {r.status_code} {r.text}")


def db_create_student(payload: dict):
    """
    Cria estudante em public.students usando REST.
    - injeta workspace_id automaticamente
    - retorna o registro criado
    """
    ok, details = _rest_ready(debug=True)
    if not ok:
        raise RuntimeError(f"Supabase não está pronto (REST). Missing: {details.get('missing')}")

    ws_id = st.session_state.get("workspace_id") or (payload or {}).get("workspace_id")
    if not ws_id or not str(ws_id).strip():
        raise RuntimeError(
            "Nenhum workspace selecionado. Faça login na Página Inicial e selecione o workspace (PIN) antes de criar um novo aluno. "
            "O aluno só aparece nas páginas quando é criado no mesmo workspace em que você está logado."
        )
    row = dict(payload or {})
    row["workspace_id"] = ws_id

    url = f"{_sb_url()}/rest/v1/students"
    h = _headers()
    h["Prefer"] = "return=representation"

    r = requests.post(url, headers=h, json=row, timeout=20)
    if r.status_code >= 400:
        _http_error("Insert em students falhou", r)

    data = r.json()
    if isinstance(data, list) and len(data) > 0:
        return data[0]
    if isinstance(data, dict):
        return data
    return None


def db_update_student(student_id: str, payload: dict):
    """
    Atualiza estudante em public.students (por id) via REST
    - garante workspace_id no filtro
    """
    ok, details = _rest_ready(debug=True)
    if not ok:
        raise RuntimeError(f"Supabase não está pronto (REST). Missing: {details.get('missing')}")

    ws_id = st.session_state.get("workspace_id")
    row = dict(payload or {})

    url = f"{_sb_url()}/rest/v1/students?id=eq.{student_id}&workspace_id=eq.{ws_id}"
    h = _headers()
    h["Prefer"] = "return=representation"

    r = requests.patch(url, headers=h, json=row, timeout=20)
    if r.status_code >= 400:
        _http_error("Update em students falhou", r)

    data = r.json()
    if isinstance(data, list) and len(data) > 0:
        return data[0]
    if isinstance(data, dict):
        return data
    return None


def db_delete_student(student_id: str):
    """
    Deleta estudante em public.students (por id) via REST
    - garante workspace_id no filtro
    """
    ok, details = _rest_ready(debug=True)
    if not ok:
        raise RuntimeError(f"Supabase não está pronto (REST). Missing: {details.get('missing')}")

    ws_id = st.session_state.get("workspace_id")

    url = f"{_sb_url()}/rest/v1/students?id=eq.{student_id}&workspace_id=eq.{ws_id}"
    h = _headers()
    h["Prefer"] = "return=representation"

    r = requests.delete(url, headers=h, timeout=20)
    if r.status_code >= 400:
        _http_error("Delete em students falhou", r)

    return r.json()


def db_list_students(search: str | None = None):
    """
    Lista estudantes do workspace atual.
    Se search vier preenchido, filtra por nome (ilike).
    """
    ok, _ = _rest_ready(debug=False)
    if not ok:
        return []

    ws_id = st.session_state.get("workspace_id")
    base = f"{_sb_url()}/rest/v1/students?select=*&workspace_id=eq.{ws_id}&order=created_at.desc"

    if search:
        s = str(search).strip()
        if s:
            base += f"&name=ilike.*{s}*"

    r = requests.get(base, headers=_headers(), timeout=20)
    if r.status_code >= 400:
        _http_error("List students falhou", r)

    data = r.json()
    data = data if isinstance(data, list) else []
    try:
        from ui.permissions import apply_member_filter
        data = apply_member_filter(data)
    except Exception:
        pass
    return data

def db_update_pei_content(student_id: str, pei_dict: dict):
    """
    Salva o dicionário completo do PEI na coluna 'pei_data' do Supabase.
    """
    # URL para atualizar o estudante específico
    url = f"{_sb_url()}/rest/v1/students?id=eq.{student_id}"
    
    h = _headers()
    h["Prefer"] = "return=representation"
    
    # Prepara o JSON. Convertemos para string/dict puro para garantir que datas não quebrem
    import json
    payload_json = json.loads(json.dumps(pei_dict, default=str))
    
    # Envia apenas o campo pei_data e a data de atualização
    body = {
        "pei_data": payload_json,
        "updated_at": datetime.now().isoformat()
    }
    
    r = requests.patch(url, headers=h, json=body, timeout=20)
    
    if r.status_code >= 400:
        raise RuntimeError(f"Erro ao salvar conteúdo do PEI: {r.text}")
        
    return r.json()

# ==============================================================================
# 
# ==============================================================================



# ==============================================================================
# 4. LISTAS DE DADOS
# ==============================================================================
SERIES_EF_EM = [
    "1º Ano (EFAI)", "2º Ano (EFAI)", "3º Ano (EFAI)", "4º Ano (EFAI)", "5º Ano (EFAI)",
    "6º Ano (EFAF)", "7º Ano (EFAF)", "8º Ano (EFAF)", "9º Ano (EFAF)",
    "1ª Série (EM)", "2ª Série (EM)", "3ª Série (EM)", "EJA (Educação de Jovens e Adultos)"
]

def _construir_lista_series(serie_atual: str | None = None) -> list[str]:
    """
    Constrói lista de séries com faixas EI dinâmicas do bncc_ei.csv.
    Quando marcado Educação Infantil, usa Idade do CSV (2, 3, 4, 5 anos).
    Fallback para (0-2) e (3-5) se CSV vazio. Inclui serie_atual se for legado EI.
    """
    try:
        from services.bncc_service import faixas_idade_ei
        faixas = faixas_idade_ei()
    except Exception:
        faixas = []
    ei_opcoes = []
    if faixas:
        ei_opcoes = [f"Educação Infantil ({f})" for f in faixas]
    else:
        ei_opcoes = ["Educação Infantil (0-2 anos)", "Educação Infantil (3-5 anos)"]
    # Retrocompatibilidade: se aluno salvo com EI legado e não está na lista, incluir
    LEGADO_EI = ["Educação Infantil (0-2 anos)", "Educação Infantil (3-5 anos)"]
    if serie_atual and serie_atual in LEGADO_EI and serie_atual not in ei_opcoes:
        ei_opcoes = [serie_atual] + [e for e in ei_opcoes if e != serie_atual]
    return ei_opcoes + SERIES_EF_EM

LISTA_ALFABETIZACAO = [
    "Não se aplica (Educação Infantil)",
    "Pré-Silábico (Garatuja/Desenho sem letras)",
    "Pré-Silábico (Letras aleatórias sem valor sonoro)",
    "Silábico (Sem valor sonoro convencional)",
    "Silábico (Com valor sonoro vogais/consoantes)",
    "Silábico-Alfabético (Transição)",
    "Alfabético (Escrita fonética, com erros ortográficos)",
    "Ortográfico (Escrita convencional consolidada)"
]

LISTAS_BARREIRAS = {
    "Funções Cognitivas": ["Atenção Sustentada/Focada", "Memória de Trabalho (Operacional)", "Flexibilidade Mental", "Planejamento e Organização", "Velocidade de Processamento", "Abstração e Generalização"],
    "Comunicação e Linguagem": ["Linguagem Expressiva (Fala)", "Linguagem Receptiva (Compreensão)", "Pragmática (Uso social da língua)", "Processamento Auditivo", "Intenção Comunicativa"],
    "Socioemocional": ["Regulação Emocional (Autocontrole)", "Tolerância à Frustração", "Interação Social com Pares", "Autoestima e Autoimagem", "Reconhecimento de Emoções"],
    "Sensorial e Motor": ["Praxias Globais (Coordenação Grossa)", "Praxias Finas (Coordenação Fina)", "Hipersensibilidade Sensorial", "Hipossensibilidade (Busca Sensorial)", "Planejamento Motor"],
    "Acadêmico": ["Decodificação Leitora", "Compreensão Textual", "Raciocínio Lógico-Matemático", "Grafomotricidade (Escrita manual)", "Produção Textual"]
}

LISTA_POTENCIAS = ["Memória Visual", "Musicalidade/Ritmo", "Interesse em Tecnologia", "Hiperfoco Construtivo", "Liderança Natural", "Habilidades Cinestésicas (Esportes)", "Expressão Artística (Desenho)", "Cálculo Mental Rápido", "Oralidade/Vocabulário", "Criatividade/Imaginação", "Empatia/Cuidado com o outro", "Resolução de Problemas", "Curiosidade Investigativa"]

LISTA_PROFISSIONAIS = ["Psicólogo Clínico", "Neuropsicólogo", "Fonoaudiólogo", "Terapeuta Ocupacional", "Neuropediatra", "Psiquiatra Infantil", "Psicopedagogo Clínico", "Professor de Apoio (Mediador)", "Acompanhante Terapêutico (AT)", "Musicoterapeuta", "Equoterapeuta", "Oftalmologista"]

LISTA_FAMILIA = ["Mãe", "Mãe 2", "Pai", "Pai 2", "Madrasta", "Padrasto", "Avó Materna", "Avó Paterna", "Avô Materno", "Avô Paterno", "Irmãos", "Tios", "Primos", "Tutor Legal", "Abrigo Institucional"]


# ==============================================================================
# 5. ESTADO DEFAULT (RASCUNHO)
# ==============================================================================
default_state = {
    "nome": "",
    "nasc": date(2015, 1, 1),
    "serie": None,
    "turma": "",
    "diagnostico": "",
    "lista_medicamentos": [],
    "composicao_familiar_tags": [],
    "historico": "",
    "familia": "",
    "hiperfoco": "",
    "potencias": [],
    "rede_apoio": [],
    "orientacoes_especialistas": "",
    "orientacoes_por_profissional": {},
    "checklist_evidencias": {},
    "nivel_alfabetizacao": "Não se aplica (Educação Infantil)",
    "barreiras_selecionadas": {k: [] for k in LISTAS_BARREIRAS.keys()},
    "niveis_suporte": {},
    "observacoes_barreiras": {},
    "estrategias_acesso": [],
    "estrategias_ensino": [],
    "estrategias_avaliacao": [],
    "ia_sugestao": "",
    "consultoria_engine": "red",  # motor que gerou o relatório (red/blue/green/yellow/orange)
    "ia_mapa_texto": "",
    "outros_acesso": "",
    "outros_ensino": "",
    "monitoramento_data": date.today(),
    "status_meta": "Não Iniciado",
    "parecer_geral": "Manter Estratégias",
    "proximos_passos_select": [],
    "status_validacao_pei": "rascunho",
    "feedback_ajuste": "",
    "status_validacao_game": "rascunho",
    "feedback_ajuste_game": "",
    "matricula": "",
    "meds_extraidas_tmp": [],
    "status_meds_extraidas": "idle",
    "habilidades_bncc_selecionadas": [],  # lista da aba BNCC (EF/EM)
    "habilidades_bncc_validadas": None,   # lista validada pelo professor (Consultoria IA usa só esta quando existir)
    "bncc_ei_idade": "",                  # EI: faixa de idade
    "bncc_ei_campo": "",                  # EI: campo de experiência
    "bncc_ei_objetivos": [],              # EI: objetivos de aprendizagem selecionados
}

if "dados" not in st.session_state:
    st.session_state.dados = default_state
else:
    for k, v in default_state.items():
        if k not in st.session_state.dados:
            st.session_state.dados[k] = v

st.session_state.setdefault("pdf_text", "")

# vínculo supabase
st.session_state.setdefault("selected_student_id", None)
st.session_state.setdefault("selected_student_name", "")




# ==============================================================================
# 7. UTILITÁRIOS
# ==============================================================================
def _is_filled(v):
    if v is None:
        return False
    if isinstance(v, str):
        return bool(v.strip())
    if isinstance(v, (list, tuple, set, dict)):
        return len(v) > 0
    return True

def _aba_ok(d, key):
    # Define o que significa "aba preenchida"
    if key == "INICIO":
        # “Início” é mais informativa; não precisa contar ou conta quando tiver nome
        return _is_filled(d.get("nome"))

    if key == "ESTUDANTE":
        return _is_filled(d.get("nome")) and _is_filled(d.get("serie")) and _is_filled(d.get("turma"))

    if key == "EVIDENCIAS":
        chk = d.get("checklist_evidencias", {}) or {}
        # pelo menos 1 evidência marcada OU texto de orientação
        return any(bool(v) for v in chk.values()) or _is_filled(d.get("orientacoes_especialistas"))

    if key == "REDE":
        # ao menos 1 profissional OU alguma orientação geral
        return _is_filled(d.get("rede_apoio")) or _is_filled(d.get("orientacoes_especialistas")) or _is_filled(d.get("orientacoes_por_profissional"))

    if key == "MAPEAMENTO":
        barreiras = d.get("barreiras_selecionadas", {}) or {}
        n_bar = sum(len(v) for v in barreiras.values()) if isinstance(barreiras, dict) else 0
        return _is_filled(d.get("hiperfoco")) or _is_filled(d.get("potencias")) or (n_bar > 0)

    if key == "PLANO":
        return _is_filled(d.get("estrategias_acesso")) or _is_filled(d.get("estrategias_ensino")) or _is_filled(d.get("estrategias_avaliacao")) \
               or _is_filled(d.get("outros_acesso")) or _is_filled(d.get("outros_ensino"))

    if key == "MONITORAMENTO":
        return _is_filled(d.get("monitoramento_data")) and _is_filled(d.get("status_meta"))

    if key == "IA":
        return _is_filled(d.get("ia_sugestao")) and d.get("status_validacao_pei") in ["revisao", "aprovado"]

    if key == "DASH":
        # dashboard depende do IA
        return _is_filled(d.get("ia_sugestao"))

    if key == "JORNADA":
        return _is_filled(d.get("ia_mapa_texto")) and d.get("status_validacao_game") in ["revisao", "aprovado"]

    return False

def calcular_progresso() -> int:
    d = st.session_state.get("dados", {}) or {}

    # PEI: abas que contam no progresso (sem JORNADA — relatório gerado = 100% com DASH)
    checkpoints = ["ESTUDANTE", "EVIDENCIAS", "REDE", "MAPEAMENTO", "PLANO", "MONITORAMENTO", "IA", "DASH"]

    done = sum(1 for k in checkpoints if _aba_ok(d, k))
    total = len(checkpoints)
    return int(round((done / total) * 100)) if total else 0

# ==============================================================================
# 7B. UTILITÁRIOS AVANÇADOS (idade, segmento, metas, radar, etc.)
# ==============================================================================

def calcular_idade(data_nasc):
    if not data_nasc:
        return ""
    hoje = date.today()
    idade = hoje.year - data_nasc.year - ((hoje.month, hoje.day) < (data_nasc.month, data_nasc.day))
    return f"{idade} anos"

def detectar_nivel_ensino(serie_str: str | None):
    if not serie_str:
        return "INDEFINIDO"
    s = serie_str.lower()
    if "infantil" in s:
        return "EI"
    if "1º ano" in s or "2º ano" in s or "3º ano" in s or "4º ano" in s or "5º ano" in s:
        return "FI"
    if "6º ano" in s or "7º ano" in s or "8º ano" in s or "9º ano" in s:
        return "FII"
    if "série" in s or "médio" in s or "eja" in s:
        return "EM"
    return "INDEFINIDO"


def _extrair_ano_serie_bncc(serie: str) -> str | None:
    """
    Extrai o ano/série no formato do CSV BNCC (ex: 1º, 2º, 6º, 1EM) a partir da série do estudante.
    Suporta: "1º ano", "7º Ano (EFAF)", "1ª Série (EM)" -> "1º", "7º", "1EM"
    """
    if not serie or not isinstance(serie, str):
        return None
    s = serie.strip()
    # 1ª Série, 2ª Série (EM) -> 1EM, 2EM, 3EM
    m_em = re.search(r"(\d)\s*ª?\s*série", s, re.IGNORECASE)
    if m_em or "em" in s.lower() or "médio" in s.lower():
        num = m_em.group(1) if m_em else (re.search(r"(\d)", s))
        if num is not None:
            n = int(num) if isinstance(num, str) else int(num.group(1))
            return f"{n}EM"
    # 1º ano, 7º Ano (EFAF), etc.
    m = re.search(r"(\d\s*º)", s)
    return m.group(1).replace(" ", "") if m else None


def carregar_habilidades_bncc_por_serie(serie: str, max_caracteres: int = 11000) -> str:
    """
    Carrega habilidades BNCC do ano/série (EF via bncc_service).
    Retorna texto formatado para injetar no prompt, para a IA usar APENAS essas habilidades
    e não inventar. Se o arquivo não existir ou a série não bater, retorna string vazia.
    """
    d = carregar_habilidades_bncc_ano_atual_e_anteriores(serie, max_caracteres, max_caracteres)
    return d.get("ano_atual", "")


def carregar_habilidades_bncc_ano_atual_e_anteriores(
    serie: str,
    max_ano_atual: int = 10000,
    max_anos_anteriores: int = 8000,
) -> dict:
    """
    Carrega BNCC do ano/série atual e anteriores (EF via bncc_service).
    Retorna {"ano_atual": "...", "anos_anteriores": "..."} para injetar no prompt.
    """
    from services.bncc_service import carregar_habilidades_ef_ano_atual_e_anteriores
    return carregar_habilidades_ef_ano_atual_e_anteriores(serie, max_ano_atual, max_anos_anteriores)


def _parse_hab_row(hab: str) -> tuple:
    """Extrai código e descrição de uma linha de habilidade (ex: (EF01LP02) Descrição)."""
    codigo, descricao = "", hab
    m = re.match(r"\(([A-Za-z0-9]+)\)\s*(.*)", hab)
    if m:
        codigo, resto = m.group(1), m.group(2)
        descricao = resto.strip() if resto else hab
    return (codigo or "(sem código)", descricao[:200] + ("..." if len(descricao) > 200 else ""))


def _codigos_permitidos(lista_hab: list) -> set:
    """Retorna set de códigos BNCC em maiúscula a partir da lista de dicts (validadas ou selecionadas)."""
    if not lista_hab or not isinstance(lista_hab, list):
        return set()
    return {(h.get("codigo") or "").strip().upper() for h in lista_hab if isinstance(h, dict) and h.get("codigo")}


def filtrar_avaliacao_repertorio_apenas_permitidas(texto: str, lista_hab_permitidas: list) -> str:
    """
    Na seção 'Avaliação de Repertório', remove linhas que citam habilidades com código BNCC
    que NÃO estejam na lista permitida (validadas pelo professor), evitando que a IA tenha inventado.
    """
    if not texto or not lista_hab_permitidas:
        return texto
    permitidos = _codigos_permitidos(lista_hab_permitidas)
    if not permitidos:
        return texto
    # Padrão para códigos BNCC (EF01LP02, EM13CHT301, etc.)
    padrao_codigo = re.compile(r"\b(EF\d+[A-Z0-9]+|EM\d+[A-Z0-9]+)\b", re.IGNORECASE)
    # Encontrar seção Avaliação de Repertório até o próximo ###
    inicio_marker = "AVALIAÇÃO DE REPERTÓRIO"
    partes = re.split(r"(?=^###\s)", texto, flags=re.MULTILINE)
    resultado = []
    for bloco in partes:
        if inicio_marker.upper() not in bloco.upper():
            resultado.append(bloco)
            continue
        linhas = bloco.split("\n")
        linhas_filtradas = []
        for linha in linhas:
            codigos_na_linha = set(c.upper() for c in padrao_codigo.findall(linha))
            if not codigos_na_linha:
                linhas_filtradas.append(linha)
                continue
            if codigos_na_linha.issubset(permitidos):
                linhas_filtradas.append(linha)
            # else: linha cita código não permitido (IA inventou) — não incluir
        resultado.append("\n".join(linhas_filtradas))
    return "".join(resultado)


def carregar_habilidades_bncc_por_componente(serie: str) -> dict:
    """
    Carrega do bncc.csv as habilidades do ano/série atual agrupadas por Disciplina.
    Retorna { "Disciplina": [ {"codigo", "descricao", "habilidade_completa"}, ... ], ... }
    """
    d = carregar_habilidades_bncc_por_componente_ano_e_anteriores(serie)
    atual = d.get("ano_atual") or {}
    # Junta ano_atual para compatibilidade com quem só usa ano atual
    return atual


def carregar_habilidades_bncc_por_componente_ano_e_anteriores(serie: str) -> dict:
    """
    Carrega BNCC do ano/série atual e anteriores agrupadas por Disciplina (EF via bncc_service).
    Retorna { "ano_atual": { "Disciplina": [...] }, "anos_anteriores": {...} }
    """
    from services.bncc_service import carregar_habilidades_ef_por_componente
    return carregar_habilidades_ef_por_componente(serie)


def get_segmento_info_visual(serie: str | None):
    nivel = detectar_nivel_ensino(serie or "")
    if nivel == "EI":
        return "Educação Infantil", "#4299e1", "Foco: Campos de Experiência (BNCC)."
    if nivel == "FI":
        return "Ensino Fundamental Anos Iniciais (EFAI)", "#48bb78", "Foco: Alfabetização e BNCC."
    if nivel == "FII":
        return "Ensino Fundamental Anos Finais (EFAF)", "#ed8936", "Foco: Autonomia e Identidade."
    if nivel == "EM":
        return "Ensino Médio / EJA", "#9f7aea", "Foco: Projeto de Vida."
    return "Selecione a Série", "grey", "Aguardando seleção..."

def get_hiperfoco_emoji(texto: str | None):
    if not texto:
        return "🚀"
    t = texto.lower()
    if "jogo" in t or "game" in t or "minecraft" in t or "roblox" in t:
        return "🎮"
    if "dino" in t:
        return "🦖"
    if "fute" in t or "bola" in t:
        return "⚽"
    if "desenho" in t or "arte" in t:
        return "🎨"
    if "músic" in t or "music" in t:
        return "🎵"
    if "anim" in t or "gato" in t or "cachorro" in t:
        return "🐾"
    if "carro" in t:
        return "🏎️"
    if "espaço" in t or "espaco" in t:
        return "🪐"
    return "🚀"

def calcular_complexidade_pei(dados: dict):
    n_bar = sum(len(v) for v in (dados.get("barreiras_selecionadas") or {}).values())
    n_suporte_alto = sum(
        1 for v in (dados.get("niveis_suporte") or {}).values()
        if v in ["Substancial", "Muito Substancial"]
    )
    recursos = 0
    if dados.get("rede_apoio"):
        recursos += 3
    if dados.get("lista_medicamentos"):
        recursos += 2
    saldo = (n_bar + n_suporte_alto) - recursos
    if saldo <= 2:
        return "FLUIDA", "#F0FFF4", "#276749"
    if saldo <= 7:
        return "ATENÇÃO", "#FFFFF0", "#D69E2E"
    return "CRÍTICA", "#FFF5F5", "#C53030"

def extrair_tag_ia(texto: str, tag: str):
    if not texto:
        return ""
    padrao = fr"\[{tag}\](.*?)(\[|$)"
    match = re.search(padrao, texto, re.DOTALL)
    return match.group(1).strip() if match else ""

def extrair_metas_estruturadas(texto: str):
    bloco = extrair_tag_ia(texto or "", "METAS_SMART")
    metas = {"Curto": "Definir...", "Medio": "Definir...", "Longo": "Definir..."}
    if bloco:
        linhas = bloco.split("\n")
        for l in linhas:
            l_clean = re.sub(r"^[\-\*]+", "", l).strip()
            if not l_clean:
                continue
            if "Curto" in l or "2 meses" in l:
                metas["Curto"] = l_clean.split(":")[-1].strip()
            elif "Médio" in l or "Semestre" in l or "Medio" in l:
                metas["Medio"] = l_clean.split(":")[-1].strip()
            elif "Longo" in l or "Ano" in l:
                metas["Longo"] = l_clean.split(":")[-1].strip()
    return metas

def get_pro_icon(nome_profissional: str | None):
    p = (nome_profissional or "").lower()
    if "psic" in p:
        return "🧠"
    if "fono" in p:
        return "🗣️"
    if "terapeuta" in p or "equo" in p or "musico" in p:
        return "🧩"
    if "neuro" in p or "psiq" in p or "medico" in p:
        return "🩺"
    return "👨‍⚕️"

def inferir_componentes_impactados(dados: dict):
    barreiras = dados.get("barreiras_selecionadas", {}) or {}
    serie = (dados.get("serie") or "")
    nivel = detectar_nivel_ensino(serie)
    impactados = set()

    # Leitura
    if barreiras.get("Acadêmico") and any("Leitora" in b for b in barreiras["Acadêmico"]):
        impactados.add("Língua Portuguesa")
        impactados.add("História/Sociologia/Filosofia" if nivel == "EM" else "História/Geografia")

    # Matemática
    if barreiras.get("Acadêmico") and any("Matemático" in b for b in barreiras["Acadêmico"]):
        impactados.add("Matemática")
        if nivel == "EM":
            impactados.add("Física/Química/Biologia")
        elif nivel == "FII":
            impactados.add("Ciências")

    # Cognitivas (transversal)
    if barreiras.get("Funções Cognitivas"):
        impactados.add("Transversal (Todas as áreas)")

    # Motor fino
    if barreiras.get("Sensorial e Motor") and any("Fina" in b for b in barreiras["Sensorial e Motor"]):
        impactados.add("Arte")
        impactados.add("Geometria")

    if not impactados and dados.get("diagnostico"):
        return ["Análise Geral (Baseada no Diagnóstico)"]

    return list(impactados) if impactados else ["Nenhum componente específico detectado automaticamente"]


# ==============================================================================
# 7C. PDF / DOCX (Exportação)
# ==============================================================================

def limpar_texto_pdf(texto: str):
    if not texto:
        return ""
    t = texto.replace("**", "").replace("__", "").replace("#", "").replace("•", "-")
    t = t.replace("“", '"').replace("”", '"').replace("‘", "'").replace("’", "'")
    t = t.replace("–", "-").replace("—", "-")
    # Usar encode/decode com replace para tratar todos os caracteres não-latin-1
    try:
        return t.encode("latin-1", "replace").decode("latin-1")
    except Exception:
        # Fallback: remover todos os caracteres não-ASCII manualmente
        return ''.join(c if ord(c) < 256 else '?' for c in t)

def finding_logo():
    """Retorna o caminho do arquivo de logo se existir, caso contrário None"""
    for logo_file in ["omni_icone.png", "logo.png", "iconeaba.png"]:
        if os.path.exists(logo_file):
            return logo_file
    return None

# Margens e largura útil do PDF (layout mais respirável)
PDF_LEFT = 15
PDF_RIGHT = 15
PDF_TOP = 20
PDF_BOTTOM = 22
PDF_PAGE_WIDTH = 210
PDF_CONTENT_WIDTH = PDF_PAGE_WIDTH - PDF_LEFT - PDF_RIGHT  # 180

class PDF_Classic(FPDF):
    def header(self):
        self.set_fill_color(248, 250, 252)
        self.rect(0, 0, PDF_PAGE_WIDTH, 42, "F")
        logo = finding_logo()
        x_offset = 45 if logo else PDF_LEFT
        if logo:
            self.image(logo, PDF_LEFT, 10, 28)
        self.set_xy(x_offset, 14)
        self.set_font("Arial", "B", 15)
        self.set_text_color(30, 41, 59)
        self.cell(0, 8, "PEI - PLANO DE ENSINO INDIVIDUALIZADO", 0, 1, "L")
        self.set_xy(x_offset, 22)
        self.set_font("Arial", "", 9)
        self.set_text_color(100, 116, 139)
        self.cell(0, 5, "Documento de Planejamento e Flexibiliza\u00e7\u00e3o Curricular", 0, 1, "L")
        self.set_draw_color(226, 232, 240)
        self.line(PDF_LEFT, 42, PDF_PAGE_WIDTH - PDF_RIGHT, 42)
        self.ln(18)

    def footer(self):
        self.set_y(-PDF_BOTTOM)
        self.set_font("Arial", "I", 8)
        self.set_text_color(148, 163, 184)
        self.cell(0, 10, f"P\u00e1gina {self.page_no()} | Gerado via Omnisfera", 0, 0, "C")

    def section_title(self, label):
        self.ln(8)
        self.set_fill_color(241, 245, 249)
        self.rect(PDF_LEFT, self.get_y(), PDF_CONTENT_WIDTH, 10, "F")
        self.set_font("Arial", "B", 11)
        self.set_text_color(30, 41, 59)
        self.set_xy(PDF_LEFT + 4, self.get_y() + 2)
        self.cell(0, 6, label.upper(), 0, 1, "L")
        self.ln(6)

    def add_flat_icon_item(self, texto, bullet_type="check"):
        self.set_font("ZapfDingbats", "", 9)
        self.set_text_color(100, 116, 139)
        char = "3" if bullet_type == "check" else "l"
        self.cell(6, 6, char, 0, 0)
        self.set_font("Arial", "", 10)
        self.set_text_color(51, 65, 85)
        self.multi_cell(0, 6, texto)
        self.ln(1)

class PDF_Simple_Text(FPDF):
    def header(self):
        self.set_font("Arial", "B", 16)
        self.set_text_color(50)
        self.cell(0, 10, "ROTEIRO DE MISSÃO", 0, 1, "C")
        self.set_draw_color(150)
        self.line(10, 25, 200, 25)
        self.ln(10)

def gerar_pdf_final(dados: dict):
    """Gera PDF completo do PEI com todas as seções - Documento Oficial (layout otimizado)"""
    pdf = PDF_Classic()
    pdf.set_margins(PDF_LEFT, PDF_TOP, PDF_RIGHT)
    pdf.set_auto_page_break(auto=True, margin=PDF_BOTTOM)
    pdf.add_page()

    # ======================================================================
    # 1. IDENTIFICAÇÃO E CONTEXTO
    # ======================================================================
    pdf.section_title("1. IDENTIFICAÇÃO E CONTEXTO")
    pdf.set_font("Arial", "B", 10)
    pdf.cell(40, 6, "Estudante:", 0, 0)
    pdf.set_font("Arial", "", 10)
    pdf.cell(0, 6, limpar_texto_pdf(dados.get("nome", "")), 0, 1)

    pdf.set_font("Arial", "B", 10)
    pdf.cell(40, 6, "Data de Nascimento:", 0, 0)
    pdf.set_font("Arial", "", 10)
    nasc = dados.get("nasc")
    if nasc:
        if isinstance(nasc, str):
            pdf.cell(0, 6, limpar_texto_pdf(nasc), 0, 1)
        elif hasattr(nasc, 'strftime'):
            pdf.cell(0, 6, nasc.strftime("%d/%m/%Y"), 0, 1)
        else:
            pdf.cell(0, 6, limpar_texto_pdf(str(nasc)), 0, 1)
    else:
        pdf.cell(0, 6, "-", 0, 1)

    pdf.set_font("Arial", "B", 10)
    pdf.cell(40, 6, "Série/Ano:", 0, 0)
    pdf.set_font("Arial", "", 10)
    pdf.cell(0, 6, limpar_texto_pdf(dados.get("serie", "")), 0, 1)

    pdf.set_font("Arial", "B", 10)
    pdf.cell(40, 6, "Turma:", 0, 0)
    pdf.set_font("Arial", "", 10)
    pdf.cell(0, 6, limpar_texto_pdf(dados.get("turma", "")), 0, 1)

    pdf.set_font("Arial", "B", 10)
    pdf.cell(40, 6, "Matrícula/RA:", 0, 0)
    pdf.set_font("Arial", "", 10)
    pdf.cell(0, 6, limpar_texto_pdf(dados.get("matricula", dados.get("ra", ""))), 0, 1)
    pdf.ln(3)

    # Histórico Escolar
    if dados.get("historico"):
        pdf.set_font("Arial", "B", 10)
        pdf.cell(0, 6, "Histórico Escolar:", 0, 1)
        pdf.set_font("Arial", "", 10)
        pdf.multi_cell(0, 6, limpar_texto_pdf(dados.get("historico", "")))
        pdf.ln(2)

    # Dinâmica Familiar
    if dados.get("familia"):
        pdf.set_font("Arial", "B", 10)
        pdf.cell(0, 6, "Dinâmica Familiar:", 0, 1)
        pdf.set_font("Arial", "", 10)
        pdf.multi_cell(0, 6, limpar_texto_pdf(dados.get("familia", "")))
        pdf.ln(2)

    # Composição Familiar
    comp_fam = dados.get("composicao_familiar_tags", [])
    if comp_fam:
        pdf.set_font("Arial", "B", 10)
        pdf.cell(0, 6, "Composição Familiar:", 0, 1)
        pdf.set_font("Arial", "", 10)
        pdf.cell(0, 6, limpar_texto_pdf(", ".join(comp_fam)), 0, 1)
        pdf.ln(2)

    # ======================================================================
    # 2. DIAGNÓSTICO E CONTEXTO CLÍNICO
    # ======================================================================
    pdf.add_page()
    pdf.section_title("2. DIAGNÓSTICO E CONTEXTO CLÍNICO")
    
    pdf.set_font("Arial", "B", 10)
    pdf.cell(0, 6, "Diagnóstico:", 0, 1)
    pdf.set_font("Arial", "", 10)
    diagnostico_limpo = limpar_texto_pdf(dados.get("diagnostico", ""))
    pdf.multi_cell(0, 6, diagnostico_limpo)
    pdf.ln(3)

    # Medicações
    meds = dados.get("lista_medicamentos", [])
    if meds:
        pdf.set_font("Arial", "B", 10)
        pdf.cell(0, 6, "Medicações em Uso:", 0, 1)
        pdf.set_font("Arial", "", 10)
        for med in meds:
            nome = med.get("nome", "")
            posologia = med.get("posologia", "")
            escola = med.get("escola", False)
            escola_txt = " (Administração na escola)" if escola else ""
            texto_med = limpar_texto_pdf(f"- {nome} - {posologia}{escola_txt}")
            pdf.multi_cell(0, 6, texto_med)
        pdf.ln(2)

    # ======================================================================
    # 3. POTENCIALIDADES E HIPERFOCO
    # ======================================================================
    potencias = dados.get("potencias", [])
    hiperfoco = dados.get("hiperfoco", "")
    
    if potencias or hiperfoco:
        pdf.section_title("3. POTENCIALIDADES E INTERESSES")
        
        if hiperfoco:
            pdf.set_font("Arial", "B", 10)
            pdf.cell(0, 6, "Hiperfoco/Interesse Principal:", 0, 1)
            pdf.set_font("Arial", "", 10)
            hiperfoco_limpo = limpar_texto_pdf(hiperfoco)
            pdf.multi_cell(0, 6, hiperfoco_limpo)
            pdf.ln(2)
        
        if potencias:
            pdf.set_font("Arial", "B", 10)
            pdf.cell(0, 6, "Potencialidades:", 0, 1)
            pdf.set_font("Arial", "", 10)
            for pot in potencias:
                pdf.add_flat_icon_item(limpar_texto_pdf(pot), "dot")
            pdf.ln(2)

    # ======================================================================
    # 4. REDE DE APOIO
    # ======================================================================
    rede = dados.get("rede_apoio", [])
    orientacoes = dados.get("orientacoes_especialistas", "")
    orientacoes_por_prof = dados.get("orientacoes_por_profissional", {})
    
    if rede or orientacoes or orientacoes_por_prof:
        pdf.add_page()
        pdf.section_title("4. REDE DE APOIO E ORIENTAÇÕES")
        
        if rede:
            pdf.set_font("Arial", "B", 10)
            pdf.cell(0, 6, "Profissionais da Rede:", 0, 1)
            pdf.set_font("Arial", "", 10)
            for prof in rede:
                pdf.add_flat_icon_item(limpar_texto_pdf(prof), "dot")
            pdf.ln(2)
        
        if orientacoes:
            pdf.set_font("Arial", "B", 10)
            pdf.cell(0, 6, "Orientações Gerais dos Especialistas:", 0, 1)
            pdf.set_font("Arial", "", 10)
            orientacoes_limpo = limpar_texto_pdf(orientacoes)
            pdf.multi_cell(0, 6, orientacoes_limpo)
            pdf.ln(2)
        
        if orientacoes_por_prof and isinstance(orientacoes_por_prof, dict):
            for prof, orient in orientacoes_por_prof.items():
                if orient:
                    pdf.set_font("Arial", "B", 10)
                    titulo_orient = limpar_texto_pdf(f"Orientações - {prof}:")
                    pdf.cell(0, 6, titulo_orient, 0, 1)
                    pdf.set_font("Arial", "", 10)
                    orient_limpo = limpar_texto_pdf(orient)
                    pdf.multi_cell(0, 6, orient_limpo)
                    pdf.ln(2)

    # ======================================================================
    # 5. MAPEAMENTO DE BARREIRAS E NÍVEIS DE SUPORTE
    # ======================================================================
    barreiras = dados.get("barreiras_selecionadas", {})
    if any(barreiras.values() if isinstance(barreiras, dict) else []):
        pdf.add_page()
        pdf.section_title("5. MAPEAMENTO DE BARREIRAS E NÍVEIS DE SUPORTE")
        
        for area, itens in barreiras.items():
            if itens:
                pdf.set_font("Arial", "B", 11)
                pdf.set_text_color(0, 51, 102)
                area_limpo = limpar_texto_pdf(area)
                pdf.cell(0, 8, area_limpo, 0, 1)
                pdf.set_text_color(0, 0, 0)
                pdf.set_font("Arial", "", 10)
                for item in itens:
                    nivel = (dados.get("niveis_suporte") or {}).get(f"{area}_{item}", "Monitorado")
                    item_texto = limpar_texto_pdf(f"{item} (Nível de Suporte: {nivel})")
                    pdf.add_flat_icon_item(item_texto, "check")
                pdf.ln(2)

    # ======================================================================
    # 6. PLANO DE AÇÃO - ESTRATÉGIAS
    # ======================================================================
    estrategias_acesso = dados.get("estrategias_acesso", [])
    estrategias_ensino = dados.get("estrategias_ensino", [])
    estrategias_avaliacao = dados.get("estrategias_avaliacao", [])
    outros_acesso = dados.get("outros_acesso", "")
    outros_ensino = dados.get("outros_ensino", "")
    
    if estrategias_acesso or estrategias_ensino or estrategias_avaliacao or outros_acesso or outros_ensino:
        pdf.add_page()
        pdf.section_title("6. PLANO DE AÇÃO - ESTRATÉGIAS")
        
        if estrategias_acesso:
            pdf.set_font("Arial", "B", 10)
            pdf.cell(0, 6, "Estratégias de Acesso:", 0, 1)
            pdf.set_font("Arial", "", 10)
            for est in estrategias_acesso:
                pdf.add_flat_icon_item(limpar_texto_pdf(est), "dot")
            pdf.ln(2)
        
        if outros_acesso:
            pdf.set_font("Arial", "", 10)
            pdf.multi_cell(0, 6, limpar_texto_pdf(outros_acesso))
            pdf.ln(2)
        
        if estrategias_ensino:
            pdf.set_font("Arial", "B", 10)
            pdf.cell(0, 6, "Estratégias de Ensino:", 0, 1)
            pdf.set_font("Arial", "", 10)
            for est in estrategias_ensino:
                pdf.add_flat_icon_item(limpar_texto_pdf(est), "dot")
            pdf.ln(2)
        
        if outros_ensino:
            pdf.set_font("Arial", "", 10)
            pdf.multi_cell(0, 6, limpar_texto_pdf(outros_ensino))
            pdf.ln(2)
        
        if estrategias_avaliacao:
            pdf.set_font("Arial", "B", 10)
            pdf.cell(0, 6, "Estratégias de Avaliação:", 0, 1)
            pdf.set_font("Arial", "", 10)
            for est in estrategias_avaliacao:
                pdf.add_flat_icon_item(limpar_texto_pdf(est), "dot")
            pdf.ln(2)

    # ======================================================================
    # 7. PLANEJAMENTO PEDAGÓGICO DETALHADO (IA)
    # ======================================================================
    if dados.get("ia_sugestao"):
        pdf.add_page()
        pdf.section_title("7. PLANEJAMENTO PEDAGÓGICO DETALHADO")
        _em = {"red": ou.AI_RED, "blue": ou.AI_BLUE, "green": ou.AI_GREEN, "yellow": ou.AI_YELLOW, "orange": ou.AI_ORANGE}
        _eng = dados.get("consultoria_engine", "red")
        motor_pdf = _em.get(_eng, ou.AI_RED)
        pdf.set_font("Arial", "I", 9)
        pdf.set_text_color(100, 100, 100)
        pdf.cell(0, 6, f"Gerado por {motor_pdf}. Input: série, diagnóstico e barreiras mapeadas; cruzei com BNCC + DUA.", 0, 1, "L")
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Arial", "", 10)
        pdf.ln(4)
        texto_limpo = limpar_texto_pdf(dados["ia_sugestao"])
        texto_limpo = re.sub(r"\[.*?\]", "", texto_limpo)

        for linha in texto_limpo.split("\n"):
            l = linha.strip()
            if not l:
                continue
            if l.startswith("###") or l.startswith("##"):
                pdf.ln(5)
                pdf.set_font("Arial", "B", 12)
                pdf.set_text_color(0, 51, 102)
                titulo_limpo = limpar_texto_pdf(l.replace("#", "").strip())
                pdf.cell(0, 8, titulo_limpo, 0, 1, "L")
                pdf.set_font("Arial", "", 10)
                pdf.set_text_color(0, 0, 0)
            elif l.startswith("-") or l.startswith("*"):
                texto_limpo = limpar_texto_pdf(l.replace("-", "").replace("*", "").strip())
                pdf.add_flat_icon_item(texto_limpo, "dot")
            else:
                pdf.multi_cell(0, 6, limpar_texto_pdf(l))

    # ======================================================================
    # 8. MONITORAMENTO E ACOMPANHAMENTO
    # ======================================================================
    monitoramento_data = dados.get("monitoramento_data")
    status_meta = dados.get("status_meta", "")
    parecer_geral = dados.get("parecer_geral", "")
    proximos_passos = dados.get("proximos_passos_select", [])
    
    if monitoramento_data or status_meta or parecer_geral or proximos_passos:
        pdf.add_page()
        pdf.section_title("8. MONITORAMENTO E ACOMPANHAMENTO")
        
        if monitoramento_data:
            pdf.set_font("Arial", "B", 10)
            pdf.cell(50, 6, "Data do Monitoramento:", 0, 0)
            pdf.set_font("Arial", "", 10)
            if isinstance(monitoramento_data, str):
                pdf.cell(0, 6, monitoramento_data, 0, 1)
            elif hasattr(monitoramento_data, 'strftime'):
                pdf.cell(0, 6, monitoramento_data.strftime("%d/%m/%Y"), 0, 1)
            else:
                pdf.cell(0, 6, str(monitoramento_data), 0, 1)
            pdf.ln(2)
        
        if status_meta:
            pdf.set_font("Arial", "B", 10)
            pdf.cell(50, 6, "Status da Meta:", 0, 0)
            pdf.set_font("Arial", "", 10)
            pdf.cell(0, 6, limpar_texto_pdf(status_meta), 0, 1)
            pdf.ln(2)
        
        if parecer_geral:
            pdf.set_font("Arial", "B", 10)
            pdf.cell(0, 6, "Parecer Geral:", 0, 1)
            pdf.set_font("Arial", "", 10)
            parecer_limpo = limpar_texto_pdf(parecer_geral)
            pdf.multi_cell(0, 6, parecer_limpo)
            pdf.ln(2)
        
        if proximos_passos:
            pdf.set_font("Arial", "B", 10)
            pdf.cell(0, 6, "Próximos Passos:", 0, 1)
            pdf.set_font("Arial", "", 10)
            for passo in proximos_passos:
                pdf.add_flat_icon_item(limpar_texto_pdf(passo), "dot")

    return pdf.output(dest="S").encode("latin-1", "replace")

def gerar_pdf_tabuleiro_simples(texto: str):
    pdf = PDF_Simple_Text()
    pdf.add_page()
    pdf.set_font("Arial", size=11)
    for linha in limpar_texto_pdf(texto).split("\n"):
        l = linha.strip()
        if not l:
            continue
        if l.isupper() or "**" in linha:
            pdf.ln(4)
            pdf.set_font("Arial", "B", 11)
            pdf.set_fill_color(240, 240, 240)
            pdf.cell(0, 8, l.replace("**", ""), 0, 1, "L", fill=True)
            pdf.set_font("Arial", "", 11)
        else:
            pdf.multi_cell(0, 6, l)
    return pdf.output(dest="S").encode("latin-1", "ignore")

def gerar_docx_final(dados: dict):
    """Gera documento Word completo do PEI com todas as seções"""
    doc = Document()
    doc.add_heading("PLANO EDUCACIONAL INDIVIDUALIZADO (PEI)", 0)
    doc.add_heading((dados.get("nome") or "Sem Nome"), level=1)
    doc.add_paragraph("")
    
    # 1. IDENTIFICAÇÃO
    doc.add_heading("1. IDENTIFICAÇÃO E CONTEXTO", level=2)
    p = doc.add_paragraph()
    p.add_run("Estudante: ").bold = True
    p.add_run(dados.get("nome", ""))
    
    p = doc.add_paragraph()
    p.add_run("Série/Ano: ").bold = True
    p.add_run(dados.get("serie", ""))
    
    p = doc.add_paragraph()
    p.add_run("Turma: ").bold = True
    p.add_run(dados.get("turma", ""))
    
    p = doc.add_paragraph()
    p.add_run("Matrícula/RA: ").bold = True
    p.add_run(dados.get("matricula", dados.get("ra", "")))
    
    if dados.get("historico"):
        p = doc.add_paragraph()
        p.add_run("Histórico Escolar: ").bold = True
        doc.add_paragraph(dados.get("historico", ""))
    
    if dados.get("familia"):
        p = doc.add_paragraph()
        p.add_run("Dinâmica Familiar: ").bold = True
        doc.add_paragraph(dados.get("familia", ""))
    
    # 2. DIAGNÓSTICO
    doc.add_heading("2. DIAGNÓSTICO E CONTEXTO CLÍNICO", level=2)
    if dados.get("diagnostico"):
        doc.add_paragraph(dados.get("diagnostico", ""))
    
    meds = dados.get("lista_medicamentos", [])
    if meds:
        p = doc.add_paragraph()
        p.add_run("Medicações em Uso:").bold = True
        for med in meds:
            escola_txt = " (Administração na escola)" if med.get("escola") else ""
            texto_med = limpar_texto_pdf(f"{med.get('nome', '')} - {med.get('posologia', '')}{escola_txt}")
            doc.add_paragraph(texto_med, style='List Bullet')
    
    # 3. POTENCIALIDADES
    if dados.get("hiperfoco") or dados.get("potencias"):
        doc.add_heading("3. POTENCIALIDADES E INTERESSES", level=2)
        if dados.get("hiperfoco"):
            p = doc.add_paragraph()
            p.add_run("Hiperfoco/Interesse Principal: ").bold = True
            doc.add_paragraph(dados.get("hiperfoco", ""))
        potencias = dados.get("potencias", [])
        if potencias:
            p = doc.add_paragraph()
            p.add_run("Potencialidades:").bold = True
            for pot in potencias:
                doc.add_paragraph(pot, style='List Bullet')
    
    # 4. REDE DE APOIO
    rede = dados.get("rede_apoio", [])
    orientacoes = dados.get("orientacoes_especialistas", "")
    if rede or orientacoes:
        doc.add_heading("4. REDE DE APOIO E ORIENTAÇÕES", level=2)
        if rede:
            p = doc.add_paragraph()
            p.add_run("Profissionais da Rede:").bold = True
            for prof in rede:
                doc.add_paragraph(prof, style='List Bullet')
        if orientacoes:
            p = doc.add_paragraph()
            p.add_run("Orientações dos Especialistas:").bold = True
            doc.add_paragraph(orientacoes)
    
    # 5. BARREIRAS
    barreiras = dados.get("barreiras_selecionadas", {})
    if any(barreiras.values() if isinstance(barreiras, dict) else []):
        doc.add_heading("5. MAPEAMENTO DE BARREIRAS E NÍVEIS DE SUPORTE", level=2)
        for area, itens in barreiras.items():
            if itens:
                doc.add_heading(area, level=3)
                for item in itens:
                    nivel = (dados.get("niveis_suporte") or {}).get(f"{area}_{item}", "Monitorado")
                    doc.add_paragraph(f"{item} (Nível de Suporte: {nivel})", style='List Bullet')
    
    # 6. ESTRATÉGIAS
    estrategias_acesso = dados.get("estrategias_acesso", [])
    estrategias_ensino = dados.get("estrategias_ensino", [])
    estrategias_avaliacao = dados.get("estrategias_avaliacao", [])
    if estrategias_acesso or estrategias_ensino or estrategias_avaliacao:
        doc.add_heading("6. PLANO DE AÇÃO - ESTRATÉGIAS", level=2)
        if estrategias_acesso:
            p = doc.add_paragraph()
            p.add_run("Estratégias de Acesso:").bold = True
            for est in estrategias_acesso:
                doc.add_paragraph(est, style='List Bullet')
        if estrategias_ensino:
            p = doc.add_paragraph()
            p.add_run("Estratégias de Ensino:").bold = True
            for est in estrategias_ensino:
                doc.add_paragraph(est, style='List Bullet')
        if estrategias_avaliacao:
            p = doc.add_paragraph()
            p.add_run("Estratégias de Avaliação:").bold = True
            for est in estrategias_avaliacao:
                doc.add_paragraph(est, style='List Bullet')
    
    # 7. PLANEJAMENTO PEDAGÓGICO
    if dados.get("ia_sugestao"):
        doc.add_heading("7. PLANEJAMENTO PEDAGÓGICO DETALHADO", level=2)
        _em = {"red": ou.AI_RED, "blue": ou.AI_BLUE, "green": ou.AI_GREEN, "yellow": ou.AI_YELLOW, "orange": ou.AI_ORANGE}
        _eng = dados.get("consultoria_engine", "red")
        motor_docx = _em.get(_eng, ou.AI_RED)
        p_motor = doc.add_paragraph()
        p_motor.add_run(f"Gerado por {motor_docx}. ").italic = True
        p_motor.add_run("Input: série, diagnóstico e barreiras mapeadas; cruzei com BNCC + DUA.")
        doc.add_paragraph()
        texto_limpo = re.sub(r"\[.*?\]", "", dados["ia_sugestao"])
        for linha in texto_limpo.split("\n"):
            l = linha.strip()
            if not l:
                continue
            if l.startswith("###"):
                doc.add_heading(l.replace("###", "").strip(), level=3)
            elif l.startswith("##"):
                doc.add_heading(l.replace("##", "").strip(), level=2)
            elif l.startswith("-") or l.startswith("*"):
                doc.add_paragraph(l.replace("-", "").replace("*", "").strip(), style='List Bullet')
            else:
                doc.add_paragraph(l)
    
    # 8. MONITORAMENTO
    if dados.get("monitoramento_data") or dados.get("status_meta") or dados.get("parecer_geral"):
        doc.add_heading("8. MONITORAMENTO E ACOMPANHAMENTO", level=2)
        if dados.get("monitoramento_data"):
            p = doc.add_paragraph()
            p.add_run("Data do Monitoramento: ").bold = True
            data_mon = dados.get("monitoramento_data")
            if hasattr(data_mon, 'strftime'):
                p.add_run(data_mon.strftime("%d/%m/%Y"))
            else:
                p.add_run(str(data_mon))
        if dados.get("status_meta"):
            p = doc.add_paragraph()
            p.add_run("Status da Meta: ").bold = True
            p.add_run(dados.get("status_meta", ""))
        if dados.get("parecer_geral"):
            p = doc.add_paragraph()
            p.add_run("Parecer Geral: ").bold = True
            doc.add_paragraph(dados.get("parecer_geral", ""))
    
    b = BytesIO()
    doc.save(b)
    b.seek(0)
    return b


# ==============================================================================
# 7D. IA (Extração PDF + Consultoria + Gamificação)
# ==============================================================================

def ler_pdf(uploaded_file, max_pages=6):
    """
    Lê um arquivo PDF e extrai o texto (até max_pages páginas).
    Retorna string vazia se houver erro.
    """
    if not uploaded_file:
        return ""
    
    try:
        pdf_reader = PdfReader(uploaded_file)
        texto_completo = []
        
        # Limitar número de páginas para não sobrecarregar
        num_pages = min(len(pdf_reader.pages), max_pages)
        
        for i in range(num_pages):
            try:
                page = pdf_reader.pages[i]
                texto = page.extract_text()
                if texto:
                    texto_completo.append(texto)
            except Exception as e:
                # Se uma página falhar, continua com as próximas
                continue
        
        return "\n\n".join(texto_completo) if texto_completo else ""
    except Exception as e:
        # Retorna string vazia em caso de erro
        return ""

def extrair_dados_pdf_ia(api_key: str, texto_pdf: str):
    if not api_key:
        return None, f"Configure a chave da IA ({ou.AI_RED}) nas configurações."
    try:
        client = OpenAI(api_key=api_key)
        prompt = (
            "Analise este laudo médico/escolar. Extraia: 1) Diagnóstico; 2) Medicamentos. "
            'Responda em JSON no formato: { "diagnostico": "...", "medicamentos": [ {"nome": "...", "posologia": "..."} ] }. '
            f"Texto: {texto_pdf[:4000]}"
        )
        res = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
        )
        return json.loads(res.choices[0].message.content), None
    except Exception as e:
        return None, str(e)

def consultar_gpt_pedagogico(api_key: str, dados: dict, contexto_pdf: str = "", modo_pratico: bool = False, feedback_usuario: str = "", engine: str = "red"):
    """
    Gera relatório da Consultoria IA.
    engine: red (DeepSeek), blue (Kimi), green (Claude), yellow (Gemini), orange (GPT fallback)
    """
    engine = (engine or "red").strip().lower()
    if engine not in ("red", "blue", "green", "yellow", "orange"):
        engine = "red"

    if engine == "red" and not ou.get_deepseek_api_key():
        return None, f"⚠️ Configure DEEPSEEK_API_KEY ({ou.AI_RED}) nas configurações."
    if engine == "blue":
        kimi_key = ou.get_kimi_api_key() or (st.session_state.get("_kimi_key_cached") if hasattr(st, "session_state") else None)
        try:
            kimi_key = kimi_key or (st.secrets.get("OPENROUTER_API_KEY", "") or st.secrets.get("KIMI_API_KEY", "") or "").strip() or None
        except Exception:
            pass
        if not kimi_key:
            return None, f"⚠️ Configure OPENROUTER_API_KEY ou KIMI_API_KEY ({ou.AI_BLUE}) nas configurações."
    if engine == "green" and not ou.get_anthropic_api_key():
        return None, f"⚠️ Configure ANTHROPIC_API_KEY ({ou.AI_GREEN}) nas configurações."
    if engine == "yellow" and not ou.get_gemini_api_key():
        return None, f"⚠️ Configure GEMINI_API_KEY ({ou.AI_YELLOW}) nas configurações."
    if engine == "orange" and not (ou.get_openai_api_key() or api_key):
        return None, f"⚠️ Configure OPENAI_API_KEY ({ou.AI_ORANGE}) para usar fallback."

    try:

        evid = "\n".join([f"- {k.replace('?', '')}" for k, v in (dados.get("checklist_evidencias") or {}).items() if v])
        meds_info = "\n".join(
            [f"- {m.get('nome','')} ({m.get('posologia','')})." for m in (dados.get("lista_medicamentos") or [])]
        ) if dados.get("lista_medicamentos") else "Nenhuma medicação informada."

        hiperfoco_txt = f"HIPERFOCO DO ESTUDANTE: {dados.get('hiperfoco','')}" if dados.get("hiperfoco") else "Hiperfoco: Não identificado."
        serie = dados.get("serie") or ""
        nivel_ensino = detectar_nivel_ensino(serie)
        alfabetizacao = dados.get("nivel_alfabetizacao", "Não Avaliado")

        prompt_identidade = f"""
[PERFIL_NARRATIVO]
Inicie com "👤 QUEM É O ESTUDANTE?". Crie um parágrafo humanizado. {hiperfoco_txt}.
Use o hiperfoco para conectar com a aprendizagem.
[/PERFIL_NARRATIVO]
""".strip()

        prompt_diagnostico = """
### 1. 🏥 DIAGNÓSTICO E IMPACTO (FUNDAMENTAL):
- Cite o Diagnóstico (e o CID se disponível).
- Descreva os impactos diretos na aprendizagem.
- Liste cuidados/pontos de atenção.
""".strip()

        prompt_literacia = ""
        if "Alfabético" not in alfabetizacao and alfabetizacao != "Não se aplica (Educação Infantil)":
            prompt_literacia = f"""[ATENÇÃO CRÍTICA: ALFABETIZAÇÃO] Fase: {alfabetizacao}. Inclua 2 ações de consciência fonológica.[/ATENÇÃO CRÍTICA]"""

        prompt_hub = """
### 7. 🧩 CHECKLIST DE ADAPTAÇÃO E ACESSIBILIDADE:
**A. Mediação (Triângulo de Ouro):**
1) Instruções passo a passo
2) Fragmentação de tarefas
3) Scaffolding

**B. Acessibilidade:**
4) Inferências/figuras de linguagem
5) Descrição de imagens (alt text)
6) Adaptação visual (fonte/espaçamento)
7) Adequação de desafio
""".strip()

        # Não há mais seção separada "Componentes de atenção" — fica só o MAPA POR COMPONENTE (único)
        prompt_metas = """
[METAS_SMART — TÓPICO OBRIGATÓRIO E PRIORITÁRIO]
Dedique ESPECIAL ATENÇÃO a este bloco. Metas SMART são fundamentais no PEI: Específicas, Mensuráveis, Atingíveis, Relevantes e Temporais.
Para cada meta, escreva em 3 a 5 linhas: o quê exatamente será trabalhado, como será medido (indicadores concretos), prazo claro, e relação com o diagnóstico/barreiras do estudante.
- **Meta de Curto Prazo (2 meses):** Objetivo específico, indicador de sucesso, quem acompanha, data de revisão.
- **Meta de Médio Prazo (1 semestre):** Objetivo específico, indicador de sucesso, quem acompanha, data de revisão.
- **Meta de Longo Prazo (1 ano):** Objetivo específico, indicador de sucesso, quem acompanha, data de revisão.
Não use placeholders genéricos. Personalize cada meta ao estudante e ao contexto pedagógico.
[/METAS_SMART]
""".strip()

        if nivel_ensino == "EI":
            perfil_ia = "Especialista em EDUCAÇÃO INFANTIL e BNCC."
            bncc_ei_bloco = ""
            idade_ei = dados.get("bncc_ei_idade") or ""
            campo_ei = dados.get("bncc_ei_campo") or ""
            obj_ei = dados.get("bncc_ei_objetivos") or []
            if idade_ei or campo_ei or obj_ei:
                obj_txt = "\n".join(f"- {o}" for o in obj_ei) if obj_ei else "(não selecionados)"
                bncc_ei_bloco = f"""[MAPEAMENTO_BNCC_EI — REGRA OBRIGATÓRIA]
Na seção "Avaliação de Repertório", use APENAS os dados abaixo (Campos de Experiência e Objetivos de Aprendizagem BNCC EI). Não invente outros.
- Faixa de idade: {idade_ei or "não informada"}
- Campo de Experiência: {campo_ei or "não informado"}
- Objetivos de Aprendizagem (cite EXATAMENTE como abaixo, não parafraseie):
{obj_txt}
[/MAPEAMENTO_BNCC_EI]"""
            else:
                bncc_ei_bloco = "[CAMPOS_EXPERIENCIA_PRIORITARIOS] Destaque 2 ou 3 Campos BNCC da Educação Infantil. Se a escola preencheu a aba BNCC, use aqueles dados. [/CAMPOS_EXPERIENCIA_PRIORITARIOS]"
            estrutura_req = f"""
{prompt_identidade}
{prompt_diagnostico}

### 2. 🌟 AVALIAÇÃO DE REPERTÓRIO:
{bncc_ei_bloco}

### 3. 🚀 ESTRATÉGIAS DE INTERVENÇÃO:
(Estratégias de acolhimento, rotina e adaptação sensorial).

### 4. 🎯 METAS SMART (tópico independente — dedique bastante espaço e detalhe):
{prompt_metas}

### 5. ⚠️ PONTOS DE ATENÇÃO FARMACOLÓGICA:
[ANALISE_FARMA] Se houver medicação, cite efeitos colaterais para atenção pedagógica. [/ANALISE_FARMA]

{prompt_hub}
""".strip()
        else:
            perfil_ia = "Especialista em Inclusão Escolar e BNCC."
            # Prioridade: habilidades VALIDADAS pelo professor; senão selecionadas; senão carregar do CSV
            hab_selecionadas = dados.get("habilidades_bncc_validadas") or dados.get("habilidades_bncc_selecionadas") or []
            if hab_selecionadas and isinstance(hab_selecionadas, list):
                texto_hab_sel = "\n".join(
                    f"- {h.get('disciplina','')}: {h.get('codigo','')} — {h.get('habilidade_completa', h.get('descricao',''))}"
                    for h in hab_selecionadas if isinstance(h, dict)
                )
                if texto_hab_sel.strip():
                    instrucao_bncc = f"""[MAPEAMENTO_BNCC — REGRA OBRIGATÓRIA]
Na seção "Avaliação de Repertório", cite SOMENTE habilidades da lista abaixo. A lista é FECHADA e EXAUSTIVA: NÃO adicione nenhuma habilidade que não esteja nela. Ao citar cada habilidade, reproduza a linha INTEIRA EXATAMENTE como está na lista: código + descrição COMPLETA, palavra por palavra, sem parafrasear, resumir ou alterar. Proibido inventar ou modificar o texto.
[HABILIDADES_VALIDADAS_PELO_PROFESSOR — USE SOMENTE ESTAS, NA ÍNTEGRA]
{texto_hab_sel}
[/HABILIDADES_VALIDADAS_PELO_PROFESSOR]
[/MAPEAMENTO_BNCC]"""
                else:
                    hab_selecionadas = []
            if not (hab_selecionadas and isinstance(hab_selecionadas, list)):
                bncc_blocos = carregar_habilidades_bncc_ano_atual_e_anteriores(serie)
                ano_atual_txt = (bncc_blocos.get("ano_atual") or "").strip()
                anos_anteriores_txt = (bncc_blocos.get("anos_anteriores") or "").strip()
                if ano_atual_txt or anos_anteriores_txt:
                    instrucao_bncc = """[MAPEAMENTO_BNCC — REGRA OBRIGATÓRIA]
Na seção "Avaliação de Repertório", cite SOMENTE habilidades das listas abaixo. Ao citar cada habilidade, reproduza a linha INTEIRA EXATAMENTE como está na lista: código + descrição COMPLETA, palavra por palavra, sem parafrasear, resumir ou alterar. O código correto sozinho NÃO basta: a descrição deve ser idêntica à da lista. Proibido inventar ou modificar o texto. Separe por Componente Curricular."""
                    if anos_anteriores_txt:
                        instrucao_bncc += f"""
[HABILIDADES_BNCC_ANOS_ANTERIORES]
{anos_anteriores_txt}
[/HABILIDADES_BNCC_ANOS_ANTERIORES]"""
                    if ano_atual_txt:
                        instrucao_bncc += f"""
[HABILIDADES_BNCC_ANO_SERIE_ATUAL]
{ano_atual_txt}
[/HABILIDADES_BNCC_ANO_SERIE_ATUAL]"""
                    instrucao_bncc += "\n[/MAPEAMENTO_BNCC]"
                else:
                    instrucao_bncc = "[MAPEAMENTO_BNCC] Separe por Componente Curricular. Inclua código alfanumérico (ex: EF01LP02) e a descrição por escrito. [/MAPEAMENTO_BNCC]"
            instrucao_bloom = "[TAXONOMIA_BLOOM] Explique a categoria cognitiva escolhida. [/TAXONOMIA_BLOOM]"
            instrucao_mapa = """
### 4. 📊 COMPONENTES CURRICULARES QUE MERECEM ATENÇÃO (mapa único):
Construa um quadro claro que demonstre QUAIS componentes curriculares precisam de MAIOR atenção e POR QUÊ. Este é o único bloco sobre componentes — não repita em outra seção.

**Formato obrigatório:** para cada componente que tenha habilidades listadas na Avaliação de Repertório, inclua:
1. **Componente** (ex.: Língua Portuguesa, Matemática, Ciências)
2. **Nível de atenção:** Alta | Média | Monitoramento
3. **Motivos:** cruze explicitamente (a) o diagnóstico e as barreiras do PEI com (b) as habilidades desse componente. Exemplo: "O diagnóstico de X e as barreiras em [área] impactam diretamente as habilidades [códigos ou temas] deste componente, exigindo adaptações em..."

**Regra:** use APENAS os componentes que aparecem nas habilidades já listadas neste relatório (Avaliação de Repertório). Para cada um, justifique o nível de atenção ligando diagnóstico + barreiras + evidências do PEI às habilidades daquele componente. Seja objetivo."""
            prompt_metas_fundamental = """
### 5. 🎯 METAS SMART (tópico independente — PRIORIDADE ALTA, dedique bastante energia aqui):
[METAS_SMART]
Metas SMART são fundamentais no PEI: Específicas, Mensuráveis, Atingíveis, Relevantes e Temporais.
Para CADA meta, escreva em 3 a 5 linhas: o quê exatamente será trabalhado, como será medido (indicadores concretos), prazo claro, e relação com o diagnóstico/barreiras do estudante.
- **Meta de Curto Prazo (2 meses):** Objetivo específico, indicador de sucesso, quem acompanha, data de revisão.
- **Meta de Médio Prazo (1 semestre):** Objetivo específico, indicador de sucesso, quem acompanha, data de revisão.
- **Meta de Longo Prazo (1 ano):** Objetivo específico, indicador de sucesso, quem acompanha, data de revisão.
NÃO use placeholders genéricos. Personalize cada meta ao estudante e ao contexto pedagógico. Este bloco deve ser um dos mais detalhados do relatório.
[/METAS_SMART]
""".strip()
            estrutura_req = f"""
{prompt_identidade}
{prompt_diagnostico}

### 2. 🌟 AVALIAÇÃO DE REPERTÓRIO:
- Habilidades de anos anteriores que merecem atenção ou são essenciais para este estudante (pontual, com código e descrição).
- Habilidades do ano/série atual que são fundamentais para este estudante (pontual, com código e descrição).
{instrucao_bncc}
{instrucao_bloom}

### 3. 🚀 ESTRATÉGIAS DE INTERVENÇÃO:
(Adaptações curriculares e de acesso).
{prompt_literacia}

{instrucao_mapa}

{prompt_metas_fundamental}

### 6. ⚠️ PONTOS DE ATENÇÃO FARMACOLÓGICA:
[ANALISE_FARMA] Se houver medicação, cite efeitos colaterais para atenção pedagógica. [/ANALISE_FARMA]

{prompt_hub}
""".strip()

        prompt_feedback = f"AJUSTE SOLICITADO: {feedback_usuario}" if feedback_usuario else ""
        prompt_formatacao = "IMPORTANTE: Use Markdown simples. Use títulos H3 (###). Evite tabelas."
        regra_repertorio = ""
        if nivel_ensino != "EI":
            regra_repertorio = "REGRA CRÍTICA (Avaliação de Repertório): Cite SOMENTE habilidades da lista fornecida (MAPEAMENTO_BNCC). Ao citar, reproduza cada habilidade EXATAMENTE como na lista: código e descrição COMPLETA, palavra por palavra. Proibido parafrasear, resumir ou alterar o texto da descrição; deve ser idêntico à lista.\n\n"

        prompt_sys = f"""{perfil_ia}
MISSÃO: Criar PEI Técnico Oficial.
{regra_repertorio}
ESTRUTURA OBRIGATÓRIA:
{estrutura_req}

{prompt_feedback}
{prompt_formatacao}
""".strip()

        if modo_pratico:
            prompt_sys = f"""{perfil_ia}
GUIA PRÁTICO PARA SALA DE AULA.
{prompt_feedback}

{prompt_hub}
""".strip()

        prompt_user = (
            f"ALUNO: {dados.get('nome','')} | SÉRIE: {serie} | HISTÓRICO: {dados.get('historico','')} | "
            f"DIAGNÓSTICO: {dados.get('diagnostico','')} | MEDS: {meds_info} | "
            f"EVIDÊNCIAS: {evid} | LAUDO: {(contexto_pdf or '')[:3000]}"
        )
        # Injetar lista de habilidades no user message para a IA NÃO inventar (copiar exatamente)
        if nivel_ensino != "EI":
            hab_para_user = dados.get("habilidades_bncc_validadas") or dados.get("habilidades_bncc_selecionadas") or []
            if hab_para_user and isinstance(hab_para_user, list):
                lista_user = "\n".join(
                    f"- {h.get('disciplina','')}: {h.get('codigo','')} — {h.get('habilidade_completa', h.get('descricao',''))}"
                    for h in hab_para_user if isinstance(h, dict)
                )
                if lista_user.strip():
                    prompt_user += f"\n\n[LISTA ÚNICA DE HABILIDADES PERMITIDAS — na Avaliação de Repertório cite SOMENTE estas, COPIANDO a linha inteira EXATAMENTE como abaixo. NÃO invente outras:]\n{lista_user}"
            else:
                bncc_u = carregar_habilidades_bncc_ano_atual_e_anteriores(serie)
                at_u = (bncc_u.get("ano_atual") or "").strip()
                ant_u = (bncc_u.get("anos_anteriores") or "").strip()
                if at_u or ant_u:
                    prompt_user += "\n\n[LISTA ÚNICA DE HABILIDADES PERMITIDAS — na Avaliação de Repertório cite SOMENTE estas, COPIANDO a linha inteira EXATAMENTE como abaixo:]"
                    if ant_u:
                        prompt_user += f"\n\nANOS ANTERIORES:\n{ant_u}"
                    if at_u:
                        prompt_user += f"\n\nANO SÉRIE ATUAL:\n{at_u}"

        messages = [{"role": "system", "content": prompt_sys}, {"role": "user", "content": prompt_user}]
        texto_relatorio = ou.chat_completion_multi_engine(engine, messages, temperature=0.7, api_key=api_key)
        # Pós-processamento: remover da Avaliação de Repertório qualquer habilidade que a IA tenha inventado
        if nivel_ensino != "EI" and texto_relatorio:
            lista_permitida = dados.get("habilidades_bncc_validadas") or dados.get("habilidades_bncc_selecionadas") or []
            if lista_permitida:
                texto_relatorio = filtrar_avaliacao_repertorio_apenas_permitidas(texto_relatorio, lista_permitida)
        return texto_relatorio, None

    except Exception as e:
        return None, str(e)

def gerar_roteiro_gamificado(api_key: str, dados: dict, pei_tecnico: str, feedback_game: str = ""):
    if not api_key:
        return None, f"Configure a chave da IA ({ou.AI_RED})."
    try:
        client = OpenAI(api_key=api_key)
        serie = dados.get("serie") or ""
        nivel_ensino = detectar_nivel_ensino(serie)
        hiperfoco = dados.get("hiperfoco") or "brincadeiras"
        nome_curto = (dados.get("nome","").split() or ["Estudante"])[0]

        contexto_seguro = (
            f"ALUNO: {nome_curto} | HIPERFOCO: {hiperfoco} | "
            f"PONTOS FORTES: {', '.join(dados.get('potencias',[]))}"
        )

        prompt_feedback = f"AJUSTE: {feedback_game}" if feedback_game else ""

        regra_sem_diagnostico = " REGRA: NUNCA inclua diagnóstico clínico, CID ou condições médicas — este material será entregue ao estudante e à família."
        if nivel_ensino == "EI":
            prompt_sys = "Crie uma história visual (4-5 anos) com emojis. Estrutura: começo, desafio, ajuda, conquista, rotina." + regra_sem_diagnostico
        elif nivel_ensino == "FI":
            prompt_sys = "Crie um quadro de missões RPG (6-10 anos). Estrutura: mapa, missões, recompensas, superpoder." + regra_sem_diagnostico
        else:
            prompt_sys = "Crie uma ficha RPG (adolescente). Estrutura: quest, skills, buffs, checklists e metas." + regra_sem_diagnostico

        full_sys = f"{prompt_sys}\n{prompt_feedback}"
        res = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "system", "content": full_sys}, {"role": "user", "content": contexto_seguro}],
        )
        return res.choices[0].message.content, None
    except Exception as e:
        return None, str(e)


# ==============================================================================
# 7E. AÇÕES AUXILIARES (Reset)
# ==============================================================================
def limpar_formulario():
    # recria um "rascunho limpo" preservando a estrutura do dicionário
    st.session_state.dados = {
        'nome': '',
        'nasc': date(2015, 1, 1),
        'serie': None,
        'turma': '',
        'diagnostico': '',
        'lista_medicamentos': [],
        'composicao_familiar_tags': [],
        'historico': '',
        'familia': '',
        'hiperfoco': '',
        'potencias': [],
        'rede_apoio': [],
        'orientacoes_especialistas': '',
        'checklist_evidencias': {},
        'nivel_alfabetizacao': 'Não se aplica (Educação Infantil)',
        'barreiras_selecionadas': {k: [] for k in LISTAS_BARREIRAS.keys()},
        'niveis_suporte': {},
        'estrategias_acesso': [],
        'estrategias_ensino': [],
        'estrategias_avaliacao': [],
        'ia_sugestao': '',
        'consultoria_engine': 'red',
        'ia_mapa_texto': '',
        'outros_acesso': '',
        'outros_ensino': '',
        'monitoramento_data': date.today(),
        'status_meta': 'Não Iniciado',
        'parecer_geral': 'Manter Estratégias',
        'proximos_passos_select': [],
        'status_validacao_pei': 'rascunho',
        'feedback_ajuste': '',
        'status_validacao_game': 'rascunho',
        'feedback_ajuste_game': '',
        'habilidades_bncc_selecionadas': [],
        'habilidades_bncc_validadas': None,
    }
    st.session_state.pdf_text = ""


# ==============================================================================
# 8. ESTILO VISUAL - BLOCO UNIFICADO (FINAL / ESTÁVEL)
# ==============================================================================
st.markdown("""
<style>
/* ===============================
   FONTES E FUNDO GLOBAL
================================ */
@import url('https://fonts.googleapis.com/css2?family=Nunito:wght@400;600;700;800&display=swap');

html, body, [class*="css"] { 
    font-family: 'Nunito', sans-serif; 
    color: #2D3748; 
    background-color: #F7FAFC; 
    margin: 0 !important; 
    padding: 0 !important; 
}

/* ===============================
   HERO CARD — PADRÃO OMNISFERA
================================ */
.mod-card-wrapper {
    margin-bottom: 4px;
    border-radius: 16px;
    overflow: hidden;
    /* margin-top já aplicado no forcar_layout_hub() - não duplicar aqui */
}

.mod-card-rect {
    background: #FFFFFF;
    border-radius: 16px;
    padding: 0;
    border: 1px solid #E2E8F0;
    display: flex;
    flex-direction: row;
    align-items: center;
    height: 130px !important;  /* 🔒 ALTURA FIXA PADRONIZADA */
    width: 100%;
    position: relative;
    overflow: hidden;
    transition: all 0.25s cubic-bezier(0.4, 0, 0.2, 1);
}

.mod-card-rect:hover {
    transform: translateY(-4px) !important;
    box-shadow: 0 12px 24px rgba(0, 0, 0, 0.08) !important;
    border-color: #CBD5E1 !important;
}

.mod-bar {
    width: 6px;
    height: 100%;
    flex-shrink: 0;
}

.mod-icon-area {
    width: 90px;
    height: 100%;
    display: flex;
    align-items: center;
    justify-content: center;
    font-size: 1.8rem;
    flex-shrink: 0;
    background: #FAFAFA !important;
    border-right: 1px solid #F1F5F9;
    transition: all 0.3s ease;
}

.mod-card-rect:hover .mod-icon-area {
    background: white !important;
    transform: scale(1.05) !important;
}

.mod-content {
    flex-grow: 1;
    padding: 0 24px;
    display: flex;
    flex-direction: column;
    justify-content: center;
}

.mod-title {
    font-weight: 800;
    font-size: 1.1rem;
    color: #1E293B;
    margin-bottom: 6px;
    letter-spacing: -0.2px;
}

.mod-desc {
    font-size: 0.8rem;
    color: #64748B;
    line-height: 1.4;
    display: -webkit-box;
    -webkit-line-clamp: 2;
    -webkit-box-orient: vertical;
    overflow: hidden;
    margin: 0;
}

/* ===============================
   CORES (HERDA ACCENT DO PEI)
================================ */
.c-blue {
    background: #0EA5E9 !important;
}

.bg-blue-soft {
    background: #E0F2FE !important;
    color: #0284C7 !important;
}

.mod-icon-area i {
    color: inherit !important;
}

/* Garante que o ícone use a cor do bg-blue-soft */
.bg-blue-soft i {
    color: #0284C7 !important;
}

/* ===============================
   TABS — PADRÃO VIA omni_utils.inject_unified_ui_css()
================================ */
/* Estilos de tabs, botões, selects, etc. são aplicados via função padronizada */

/* ===============================
   FORMULÁRIOS — PADRÃO VIA omni_utils.inject_unified_ui_css()
================================ */
/* Estilos de inputs, selects, etc. são aplicados via função padronizada */

/* ===============================
   FOOTER
================================ */
.footer-signature { 
    text-align: center; 
    opacity: 0.55; 
    font-size: 0.75rem; 
    padding: 30px 0 10px 0; 
}

/* ===============================
   LOGO GIRANDO
================================ */
@keyframes spin-slow {
    from { transform: rotate(0deg); }
    to { transform: rotate(360deg); }
}

.omni-logo-spin {
    animation: spin-slow 10s linear infinite;
}

/* ===============================
   RESPONSIVIDADE
================================ */
@media (max-width: 768px) {
    .mod-card-rect {
        flex-direction: column;
        height: auto;
        padding: 12px;
    }

    .mod-bar {
        width: 100%;
        height: 4px;
    }

    .mod-icon-area {
        width: 100%;
        height: 50px;
        border-right: none;
        border-bottom: 1px solid #F1F5F9;
    }

    .mod-content {
        padding: 12px 0 0 0;
    }

    
}

/* ===============================
AJUSTE ENTRE MENU SUPERIOR E HERO (PADRONIZADO)
================================ */
/* margin-top já aplicado no forcar_layout_hub() - não duplicar aqui */

/* ===============================
   ESCONDER HEADER STREAMLIT
================================ */
section[data-testid="stHeader"] {
    display: none !important;
}
</style>

<link href="https://cdn.jsdelivr.net/npm/remixicon@4.1.0/fonts/remixicon.css" rel="stylesheet">
""", unsafe_allow_html=True)

# ============================================================================== 
# 9. LOGO E PROGRESSO
# ============================================================================== 

def get_logo_base64() -> str | None:
    """Obtém o logo em base64"""
    for c in ["omni_icone.png", "logo.png", "iconeaba.png"]:
        if os.path.exists(c):
            with open(c, "rb") as f:
                return f"data:image/png;base64,{base64.b64encode(f.read()).decode()}"
    return None

src_logo_giratoria = get_logo_base64()

def render_progresso():
    """Renderiza a barra de progresso compacta"""
    p = max(0, min(100, int(calcular_progresso())))
    icon_html = ""
    
    if src_logo_giratoria:
        icon_html = f'<img src="{src_logo_giratoria}" class="omni-logo-spin" style="width:25px;height:25px;">'
    
    bar_color = "linear-gradient(90deg, #FF6B6B 0%, #FF8E53 100%)"
    if p >= 100:
        bar_color = "linear-gradient(90deg, #34D399 0%, #059669 100%)"
    
    st.markdown(f"""
    <div class="progress-container">
        <div style="width:100%; height:3px; background:#E2E8F0; border-radius:2px; position:relative;">
            <div style="height:3px; width:{p}%; background:{bar_color}; border-radius:2px;"></div>
            <div style="position:absolute; top:-14px; left:{p}%; transform:translateX(-50%);">{icon_html}</div>
        </div>
    </div>
    """, unsafe_allow_html=True)


# ==============================================================================
# ABAS DO PEI (TEXTO EM MAIÚSCULAS, SEM EMOJIS)
# ==============================================================================
st.caption("📍 **PEI** — Navegue pelas abas: Início | Estudante | Evidências | Rede de Apoio | Mapeamento | Plano de Ação | Monitoramento | BNCC | Consultoria IA | Dashboard")
abas = [
    "INÍCIO", "ESTUDANTE", "EVIDÊNCIAS", "REDE DE APOIO", "MAPEAMENTO",
    "PLANO DE AÇÃO", "MONITORAMENTO", "BNCC", "CONSULTORIA IA", "DASHBOARD & DOCS"
]

tab0, tab1, tab2, tab3, tab4, tab5, tab6, tab7_hab, tab8, tab9 = st.tabs(abas)

# ==============================================================================
# 11. ABA INÍCIO — CENTRAL (Gestão de Estudantes + Backups)
# ==============================================================================
with tab0:
    st.markdown(f"### {icon_title('Central de Fundamentos e Gestão', 'pei', 24, '#0EA5E9')}", unsafe_allow_html=True)
    st.caption("Aqui você gerencia estudantes (backup local e nuvem/Supabase) e acessa fundamentos do PEI.")

    # -------------------------
    # Helpers locais (somente UI)
    # -------------------------
    def _coerce_dates_in_payload(d: dict):
        """Converte campos de data salvos como string de volta para date (sem depender de Supabase)."""
        if not isinstance(d, dict):
            return d
        for k in ["nasc", "monitoramento_data"]:
            try:
                if k in d and isinstance(d[k], str) and d[k]:
                    d[k] = date.fromisoformat(d[k])
            except Exception:
                pass
        return d

    # -------------------------
    # LAYOUT 2 COLUNAS
    # -------------------------
    col_left, col_right = st.columns([1.15, 0.85])

    # =========================
    # ESQUERDA: Fundamentos
    # =========================
    with col_left:
        with st.container(border=True):
            st.markdown(f"#### {icon_title('Fundamentos do PEI', 'pei', 20, '#0EA5E9')}", unsafe_allow_html=True)
            st.markdown(
                """
- O **PEI** organiza o planejamento individualizado com foco em **barreiras** e **apoios**.
- A lógica é **equidade**: ajustar **acesso, ensino e avaliação**, sem baixar expectativas.
- Base: **LBI (Lei 13.146/2015)**, LDB e diretrizes de Educação Especial na Perspectiva Inclusiva.
                """
            )

        with st.container(border=True):
            st.markdown(f"#### {icon_title('Como usar a Omnisfera', 'info', 20, '#0EA5E9')}", unsafe_allow_html=True)
            st.markdown(
                """
1) **Estudante**: identificação + contexto + laudo (opcional)  
2) **Evidências**: o que foi observado e como aparece na rotina  
3) **Mapeamento**: barreiras + nível de apoio + potências  
4) **Plano de Ação**: acesso/ensino/avaliação  
5) **Consultoria IA**: gerar o documento técnico (validação do educador)  
6) **Dashboard**: KPIs + exportações + sincronização  
                """
            )

        with st.expander("📘 PEI/PDI e a Prática Inclusiva — Amplie o conhecimento", expanded=False):
            st.markdown("""
            O **Plano Educacional Individualizado (PEI)**, também denominado **Plano de Desenvolvimento Individual (PDI)**, é um roteiro de intervenção pedagógica personalizado e flexível que norteia o processo de aprendizagem em sala comum para público-alvo da educação inclusiva. Tem o objetivo de **remover obstáculos** e **promover a escolarização**.

            O PEI/PDI leva em conta as particularidades do(a) aluno(a), incluindo-o no repertório da classe que frequenta e tendo como referência a **mesma matriz curricular** do ano a ser cursado.

            **Caráter obrigatório:** deve ser atualizado sistematicamente e compor a documentação escolar de alunos com deficiência, transtorno global do desenvolvimento e altas habilidades/superdotação. Respeita as orientações do laudo médico, quando houver. Se o aluno apresentar condições descritas e não possuir acompanhamento externo, cabe à escola orientar a família na busca de recursos para diagnóstico e obtenção de dados.

            **Elaboração:** pela equipe multidisciplinar da escola; discutido com a família e profissionais externos no início do ano letivo; replanejado ao final de cada unidade e/ou período de avaliação. Em transferência, o PEI deve compor a documentação do aluno egresso.
            """)
            st.markdown("**Registros fundamentais:**")
            st.markdown("""
            - Identidade do aluno
            - Necessidades específicas (características mais recorrentes)
            - Dados sobre autonomia
            - Dados atualizados sobre atendimentos externos
            - Desenvolvimento escolar (leitura e raciocínio lógico-matemático)
            - Necessidades de material pedagógico e tecnologias assistivas
            """)
            st.markdown("**Avaliação da aprendizagem:**")
            st.markdown("""
            Parte de objetivos mensuráveis e metas educacionais específicas. Deve considerar: implicações das funções e estrutura corporal, limitações no desempenho, condições socioemocionais. A avaliação **não incide apenas sobre o resultado**, mas sobre o processo — identificando conhecimentos construídos, dificuldades superadas e as que necessitam de mais tempo. A avaliação inclusiva **não classifica** em relação à turma; o que importa é o quanto e o como o aluno avança em relação ao seu conhecimento inicial. Nota/conceito segundo o avanço constatado.
            """)
            st.caption("A família deve acompanhar a elaboração do PEI/PDI e consentir formalmente, participando da análise das avaliações sistemáticas.")

    # =========================
    # DIREITA: Gestão de alunos
    # =========================
    with col_right:
        st.markdown(f"#### {icon_title('Gestão de Estudantes', 'estudantes', 20, '#0EA5E9')}", unsafe_allow_html=True)

        # garante d (se seu código já define antes, isso não atrapalha)
        d = st.session_state.get("dados", {})
        if not isinstance(d, dict):
            d = {}

        # Status vínculo
        student_id = st.session_state.get("selected_student_id")
        if student_id:
            st.success("✅ Estudante vinculado ao Supabase (nuvem)")
            st.caption(f"student_id: {str(student_id)[:8]}...")
        else:
            st.warning("📝 Modo rascunho (sem vínculo na nuvem)")

        # ------------------------------------------------------------------
        # (1) BACKUP LOCAL: upload JSON NÃO aplica sozinho (evita loop)
        # ------------------------------------------------------------------
        with st.container(border=True):
            st.markdown("##### 1) Carregar Backup Local (.JSON)")
            st.caption("✅ Não comunica com Supabase. Envie o arquivo e clique em **Carregar no formulário**.")

            # estados do fluxo local (cache em memória)
            if "local_json_pending" not in st.session_state:
                st.session_state["local_json_pending"] = None
            if "local_json_name" not in st.session_state:
                st.session_state["local_json_name"] = ""

            up_json = st.file_uploader(
                "Envie um arquivo .json",
                type="json",
                key="inicio_uploader_json",
            )

            # 1) Ao enviar: só guardar em memória (não aplicar)
            if up_json is not None:
                try:
                    payload = json.load(up_json)
                    payload = _coerce_dates_in_payload(payload)

                    st.session_state["local_json_pending"] = payload
                    st.session_state["local_json_name"] = getattr(up_json, "name", "") or "backup.json"

                    st.success(f"Arquivo pronto ✅ ({st.session_state['local_json_name']})")
                    st.caption("Agora clique no botão abaixo para aplicar os dados no formulário.")
                except Exception as e:
                    st.session_state["local_json_pending"] = None
                    st.session_state["local_json_name"] = ""
                    st.error(f"Erro ao ler JSON: {e}")

            pending = st.session_state.get("local_json_pending")

            # 2) Prévia (opcional)
            if isinstance(pending, dict) and pending:
                with st.expander("👀 Prévia do backup", expanded=False):
                    st.write({
                        "nome": pending.get("nome"),
                        "serie": pending.get("serie"),
                        "turma": pending.get("turma"),
                        "diagnostico": pending.get("diagnostico"),
                        "tem_ia_sugestao": bool(pending.get("ia_sugestao")),
                    })

            # 3) Botões
            b1, b2 = st.columns(2)

            with b1:
                if st.button(
                    "📥 Carregar no formulário",
                    type="primary",
                    use_container_width=True,
                    disabled=not isinstance(pending, dict),
                    key="inicio_btn_aplicar_json_local",
                ):
                    # aplica no estado do formulário
                    if "dados" in st.session_state and isinstance(st.session_state.dados, dict):
                        st.session_state.dados.update(pending)
                    else:
                        st.session_state.dados = pending

                    # JSON local NÃO cria vínculo com nuvem
                    st.session_state["selected_student_id"] = None
                    st.session_state["selected_student_name"] = ""

                    # limpa pendência pra não reaplicar
                    st.session_state["local_json_pending"] = None
                    st.session_state["local_json_name"] = ""

                    st.success("Backup aplicado ao formulário ✅")
                    st.toast("Dados aplicados.", icon="✅")
                    st.rerun()

            with b2:
                if st.button(
                    "🧹 Limpar pendência",
                    use_container_width=True,
                    key="inicio_btn_limpar_json_local",
                ):
                    st.session_state["local_json_pending"] = None
                    st.session_state["local_json_name"] = ""
                    st.rerun()

        # ------------------------------------------------------------------
        # (2) CLOUD — SINCRONIZAÇÃO COMPLETA
        # ------------------------------------------------------------------
        with st.container(border=True):
            st.caption("🌐 Omnisfera Cloud")
            st.markdown(
                "<div style='font-size:.85rem; color:#4A5568; margin-bottom:8px;'>"
                "Sincroniza o cadastro e <b>salva todo o conteúdo do PEI</b> na nuvem (coluna pei_data)."
                "</div>",
                unsafe_allow_html=True
            )

            def _cloud_ready_check():
                try:
                    url = str(ou.get_setting("SUPABASE_URL", "")).strip()
                    key = str(
                        ou.get_setting("SUPABASE_SERVICE_KEY", "")
                        or ou.get_setting("SUPABASE_ANON_KEY", "")
                        or ""
                    ).strip()
                    return bool(url and key)
                except Exception:
                    return False

            if st.button("🔗 Sincronizar Tudo", type="primary", use_container_width=True, key="btn_sync_full_final"):
                if not _cloud_ready_check():
                    st.error("⚠️ Configure os Secrets do Supabase.")
                else:
                    try:
                        with st.spinner("Sincronizando dados completos..."):
                            # 1) Datas
                            nasc_iso = d.get("nasc").isoformat() if hasattr(d.get("nasc"), "isoformat") else None

                            # 2) Payload básico (tabela students)
                            student_payload = {
                                "name": d.get("nome"),
                                "birth_date": nasc_iso,
                                "grade": d.get("serie"),
                                "class_group": d.get("turma") or None,
                                "diagnosis": d.get("diagnostico") or None,
                                "workspace_id": st.session_state.get("workspace_id"),
                            }

                            # 3) Identificar / Criar
                            sid = st.session_state.get("selected_student_id")

                            if not sid:
                                created = db_create_student(student_payload)
                                if created and isinstance(created, dict):
                                    sid = created.get("id")
                                    st.session_state["selected_student_id"] = sid
                                    st.session_state["students_cache_invalid"] = True  # outras páginas recarregam a lista
                            else:
                                db_update_student(sid, student_payload)

                            # 4) SALVAR conteúdo completo (JSONB pei_data)
                            if sid:
                                db_update_pei_content(sid, d)

                                # 5) Backup local pós-sync
                                st.session_state["ultimo_backup_json"] = json.dumps(d, default=str, ensure_ascii=False)
                                st.session_state["sync_sucesso"] = True

                                st.toast("PEI completo salvo na nuvem com sucesso!", icon="☁️")
                                st.rerun()
                            else:
                                st.error("Erro: Não foi possível obter o ID do estudante no banco.")

                    except Exception as e:
                        st.error(f"Erro na sincronização: {e}")

            # Pós sucesso: botão de download
            if st.session_state.get("sync_sucesso"):
                st.toast("Tudo salvo no Supabase!")

                timestamp = datetime.now().strftime("%d-%m_%Hh%M")
                nome_clean = (d.get("nome") or "Estudante").replace(" ", "_")

                st.download_button(
                    label="📂 BAIXAR BACKUP (.JSON)",
                    data=st.session_state.get("ultimo_backup_json", "{}"),
                    file_name=f"PEI_{nome_clean}_{timestamp}.json",
                    mime="application/json",
                    type="secondary",
                    use_container_width=True,
                    key="btn_post_sync_download_final"
                )

# ==============================================================================
# 12. ABA ESTUDANTE
# ==============================================================================
  
with tab1:
    render_progresso()
    st.markdown("### <i class='ri-user-smile-line'></i> Dossiê do Estudante", unsafe_allow_html=True)

    # Garantias (caso algo não tenha entrado no default_state)
    st.session_state.dados.setdefault("matricula", "")
    st.session_state.dados.setdefault("meds_extraidas_tmp", [])
    st.session_state.dados.setdefault("status_meds_extraidas", "idle")

    # =========================
    # Funções de apoio da aba
    # =========================
    def detectar_segmento(serie_str: str) -> str:
        """Retorna: EI | EFI | EFII | EM"""
        if not serie_str:
            return "INDEFINIDO"
        s = serie_str.lower()
        if "infantil" in s:
            return "EI"
        if "1º ano" in s or "2º ano" in s or "3º ano" in s or "4º ano" in s or "5º ano" in s:
            return "EFI"
        if "6º ano" in s or "7º ano" in s or "8º ano" in s or "9º ano" in s:
            return "EFII"
        if "série" in s or "médio" in s or "eja" in s:
            return "EM"
        return "INDEFINIDO"

    def get_segmento_info_visual_v2(serie: str):
        seg = detectar_segmento(serie)
        if seg == "EI":
            return "Educação Infantil", "#4299e1", "Foco: Campos de Experiência (BNCC) e rotina estruturante."
        if seg == "EFI":
            return "Ensino Fundamental Anos Iniciais (EFAI)", "#48bb78", "Foco: alfabetização, numeracia e consolidação de habilidades basais."
        if seg == "EFII":
            return "Ensino Fundamental Anos Finais (EFAF)", "#ed8936", "Foco: autonomia, funções executivas, organização e aprofundamento conceitual."
        if seg == "EM":
            return "Ensino Médio / EJA", "#9f7aea", "Foco: projeto de vida, áreas do conhecimento e estratégias de estudo."
        return "Selecione a Série/Ano", "#718096", "Aguardando seleção..."

    def _normalizar_med(m: dict):
        return {
            "nome": (m.get("nome") or "").strip(),
            "posologia": (m.get("posologia") or "").strip(),
            "escola": bool(m.get("escola", False)),
        }

    def _ja_existe_med(lista, nome):
        nome_norm = (nome or "").strip().lower()
        if not nome_norm:
            return True
        return any((x.get("nome") or "").strip().lower() == nome_norm for x in (lista or []))

    # =========================
    # Identificação
    # =========================
    c1, c2, c3, c4, c5 = st.columns([3, 2, 2, 1, 2])

    st.session_state.dados["nome"] = c1.text_input("Nome Completo", st.session_state.dados.get("nome", ""))
    st.session_state.dados["nasc"] = c2.date_input("Nascimento", value=st.session_state.dados.get("nasc", date(2015, 1, 1)))

    # Série/Ano — EI usa faixas do bncc_ei.csv (2, 3, 4, 5 anos)
    lista_series = _construir_lista_series(st.session_state.dados.get("serie"))
    try:
        serie_val = st.session_state.dados.get("serie")
        serie_idx = lista_series.index(serie_val) if serie_val in lista_series else 0
    except Exception:
        serie_idx = 0
    st.session_state.dados["serie"] = c3.selectbox("Série/Ano", lista_series, index=serie_idx, placeholder="Selecione...")

    # Segmento guiado (badge + descrição)
    if st.session_state.dados.get("serie"):
        seg_nome, seg_cor, seg_desc = get_segmento_info_visual_v2(st.session_state.dados["serie"])
        c3.markdown(
            f"<div class='segmento-badge' style='background-color:{seg_cor}'>{seg_nome}</div>",
            unsafe_allow_html=True
        )
        st.caption(seg_desc)

    st.session_state.dados["turma"] = c4.text_input("Turma", st.session_state.dados.get("turma", ""))

    # Matrícula / RA
    st.session_state.dados["matricula"] = c5.text_input("Matrícula / RA", st.session_state.dados.get("matricula", ""), placeholder="Ex: 2026-001234")

    st.divider()

    # =========================
    # Histórico & Família
    # =========================
    st.markdown("##### Histórico & Contexto Familiar")
    c_hist, c_fam = st.columns(2)
    st.session_state.dados["historico"] = c_hist.text_area("Histórico Escolar", st.session_state.dados.get("historico", ""))
    st.session_state.dados["familia"] = c_fam.text_area("Dinâmica Familiar", st.session_state.dados.get("familia", ""))

    default_familia_valido = [x for x in st.session_state.dados.get("composicao_familiar_tags", []) if x in LISTA_FAMILIA]
    st.session_state.dados["composicao_familiar_tags"] = st.multiselect(
        "Quem convive com o estudante?",
        LISTA_FAMILIA,
        default=default_familia_valido,
        help="Incluímos Mãe 1 / Mãe 2 e Pai 1 / Pai 2 para famílias diversas."
    )

    st.divider()

    # =========================
    # Laudo PDF + Extração IA
    # =========================
    st.markdown("##### 📎 Laudo (PDF) + Extração Inteligente")

    col_pdf, col_action = st.columns([2, 1], vertical_alignment="center")

    with col_pdf:
        up = st.file_uploader(
            "Arraste o arquivo aqui",
            type="pdf",
            label_visibility="collapsed",
            key="pei_laudo_pdf_uploader_tab1",
        )
        if up:
            st.session_state.pdf_text = ler_pdf(up)
            if st.session_state.pdf_text:
                st.success("PDF lido ✅ (usando até 6 páginas)")
            else:
                st.warning("Não consegui extrair texto do PDF (pode estar escaneado/imagem).")

    with col_action:
        st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)
        cbtn1, cbtn2, cbtn3 = st.columns([1, 2, 1])
        with cbtn2:
            extrair = st.button(
                "✨ Extrair Dados do Laudo",
                type="primary",
                use_container_width=True,
                disabled=(not st.session_state.get("pdf_text")),
                key="btn_extrair_laudo_tab1",
            )

        if extrair:
            with st.spinner("Analisando laudo..."):
                dados_extraidos, erro = extrair_dados_pdf_ia(api_key, st.session_state.pdf_text)

            if dados_extraidos:
                # 1) Diagnóstico: preencher o campo existente
                diag = (dados_extraidos.get("diagnostico") or "").strip()
                if diag:
                    st.session_state.dados["diagnostico"] = diag

                # 2) Medicações: preparar revisão (não inserir direto)
                meds = dados_extraidos.get("medicamentos") or []
                meds_norm = []
                for med in meds:
                    m = _normalizar_med(med)
                    if m["nome"]:
                        meds_norm.append(m)

                st.session_state.dados["meds_extraidas_tmp"] = meds_norm
                st.session_state.dados["status_meds_extraidas"] = "review" if meds_norm else "idle"

                st.success("Dados extraídos ✅ (revise as medicações abaixo)")
                st.rerun()
            else:
                st.error(f"Erro: {erro}")

    # Revisão das meds extraídas (antes de inserir na lista oficial)
    if st.session_state.dados.get("status_meds_extraidas") == "review":
        meds_tmp = st.session_state.dados.get("meds_extraidas_tmp", [])

        with st.container(border=True):
            st.markdown("**💊 Medicações encontradas no laudo (confirme antes de adicionar)**")

            if not meds_tmp:
                st.info("Nenhuma medicação identificada.")
                st.session_state.dados["status_meds_extraidas"] = "idle"
            else:
                for i, m in enumerate(meds_tmp):
                    cc1, cc2, cc3 = st.columns([3, 2, 1.5])
                    m["nome"] = cc1.text_input("Nome", value=m.get("nome", ""), key=f"tmp_med_nome_{i}")
                    m["posologia"] = cc2.text_input("Posologia", value=m.get("posologia", ""), key=f"tmp_med_pos_{i}")
                    m["escola"] = cc3.checkbox("Na escola?", value=bool(m.get("escola", False)), key=f"tmp_med_esc_{i}")

                a1, a2, a3 = st.columns([2, 2, 2])

                if a1.button("✅ Adicionar ao PEI", type="primary", use_container_width=True, key="btn_add_meds_tmp"):
                    # inserir no campo existente: lista_medicamentos (sem duplicar por nome)
                    lista_atual = st.session_state.dados.get("lista_medicamentos", [])
                    for m in meds_tmp:
                        m = _normalizar_med(m)
                        if m["nome"] and not _ja_existe_med(lista_atual, m["nome"]):
                            lista_atual.append(m)

                    st.session_state.dados["lista_medicamentos"] = lista_atual
                    st.session_state.dados["meds_extraidas_tmp"] = []
                    st.session_state.dados["status_meds_extraidas"] = "idle"
                    st.success("Medicações adicionadas ✅")
                    st.rerun()

                if a2.button("🧹 Limpar lista extraída", use_container_width=True, key="btn_clear_meds_tmp"):
                    st.session_state.dados["meds_extraidas_tmp"] = []
                    st.session_state.dados["status_meds_extraidas"] = "idle"
                    st.rerun()

                if a3.button("↩️ Voltar sem adicionar", use_container_width=True, key="btn_back_meds_tmp"):
                    st.session_state.dados["status_meds_extraidas"] = "idle"
                    st.rerun()

    st.divider()

    # =========================
    # Contexto Clínico + Medicação (campo EXISTENTE)
    # =========================
    st.markdown("##### Contexto Clínico")
    st.session_state.dados["diagnostico"] = st.text_input("Diagnóstico", st.session_state.dados.get("diagnostico", ""))

    with st.container(border=True):
        usa_med = st.toggle(
            "💊 O estudante faz uso contínuo de medicação?",
            value=len(st.session_state.dados.get("lista_medicamentos", [])) > 0,
            key="toggle_usa_med_tab1"
        )

        if usa_med:
            cmed1, cmed2, cmed3 = st.columns([3, 2, 2])
            nm = cmed1.text_input("Nome", key="nm_med_manual")
            pos = cmed2.text_input("Posologia", key="pos_med_manual")
            admin_escola = cmed3.checkbox("Na escola?", key="adm_esc_manual")

            if st.button("Adicionar", key="btn_add_med_manual"):
                if nm.strip():
                    # não duplicar por nome
                    if not _ja_existe_med(st.session_state.dados.get("lista_medicamentos", []), nm):
                        st.session_state.dados["lista_medicamentos"].append(
                            {"nome": nm.strip(), "posologia": pos.strip(), "escola": admin_escola}
                        )
                    st.rerun()

        if st.session_state.dados.get("lista_medicamentos"):
            st.write("---")
            for i, m in enumerate(st.session_state.dados["lista_medicamentos"]):
                tag = " [NA ESCOLA]" if m.get("escola") else ""
                c_txt, c_btn = st.columns([5, 1])
                c_txt.info(f"💊 **{m.get('nome','')}** ({m.get('posologia','')}){tag}")
                if c_btn.button("Excluir", key=f"del_med_{i}"):
                    st.session_state.dados["lista_medicamentos"].pop(i)
                    st.rerun()
# ==============================================================================
# 13. ABA EVIDÊNCIAS (COMPLETA)
# ==============================================================================
with tab2:
    render_progresso()
    st.markdown("### <i class='ri-search-eye-line'></i> Coleta de Evidências", unsafe_allow_html=True)

    atual = st.session_state.dados.get("nivel_alfabetizacao")
    idx = LISTA_ALFABETIZACAO.index(atual) if atual in LISTA_ALFABETIZACAO else 0
    st.session_state.dados["nivel_alfabetizacao"] = st.selectbox("Hipótese de Escrita", LISTA_ALFABETIZACAO, index=idx, help="Nível de apropriação do sistema de escrita (Emília Ferreiro).")

    st.divider()
    st.caption("Marque as evidências observadas na rotina do estudante (pedagógicas, cognitivas e comportamentais).")
    c1, c2, c3 = st.columns(3)

    def _tog(label):
        st.session_state.dados["checklist_evidencias"][label] = st.toggle(
            label,
            value=st.session_state.dados["checklist_evidencias"].get(label, False),
        )

    with c1:
        st.markdown("**Pedagógico**")
        for q in [
            "Estagnação na aprendizagem",
            "Lacuna em pré-requisitos",
            "Dificuldade de generalização",
            "Dificuldade de abstração",
        ]:
            _tog(q)

    with c2:
        st.markdown("**Cognitivo**")
        for q in [
            "Oscilação de foco",
            "Fadiga mental rápida",
            "Dificuldade de iniciar tarefas",
            "Esquecimento recorrente",
        ]:
            _tog(q)

    with c3:
        st.markdown("**Comportamental**")
        for q in [
            "Dependência de mediação (1:1)",
            "Baixa tolerância à frustração",
            "Desorganização de materiais",
            "Recusa de tarefas",
        ]:
            _tog(q)

    st.divider()
    st.markdown("##### Observações rápidas")
    st.session_state.dados["orientacoes_especialistas"] = st.text_area(
        "Registre observações de professores e especialistas (se houver)",
        st.session_state.dados.get("orientacoes_especialistas", ""),
        height=120,
    )

# ==============================================================================
# 14. ABA REDE DE APOIO (COMPLETA)
# ==============================================================================
with tab3:
    render_progresso()
    st.markdown("### <i class='ri-team-line'></i> Rede de Apoio", unsafe_allow_html=True)

    # Garantias (caso algo não tenha entrado no default_state)
    st.session_state.dados.setdefault("rede_apoio", [])
    st.session_state.dados.setdefault("orientacoes_especialistas", "")
    st.session_state.dados.setdefault("orientacoes_por_profissional", {})

    st.caption("Selecione os profissionais envolvidos e registre as orientações específicas de cada um.")

    # 1) Seleção da rede
    selecionados = st.multiselect(
        "Profissionais:",
        LISTA_PROFISSIONAIS,
        default=[p for p in st.session_state.dados.get("rede_apoio", []) if p in LISTA_PROFISSIONAIS],
        help="Ao selecionar um profissional, um campo de observação individual aparece abaixo."
    )
    st.session_state.dados["rede_apoio"] = selecionados

    # 2) Limpeza automática de chaves que não existem mais
    # (se o usuário desmarcar um profissional, removemos o texto dele do dicionário)
    orient_map = st.session_state.dados.get("orientacoes_por_profissional", {})
    orient_map = {k: v for k, v in orient_map.items() if k in selecionados}
    st.session_state.dados["orientacoes_por_profissional"] = orient_map

    st.divider()

    # 3) Campo geral (opcional) — mantém compatibilidade com o legado
    with st.expander("🗒️ Anotações gerais (opcional)", expanded=False):
        st.session_state.dados["orientacoes_especialistas"] = st.text_area(
            "Orientações clínicas gerais / resumo",
            st.session_state.dados.get("orientacoes_especialistas", ""),
            placeholder="Use para observações gerais da equipe (ex.: acordos com a família, encaminhamentos, alinhamentos).",
            height=140,
            key="txt_orientacoes_gerais_rede"
        )

    # 4) Campos individuais por profissional
    st.markdown(f"#### {icon_title('Orientações por profissional', 'info', 20, '#0EA5E9')}", unsafe_allow_html=True)
    if not selecionados:
        st.info("Selecione ao menos um profissional para habilitar os campos de observação.")
    else:
        # Layout em cards (2 colunas)
        cols = st.columns(2)
        for i, prof in enumerate(selecionados):
            alvo = cols[i % 2]
            with alvo:
                icon = get_pro_icon(prof) if "get_pro_icon" in globals() else "👤"
                with st.container(border=True):
                    st.markdown(f"**{icon} {prof}**")

                    st.session_state.dados["orientacoes_por_profissional"].setdefault(prof, "")

                    st.session_state.dados["orientacoes_por_profissional"][prof] = st.text_area(
                        "Observações / orientações",
                        value=st.session_state.dados["orientacoes_por_profissional"].get(prof, ""),
                        placeholder="Ex.: recomendações de intervenção, frequência, sinais de alerta, ajustes para sala de aula...",
                        height=140,
                        key=f"txt_orient_{prof}"
                    )

                    c1, c2 = st.columns([1, 1])
                    if c1.button("🧹 Limpar", use_container_width=True, key=f"btn_limpar_{prof}"):
                        st.session_state.dados["orientacoes_por_profissional"][prof] = ""
                        st.rerun()

                    if c2.button("🗑️ Remover profissional", use_container_width=True, key=f"btn_remove_{prof}"):
                        # remove do multiselect
                        st.session_state.dados["rede_apoio"] = [x for x in st.session_state.dados["rede_apoio"] if x != prof]
                        # remove do dicionário
                        st.session_state.dados["orientacoes_por_profissional"].pop(prof, None)
                        st.rerun()

    st.divider()

    # 5) Resumo visual rápido
    if selecionados:
        resumo = []
        for p in selecionados:
            txt = (st.session_state.dados["orientacoes_por_profissional"].get(p) or "").strip()
            resumo.append(f"- **{p}**: {'✅ preenchido' if txt else '⚠️ vazio'}")
        st.markdown(f"##### {icon_title('Checklist de preenchimento', 'validar', 18, '#16A34A')}", unsafe_allow_html=True)
        st.markdown("\n".join(resumo))


# ==============================================================================
# 15. ABA MAPEAMENTO (3 colunas | hiperfoco + potências + barreiras + nível de apoio + observações)
# ==============================================================================
with tab4:
    render_progresso()
    st.markdown("### <i class='ri-radar-line'></i> Mapeamento", unsafe_allow_html=True)
    st.caption("Mapeie forças, hiperfocos e barreiras. Para cada barreira selecionada, indique a intensidade de apoio necessária.")

    # -------------------------
    # Garantias de estado
    # -------------------------
    st.session_state.dados.setdefault("hiperfoco", "")
    st.session_state.dados.setdefault("potencias", [])
    st.session_state.dados.setdefault("barreiras_selecionadas", {k: [] for k in LISTAS_BARREIRAS.keys()})
    st.session_state.dados.setdefault("niveis_suporte", {})          # chave: f"{dominio}_{barreira}" -> valor
    st.session_state.dados.setdefault("observacoes_barreiras", {})   # texto livre por domínio

    # -------------------------
    # 1) POTENCIALIDADES + HIPERFOCO
    # -------------------------
    with st.container(border=True):
        st.markdown(
            f"#### {icon_title('Potencialidades e Hiperfoco', 'pei', 20, '#0EA5E9')}",
            unsafe_allow_html=True
        )
        c1, c2 = st.columns(2)

        st.session_state.dados["hiperfoco"] = c1.text_input(
            "Hiperfoco (se houver)",
            st.session_state.dados.get("hiperfoco", ""),
            placeholder="Ex.: Dinossauros, Minecraft, Mapas, Carros, Desenho..."
        )

        pot_validas = [p for p in st.session_state.dados.get("potencias", []) if p in LISTA_POTENCIAS]
        st.session_state.dados["potencias"] = c2.multiselect(
            "Potencialidades / Pontos fortes",
            LISTA_POTENCIAS,
            default=pot_validas
        )

    st.divider()

    st.markdown(
        f"#### {icon_title('Barreiras e nível de apoio', 'configurar', 20, '#0EA5E9')}",
        unsafe_allow_html=True
    )
    st.caption("Selecione as barreiras observadas e defina o nível de apoio para a rotina escolar (não é DUA).")

    # -------------------------
    # 2) Renderização por domínio
    # -------------------------
    def render_dominio(dominio: str, opcoes: list[str]):
        with st.container(border=True):
            st.markdown(f"**{dominio}**")

            # multiselect
            salvas = [b for b in st.session_state.dados["barreiras_selecionadas"].get(dominio, []) if b in opcoes]
            selecionadas = st.multiselect(
                "Selecione as barreiras",
                opcoes,
                default=salvas,
                key=f"ms_{dominio}",
                label_visibility="collapsed"
            )
            st.session_state.dados["barreiras_selecionadas"][dominio] = selecionadas

            # sliders por barreira (bem visível: nome + barra na mesma linha)
            if selecionadas:
                st.markdown("---")
                st.markdown("**Nível de apoio por barreira**")
                st.caption("Escala: Autônomo (faz sozinho) → Monitorado → Substancial → Muito Substancial (suporte intenso/contínuo).")

                for b in selecionadas:
                    chave = f"{dominio}_{b}"
                    st.session_state.dados["niveis_suporte"].setdefault(chave, "Monitorado")

                    colA, colB = st.columns([2.2, 2.8], vertical_alignment="center")
                    with colA:
                        st.markdown(f"**{b}**")
                    with colB:
                        st.session_state.dados["niveis_suporte"][chave] = st.select_slider(
                            "Nível de apoio",
                            options=["Autônomo", "Monitorado", "Substancial", "Muito Substancial"],
                            value=st.session_state.dados["niveis_suporte"].get(chave, "Monitorado"),
                            key=f"sl_{dominio}_{b}",
                            label_visibility="collapsed",
                            help=(
                                "Autônomo: realiza sem mediação | "
                                "Monitorado: precisa de checagens | "
                                "Substancial: precisa de mediação frequente | "
                                "Muito Substancial: precisa de suporte intenso/contínuo"
                            )
                        )

            # observação por domínio (mantido)
            st.session_state.dados["observacoes_barreiras"].setdefault(dominio, "")
            st.session_state.dados["observacoes_barreiras"][dominio] = st.text_area(
                "Observações (opcional)",
                value=st.session_state.dados["observacoes_barreiras"].get(dominio, ""),
                placeholder="Ex.: quando ocorre, gatilhos, o que ajuda, o que piora, estratégias que já funcionam...",
                height=90,
                key=f"obs_{dominio}"
            )

    # -------------------------
    # 3) 3 colunas (distribuição como era antes)
    # -------------------------
    c_bar1, c_bar2, c_bar3 = st.columns(3)

    with c_bar1:
        render_dominio("Funções Cognitivas", LISTAS_BARREIRAS.get("Funções Cognitivas", []))
        render_dominio("Sensorial e Motor", LISTAS_BARREIRAS.get("Sensorial e Motor", []))

    with c_bar2:
        render_dominio("Comunicação e Linguagem", LISTAS_BARREIRAS.get("Comunicação e Linguagem", []))
        render_dominio("Acadêmico", LISTAS_BARREIRAS.get("Acadêmico", []))

    with c_bar3:
        render_dominio("Socioemocional", LISTAS_BARREIRAS.get("Socioemocional", []))

    # -------------------------
    # 4) Limpeza automática (remove níveis de suporte de barreiras desmarcadas)
    # -------------------------
    chaves_validas = set()
    for dom, itens in st.session_state.dados["barreiras_selecionadas"].items():
        for b in itens:
            chaves_validas.add(f"{dom}_{b}")

    niveis = st.session_state.dados.get("niveis_suporte", {})
    st.session_state.dados["niveis_suporte"] = {k: v for k, v in niveis.items() if k in chaves_validas}

    st.divider()

    # -------------------------
    # 5) Resumo
    # -------------------------
    st.markdown(f"#### {icon_title('Resumo do Mapeamento', 'pei', 20, '#0EA5E9')}", unsafe_allow_html=True)

    r1, r2 = st.columns(2)

    with r1:
        hf = (st.session_state.dados.get("hiperfoco") or "").strip()
        if hf:
            st.success(f"**Hiperfoco:** {hf}")
        else:
            st.info("**Hiperfoco:** não informado")

        pots = st.session_state.dados.get("potencias", [])
        if pots:
            st.success(f"**Potencialidades:** {', '.join(pots)}")
        else:
            st.info("**Potencialidades:** não selecionadas")

    with r2:
        selecionadas = {dom: vals for dom, vals in st.session_state.dados["barreiras_selecionadas"].items() if vals}
        total_bar = sum(len(v) for v in selecionadas.values())

        if total_bar == 0:
            st.info("**Barreiras:** nenhuma selecionada")
        else:
            st.warning(f"**Barreiras selecionadas:** {total_bar}")
            for dom, vals in selecionadas.items():
                st.markdown(f"**{dom}:**")
                for b in vals:
                    chave = f"{dom}_{b}"
                    nivel = st.session_state.dados["niveis_suporte"].get(chave, "Monitorado")
                    st.markdown(f"- {b} → **{nivel}**")


# ==============================================================================
# 16. ABA PLANO DE AÇÃO (COMPLETA)
# ==============================================================================
with tab5:
    render_progresso()
    st.markdown("### <i class='ri-puzzle-line'></i> Plano de Ação", unsafe_allow_html=True)

    c1, c2, c3 = st.columns(3)

    with c1:
        st.markdown("#### 1) Acesso (DUA)")
        st.session_state.dados["estrategias_acesso"] = st.multiselect(
            "Recursos de acesso",
            [
                "Tempo Estendido",
                "Apoio Leitura/Escrita",
                "Material Ampliado",
                "Tecnologia Assistiva",
                "Sala Silenciosa",
                "Mobiliário Adaptado",
                "Pistas Visuais",
                "Rotina Estruturada",
            ],
            default=st.session_state.dados.get("estrategias_acesso", []),
        )
        st.session_state.dados["outros_acesso"] = st.text_input(
            "Personalizado (Acesso)",
            st.session_state.dados.get("outros_acesso", ""),
            placeholder="Ex: Prova em local separado, fonte 18, papel pautado ampliado…",
        )

    with c2:
        st.markdown("#### 2) Ensino (Metodologias)")
        st.session_state.dados["estrategias_ensino"] = st.multiselect(
            "Estratégias de ensino",
            [
                "Fragmentação de Tarefas",
                "Instrução Explícita",
                "Modelagem",
                "Mapas Mentais",
                "Andaimagem (Scaffolding)",
                "Ensino Híbrido",
                "Organizadores Gráficos",
                "Prática Guiada",
            ],
            default=st.session_state.dados.get("estrategias_ensino", []),
        )
        st.session_state.dados["outros_ensino"] = st.text_input(
            "Personalizado (Ensino)",
            st.session_state.dados.get("outros_ensino", ""),
            placeholder="Ex: Sequência didática com apoio de imagens + exemplo resolvido…",
        )

    with c3:
        st.markdown("#### 3) Avaliação (Formato)")
        st.session_state.dados["estrategias_avaliacao"] = st.multiselect(
            "Estratégias de avaliação",
            [
                "Prova Adaptada",
                "Prova Oral",
                "Consulta Permitida",
                "Portfólio",
                "Autoavaliação",
                "Parecer Descritivo",
                "Questões Menores por Bloco",
                "Avaliação Prática (Demonstração)",
            ],
            default=st.session_state.dados.get("estrategias_avaliacao", []),
        )
        st.caption("Dica: combine formato + acesso (tempo/ambiente) para reduzir barreiras.")

    st.divider()
    st.info("✅ O plano de ação alimenta a Consultoria IA com contexto prático (o que você já pretende fazer).")


# ==============================================================================
# 17. ABA MONITORAMENTO (COMPLETA)
# ==============================================================================
with tab6:
    render_progresso()
    st.markdown("### <i class='ri-loop-right-line'></i> Monitoramento", unsafe_allow_html=True)

    st.session_state.dados["monitoramento_data"] = st.date_input(
        "Data da Próxima Revisão",
        value=st.session_state.dados.get("monitoramento_data", date.today()),
    )

    st.divider()
    st.warning("⚠️ Preencher esta aba principalmente na REVISÃO do PEI (ciclo de acompanhamento).")

    with st.container(border=True):
        c2, c3 = st.columns(2)
        with c2:
            atual = st.session_state.dados.get("status_meta", "Não Iniciado")
            st.session_state.dados["status_meta"] = st.selectbox(
                "Status da Meta",
                ["Não Iniciado", "Em Andamento", "Parcialmente Atingido", "Atingido", "Superado"],
                index=(["Não Iniciado", "Em Andamento", "Parcialmente Atingido", "Atingido", "Superado"].index(atual) if atual in ["Não Iniciado", "Em Andamento", "Parcialmente Atingido", "Atingido", "Superado"] else 0),
            )
        with c3:
            atualp = st.session_state.dados.get("parecer_geral", "Manter Estratégias")
            st.session_state.dados["parecer_geral"] = st.selectbox(
                "Parecer Geral",
                [
                    "Manter Estratégias",
                    "Aumentar Suporte",
                    "Reduzir Suporte (Autonomia)",
                    "Alterar Metodologia",
                    "Encaminhar para Especialista",
                ],
                index=(
                    [
                        "Manter Estratégias",
                        "Aumentar Suporte",
                        "Reduzir Suporte (Autonomia)",
                        "Alterar Metodologia",
                        "Encaminhar para Especialista",
                    ].index(atualp)
                    if atualp in [
                        "Manter Estratégias",
                        "Aumentar Suporte",
                        "Reduzir Suporte (Autonomia)",
                        "Alterar Metodologia",
                        "Encaminhar para Especialista",
                    ]
                    else 0
                ),
            )

        st.session_state.dados["proximos_passos_select"] = st.multiselect(
            "Ações Futuras",
            [
                "Reunião com Família",
                "Encaminhamento Clínico",
                "Adaptação de Material",
                "Mudança de Lugar em Sala",
                "Novo PEI",
                "Observação em Sala",
            ],
            default=st.session_state.dados.get("proximos_passos_select", []),
        )


# ==============================================================================
# 17A. ABA BNCC — EI: Campos + Objetivos | EF/EM: Habilidades
# ==============================================================================
with tab7_hab:
    render_progresso()
    st.markdown("### <i class='ri-list-check-2'></i> BNCC", unsafe_allow_html=True)

    serie_hab = st.session_state.dados.get("serie") or ""
    if not serie_hab:
        st.warning("Selecione a **Série/Ano** (ou faixa de idade para EI) na aba **Estudante**.")
        st.stop()

    # Educação Infantil: Campos de Experiência + Objetivos de Aprendizagem (bncc_ei.csv — todas as faixas do CSV)
    if detectar_nivel_ensino(serie_hab) == "EI":
        try:
            from services.bncc_service import faixas_idade_ei, campos_experiencia_ei, objetivos_ei_por_idade_campo, carregar_bncc_ei
            bncc_ei = carregar_bncc_ei()
            faixas = faixas_idade_ei() or ["0 a 1 ano e 6 meses", "1 ano e 7 meses a 3 anos e 11 meses", "4 anos a 5 anos e 11 meses"]
            campos = campos_experiencia_ei() if bncc_ei else [
                "O eu, o outro e o nós", "Corpo, gestos e movimentos", "Traços, sons, cores e formas",
                "Escuta, fala, pensamento e imaginação", "Espaços, tempos, quantidades, relações e transformações"
            ]
        except Exception:
            bncc_ei = []
            faixas = ["0 a 1 ano e 6 meses", "1 ano e 7 meses a 3 anos e 11 meses", "4 anos a 5 anos e 11 meses"]
            campos = ["O eu, o outro e o nós", "Corpo, gestos e movimentos", "Traços, sons, cores e formas",
                      "Escuta, fala, pensamento e imaginação", "Espaços, tempos, quantidades, relações e transformações"]
            def objetivos_ei_por_idade_campo(idade, campo):
                return []
        st.caption("Educação Infantil: selecione faixa de idade, campo de experiência e objetivos de aprendizagem. A Consultoria IA usará estes dados.")
        col_ei1, col_ei2 = st.columns(2)
        with col_ei1:
            idade_ei = st.selectbox("Faixa de Idade", faixas, key="pei_idade_ei",
                index=faixas.index(st.session_state.dados.get("bncc_ei_idade", "") or faixas[0]) if (st.session_state.dados.get("bncc_ei_idade") or "") in faixas else 0)
            campo_ei = st.selectbox("Campo de Experiência", campos, key="pei_campo_ei",
                index=campos.index(st.session_state.dados.get("bncc_ei_campo", "") or campos[0]) if (st.session_state.dados.get("bncc_ei_campo") or "") in campos else 0)
        with col_ei2:
            obj_opcoes = objetivos_ei_por_idade_campo(idade_ei, campo_ei) if bncc_ei else []
            obj_atuais = st.session_state.dados.get("bncc_ei_objetivos") or []
            obj_selecionados = st.multiselect("Objetivos de Aprendizagem", obj_opcoes or ["Selecione idade e campo primeiro"],
                default=[o for o in obj_atuais if o in (obj_opcoes or [])], key="pei_obj_ei")
        st.session_state.dados["bncc_ei_idade"] = idade_ei
        st.session_state.dados["bncc_ei_campo"] = campo_ei
        st.session_state.dados["bncc_ei_objetivos"] = obj_selecionados
        st.info("👉 Com os campos e objetivos selecionados, siga para a aba **Consultoria IA** para gerar o relatório.")
        st.stop()

    # EF/EM: Habilidades BNCC (fluxo existente)
    # EM usa bncc_em.csv (áreas); EF usa bncc_ef.csv (componentes) via bncc_service

    st.caption("Selecione as habilidades do ano/série do estudante. A Consultoria IA usará apenas estas para o relatório.")

    st.session_state.dados.setdefault("habilidades_bncc_selecionadas", [])
    if detectar_nivel_ensino(serie_hab) == "EM":
        try:
            from services.bncc_service import carregar_habilidades_em_por_area, carregar_bncc_em
            if carregar_bncc_em():
                blocos = carregar_habilidades_em_por_area(serie_hab)
            else:
                blocos = {"ano_atual": {}, "anos_anteriores": {}}
        except Exception:
            blocos = {"ano_atual": {}, "anos_anteriores": {}}
    else:
        blocos = carregar_habilidades_bncc_por_componente_ano_e_anteriores(serie_hab)
    ano_atual = blocos.get("ano_atual") or {}
    anos_anteriores = blocos.get("anos_anteriores") or {}
    if not ano_atual and not anos_anteriores:
        msg_em = "Para Ensino Médio: adicione bncc_em.csv na raiz (Área;Componente;Série;Unidade Temática;Habilidade). " if detectar_nivel_ensino(serie_hab) == "EM" else ""
        st.info(f"Nenhuma habilidade BNCC encontrada para esta série. {msg_em}Os arquivos bncc_ef.csv (EF) ou bncc_em.csv (EM) na raiz devem conter as habilidades.")
        st.stop()

    def _opcao_label(h: dict) -> str:
        """Label na lista suspensa: código + texto completo da habilidade (na íntegra)."""
        c = h.get("codigo", "")
        texto = h.get("habilidade_completa") or h.get("descricao") or ""
        return f"{c} — {texto}" if c else texto

    # Seleção atual: set de (disciplina, codigo, origem)
    selecionadas_atuais = st.session_state.dados.get("habilidades_bncc_selecionadas") or []
    set_selecionados = set()
    for item in selecionadas_atuais:
        if isinstance(item, dict) and item.get("codigo") and item.get("disciplina"):
            set_selecionados.add((item["disciplina"], item["codigo"], item.get("origem", "ano_atual")))

    novas_selecoes = []
    st.session_state.dados.setdefault("_hab_ia_motivo_ano_atual", "")
    st.session_state.dados.setdefault("_hab_ia_motivo_anos_anteriores", "")

    # Lista de habilidades selecionadas em aba retrátil (expander)
    with st.expander("📋 Habilidades selecionadas", expanded=bool(selecionadas_atuais)):
        if selecionadas_atuais:
            motivo_atual = st.session_state.dados.get("_hab_ia_motivo_ano_atual", "").strip()
            motivo_ant = st.session_state.dados.get("_hab_ia_motivo_anos_anteriores", "").strip()
            if motivo_atual or motivo_ant:
                st.caption("**Por que a IA escolheu estas habilidades:**")
                if motivo_atual:
                    st.markdown(f"*Ano atual:* {motivo_atual}")
                if motivo_ant:
                    st.markdown(f"*Anos anteriores:* {motivo_ant}")
                st.markdown("---")
            st.caption("Revise a lista. Use **Remover** para tirar uma habilidade ou **Desmarcar todas** para limpar.")
            for idx, h in enumerate(selecionadas_atuais):
                if not isinstance(h, dict) or not h.get("codigo"):
                    continue
                disc = h.get("disciplina", "")
                cod = h.get("codigo", "")
                texto = (h.get("habilidade_completa") or h.get("descricao") or "")
                col_txt, col_btn = st.columns([5, 1])
                with col_txt:
                    st.markdown(f"**{disc}** — *{cod}* — {texto}")
                with col_btn:
                    if st.button("Remover", key=f"remover_hab_{idx}_{cod}_{disc[:20]}", type="secondary"):
                        nova_lista = [x for x in selecionadas_atuais if isinstance(x, dict) and not (x.get("codigo") == cod and x.get("disciplina") == disc)]
                        st.session_state.dados["habilidades_bncc_selecionadas"] = nova_lista
                        st.session_state.dados["habilidades_bncc_validadas"] = None
                        st.session_state["_hab_bncc_key_salt"] = (st.session_state.get("_hab_bncc_key_salt", 0) + 1)
                        st.rerun()
            if st.button("Desmarcar todas", key="desmarcar_todas_hab_bncc", type="secondary"):
                st.session_state.dados["habilidades_bncc_selecionadas"] = []
                st.session_state.dados["habilidades_bncc_validadas"] = None
                st.session_state.dados["_hab_ia_motivo_ano_atual"] = ""
                st.session_state.dados["_hab_ia_motivo_anos_anteriores"] = ""
                st.session_state["_hab_bncc_key_salt"] = (st.session_state.get("_hab_bncc_key_salt", 0) + 1)
                st.rerun()
        else:
            st.caption("Nenhuma habilidade selecionada. Use os botões nas seções abaixo (ano atual e anos anteriores) para a IA sugerir ou selecione manualmente nas listas.")

    if st.session_state.get("_run_auxilio_hab"):
        st.session_state["_run_auxilio_hab"] = False
        if not api_key:
            st.error(f"Configure a chave da IA ({ou.AI_RED}): variável OPENAI_API_KEY, secrets ou session_state.")
        elif not ano_atual:
            st.warning("Não há habilidades do ano atual para sugerir.")
        else:
            lista_para_ia = []
            componentes_atual = sorted(ano_atual.keys())
            for disc in componentes_atual:
                for h in ano_atual[disc]:
                    lista_para_ia.append(f"- {disc}: {h.get('codigo','')} — {h.get('habilidade_completa','')}")
            texto_lista = "\n".join(lista_para_ia)
            prompt_aux = f"""O estudante está no ano/série: {serie_hab}. Abaixo estão as habilidades BNCC do ano atual (código e descrição).
Indique as habilidades mais fundamentais para esse ano (máximo 3 a 5 por componente curricular).

Retorne EXATAMENTE neste formato:
CODES:
(códigos um por linha, ex: EF01LP02)
MOTIVO:
(em 2 a 4 linhas, explique por que escolheu estas habilidades para este estudante.)

HABILIDADES DO ANO ATUAL:
{texto_lista}"""
            try:
                with st.spinner("Sugerindo habilidades fundamentais para esta série..."):
                    from openai import OpenAI
                    client = OpenAI(api_key=api_key)
                    r = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[{"role": "user", "content": prompt_aux}],
                    )
                    texto_resp = (r.choices[0].message.content or "").strip()
                    codigos_sugeridos = []
                    motivo_ia = ""
                    if "MOTIVO:" in texto_resp.upper():
                        partes = re.split(r"\s*MOTIVO\s*:\s*", texto_resp, flags=re.IGNORECASE, maxsplit=1)
                        if len(partes) >= 2:
                            motivo_ia = partes[1].strip()[:500]
                        bloco_codes = partes[0] if partes else texto_resp
                    else:
                        bloco_codes = texto_resp
                    for linha in bloco_codes.splitlines():
                        cod = re.search(r"(EF\d+[A-Z0-9]+|EM\d+[A-Z0-9]+)", linha.strip())
                        if cod:
                            codigos_sugeridos.append(cod.group(1).upper())
                    for disc in componentes_atual:
                        for h in ano_atual[disc]:
                            if (h.get("codigo") or "").upper() in codigos_sugeridos:
                                set_selecionados.add((disc, h.get("codigo", ""), "ano_atual"))
                    sug_list = []
                    for disc in componentes_atual:
                        for h in ano_atual[disc]:
                            if (disc, h.get("codigo", ""), "ano_atual") in set_selecionados:
                                sug_list.append({
                                    "disciplina": disc, "codigo": h.get("codigo", ""),
                                    "descricao": h.get("descricao", ""), "habilidade_completa": h.get("habilidade_completa", ""),
                                    "origem": "ano_atual",
                                })
                    for item in selecionadas_atuais:
                        if isinstance(item, dict) and item.get("origem") == "anos_anteriores":
                            sug_list.append(item)
                    st.session_state.dados["habilidades_bncc_selecionadas"] = sug_list
                    st.session_state.dados["_hab_ia_motivo_ano_atual"] = motivo_ia
                    st.session_state["_hab_bncc_key_salt"] = (st.session_state.get("_hab_bncc_key_salt", 0) + 1)
                st.success("Habilidades sugeridas gravadas. Veja a lista em **Habilidades selecionadas**; você pode remover uma a uma ou desmarcar todas.")
                st.rerun()
            except Exception as e:
                st.error(f"Erro ao sugerir: {str(e)[:120]}")

    # Preenchimento com auxílio da IA para anos anteriores (mesma lógica que ano atual)
    if st.session_state.get("_run_auxilio_hab_anteriores"):
        st.session_state["_run_auxilio_hab_anteriores"] = False
        if not api_key:
            st.error(f"Configure a chave da IA ({ou.AI_RED}): variável OPENAI_API_KEY, secrets ou session_state.")
        elif not anos_anteriores:
            st.warning("Não há habilidades de anos anteriores para sugerir.")
        else:
            lista_para_ia_ant = []
            componentes_ant = sorted(anos_anteriores.keys())
            for disc in componentes_ant:
                for h in anos_anteriores[disc]:
                    lista_para_ia_ant.append(f"- {disc}: {h.get('codigo','')} — {h.get('habilidade_completa','')}")
            texto_lista_ant = "\n".join(lista_para_ia_ant)
            prompt_aux_ant = f"""O estudante está no ano/série: {serie_hab}. Abaixo estão habilidades BNCC de ANOS ANTERIORES (que podem merecer atenção ou reforço).
Indique as habilidades mais relevantes para esse estudante (máximo 3 a 5 por componente curricular).

Retorne EXATAMENTE neste formato:
CODES:
(códigos um por linha, ex: EF01LP02)
MOTIVO:
(em 2 a 4 linhas, explique por que escolheu estas habilidades de anos anteriores para este estudante.)

HABILIDADES DE ANOS ANTERIORES:
{texto_lista_ant}"""
            try:
                with st.spinner("Sugerindo habilidades de anos anteriores..."):
                    from openai import OpenAI
                    client = OpenAI(api_key=api_key)
                    r = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[{"role": "user", "content": prompt_aux_ant}],
                    )
                    texto_resp = (r.choices[0].message.content or "").strip()
                    codigos_sugeridos = []
                    motivo_ia_ant = ""
                    if "MOTIVO:" in texto_resp.upper():
                        partes = re.split(r"\s*MOTIVO\s*:\s*", texto_resp, flags=re.IGNORECASE, maxsplit=1)
                        if len(partes) >= 2:
                            motivo_ia_ant = partes[1].strip()[:500]
                        bloco_codes = partes[0] if partes else texto_resp
                    else:
                        bloco_codes = texto_resp
                    for linha in bloco_codes.splitlines():
                        cod = re.search(r"(EF\d+[A-Z0-9]+|EM\d+[A-Z0-9]+)", linha.strip())
                        if cod:
                            codigos_sugeridos.append(cod.group(1).upper())
                    for disc in componentes_ant:
                        for h in anos_anteriores[disc]:
                            if (h.get("codigo") or "").upper() in codigos_sugeridos:
                                set_selecionados.add((disc, h.get("codigo", ""), "anos_anteriores"))
                    sug_list_ant = []
                    for item in selecionadas_atuais:
                        if isinstance(item, dict) and item.get("origem") == "ano_atual":
                            sug_list_ant.append(item)
                    for disc in componentes_ant:
                        for h in anos_anteriores[disc]:
                            if (disc, h.get("codigo", ""), "anos_anteriores") in set_selecionados:
                                sug_list_ant.append({
                                    "disciplina": disc, "codigo": h.get("codigo", ""),
                                    "descricao": h.get("descricao", ""), "habilidade_completa": h.get("habilidade_completa", ""),
                                    "origem": "anos_anteriores",
                                })
                    st.session_state.dados["habilidades_bncc_selecionadas"] = sug_list_ant
                    st.session_state.dados["_hab_ia_motivo_anos_anteriores"] = motivo_ia_ant
                    st.session_state["_hab_bncc_key_salt"] = (st.session_state.get("_hab_bncc_key_salt", 0) + 1)
                st.success("Habilidades de anos anteriores sugeridas. Veja a lista em **Habilidades selecionadas**; você pode remover uma a uma ou desmarcar todas.")
                st.rerun()
            except Exception as e:
                st.error(f"Erro ao sugerir anos anteriores: {str(e)[:120]}")

    def _render_multiselects(por_componente: dict, titulo: str, origem: str):
        componentes_ordenados = sorted(por_componente.keys())
        if not componentes_ordenados:
            return
        st.markdown(f"#### {titulo}")
        st.caption("Habilidades marcadas são as selecionadas (pelo professor ou sugeridas pela IA). Desmarque para excluir.")
        for disc in componentes_ordenados:
            lista_hab = por_componente[disc]
            opcoes = [_opcao_label(h) for h in lista_hab]
            default_labels = [opcoes[i] for i, h in enumerate(lista_hab) if (disc, h.get("codigo", ""), origem) in set_selecionados]
            _key_salt = st.session_state.get("_hab_bncc_key_salt", 0)
            escolhidas = st.multiselect(
                f"**{disc}**",
                options=opcoes,
                default=default_labels,
                key=f"hab_bncc_{origem}_{disc}_{_key_salt}",
            )
            for label in escolhidas:
                for i, h in enumerate(lista_hab):
                    if _opcao_label(h) == label:
                        novas_selecoes.append({
                            "disciplina": disc, "codigo": h.get("codigo", ""),
                            "descricao": h.get("descricao", ""), "habilidade_completa": h.get("habilidade_completa", ""),
                            "origem": origem,
                        })
                        break

    # Seção Ano atual: botão da IA + listas em expander retrátil
    rotulo_em = "área de conhecimento" if detectar_nivel_ensino(serie_hab) == "EM" else "componente"
    with st.expander("📌 Habilidades do ano/série atual", expanded=True):
        if st.button("Preenchimento com auxílio da IA (ano atual)", type="secondary", use_container_width=True, key="btn_auxilio_hab_bncc"):
            st.session_state["_run_auxilio_hab"] = True
        st.caption(f"Use o botão acima para a IA sugerir habilidades do ano atual; ou selecione manualmente nas listas por {rotulo_em} abaixo.")
        _render_multiselects(ano_atual, f"Listas por {rotulo_em} (ano atual)", "ano_atual")

    st.markdown("---")

    # Seção Anos anteriores: botão da IA + listas em expander retrátil (EM não tem anos anteriores no mesmo sentido)
    with st.expander("📌 Habilidades de anos anteriores (que merecem atenção)", expanded=bool(anos_anteriores)):
        if st.button("Preenchimento com auxílio da IA (anos anteriores)", type="secondary", use_container_width=True, key="btn_auxilio_hab_anteriores"):
            st.session_state["_run_auxilio_hab_anteriores"] = True
        st.caption(f"Use o botão acima para a IA sugerir habilidades de anos anteriores; ou selecione manualmente nas listas por {rotulo_em} abaixo.")
        _render_multiselects(anos_anteriores, f"Listas por {rotulo_em} (anos anteriores)", "anos_anteriores")

    st.session_state.dados["habilidades_bncc_selecionadas"] = novas_selecoes
    n_hab = len(novas_selecoes)

    # Botão Validar seleção: professor confirma que a seleção atual é a que deve ir para o relatório
    st.session_state.dados.setdefault("habilidades_bncc_validadas", None)
    col_validar, _ = st.columns([1, 2])
    with col_validar:
        if st.button("✅ Validar seleção", type="primary", use_container_width=True, key="btn_validar_hab_bncc"):
            if n_hab == 0:
                st.warning("Selecione ao menos uma habilidade nas listas acima antes de validar.")
            else:
                st.session_state.dados["habilidades_bncc_validadas"] = [
                    {**h} for h in novas_selecoes if isinstance(h, dict)
                ]
                st.success("Seleção validada pelo professor. O relatório da Consultoria IA usará **apenas** estas habilidades.")
                st.rerun()

    if st.session_state.dados.get("habilidades_bncc_validadas"):
        n_val = len(st.session_state.dados["habilidades_bncc_validadas"])
        st.info(f"📌 **{n_val}** habilidade(s) validadas. Para alterar, desmarque/marque nas listas e clique em **Validar seleção** novamente.")

    if n_hab > 0 and not st.session_state.dados.get("habilidades_bncc_validadas"):
        st.success(f"**{n_hab}** habilidade(s) selecionada(s). Clique em **Validar seleção** para o professor confirmar; a Consultoria IA usará apenas as validadas no relatório.")

    st.divider()
    st.caption("Na aba **Consultoria IA**, o relatório será gerado com base nas habilidades **validadas** e incluirá um mapa por componente curricular com maior necessidade de atenção.")


# ==============================================================================
# 18. ABA CONSULTORIA IA (COMPLETA: gerar + revisar + aprovar + ajustar)
# ==============================================================================
with tab8:
    render_progresso()
    st.markdown("### <i class='ri-robot-2-line'></i> Consultoria Pedagógica", unsafe_allow_html=True)

    if not st.session_state.dados.get("serie"):
        st.warning("⚠️ Selecione a Série/Ano na aba **Estudante** para ativar o modo especialista.")
        st.stop()

    # estado default
    st.session_state.dados.setdefault("status_validacao_pei", "rascunho")
    st.session_state.dados.setdefault("feedback_ajuste", "")
    st.session_state.dados.setdefault("consultoria_engine", "red")

    seg_nome, seg_cor, seg_desc = get_segmento_info_visual(st.session_state.dados.get("serie"))
    st.markdown(
        f"<div style='background-color:#F7FAFC; border-left:5px solid {seg_cor}; padding:14px; border-radius:8px; margin-bottom:16px;'>"
        f"<b style='color:{seg_cor};'>ℹ️ Modo Especialista: {seg_nome}</b><br>"
        f"<span style='color:#4A5568;'>{seg_desc}</span></div>",
        unsafe_allow_html=True,
    )



    # 1) Se ainda não tem texto, ou voltou para rascunho: botões de geração
    if (not st.session_state.dados.get("ia_sugestao")) or (st.session_state.dados.get("status_validacao_pei") == "rascunho"):
        with st.expander("🔧 Escolher motor de IA (Red, Blue, Green, Yellow ou Orange)", expanded=True):
            st.caption("Selecione qual IA gerará o relatório. Orange = fallback (GPT) se outros falharem.")
            engine_map = {
                "red": f"🔴 {ou.AI_RED}",
                "blue": f"🔵 {ou.AI_BLUE}",
                "green": f"🟢 {ou.AI_GREEN}",
                "yellow": f"🟡 {ou.AI_YELLOW}",
                "orange": f"🟠 {ou.AI_ORANGE}",
            }
            engine_sel = st.radio(
                "Motor de IA",
                options=["red", "blue", "green", "yellow", "orange"],
                format_func=lambda x: engine_map.get(x, x),
                index={"red": 0, "blue": 1, "green": 2, "yellow": 3, "orange": 4}.get(st.session_state.dados.get("consultoria_engine", "red"), 0),
                key="consultoria_engine_radio",
                horizontal=True,
            )
            st.session_state.dados["consultoria_engine"] = engine_sel

        col_btn, col_info = st.columns([1, 2])

        with col_btn:
            engine = st.session_state.dados.get("consultoria_engine", "red")
            if st.button("✨ Gerar Estratégia Técnica", type="primary", use_container_width=True):
                with st.spinner(ou.get_loading_message(engine)):
                    res, err = consultar_gpt_pedagogico(
                        api_key,
                        st.session_state.dados,
                        st.session_state.get("pdf_text", ""),
                        modo_pratico=False,
                        engine=engine,
                    )
                if res:
                    st.session_state.dados["ia_sugestao"] = res
                    st.session_state.dados["status_validacao_pei"] = "revisao"
                    st.rerun()
                else:
                    st.error(err or "Erro ao gerar.")

            st.write("")
            if st.button("🧰 Gerar Guia Prático (Sala de Aula)", use_container_width=True):
                with st.spinner(ou.get_loading_message(engine)):
                    res, err = consultar_gpt_pedagogico(
                        api_key,
                        st.session_state.dados,
                        st.session_state.get("pdf_text", ""),
                        modo_pratico=True,
                        engine=engine,
                    )
                if res:
                    st.session_state.dados["ia_sugestao"] = res
                    st.session_state.dados["status_validacao_pei"] = "revisao"
                    st.rerun()
                else:
                    st.error(err or "Erro ao gerar.")

        with col_info:
            n_bar = sum(len(v) for v in (st.session_state.dados.get("barreiras_selecionadas") or {}).values())
            n_hab = len(st.session_state.dados.get("habilidades_bncc_selecionadas") or [])
            hab_validadas = st.session_state.dados.get("habilidades_bncc_validadas") or []
            st.info(
                "Quanto mais completo o **Mapeamento** (barreiras + nível de suporte + hiperfoco) "
                "e o **Plano de Ação**, melhor a precisão.\n\n"
                f"📌 Barreiras mapeadas agora: **{n_bar}**"
            )
            if n_hab > 0 and not hab_validadas:
                st.warning("⚠️ Há habilidades selecionadas na aba **Habilidades BNCC** mas ainda não validadas. Clique em **Validar seleção** naquela aba para o relatório usar apenas as habilidades confirmadas pelo professor.")

    # 2) Revisão / Aprovado: mostrar e permitir aprovar/ajustar
    elif st.session_state.dados.get("status_validacao_pei") in ["revisao", "aprovado"]:
        n_barreiras = sum(len(v) for v in (st.session_state.dados.get("barreiras_selecionadas") or {}).values())
        diag_show = st.session_state.dados.get("diagnostico") or "em observação"

        with st.expander("🧠 Como a IA construiu este relatório (transparência)"):
            exemplo_barreira = "geral"
            try:
                for area, lst in (st.session_state.dados.get("barreiras_selecionadas") or {}).items():
                    if lst:
                        exemplo_barreira = lst[0]
                        break
            except Exception:
                pass

            _engine_map = {
                "red": ou.AI_RED, "blue": ou.AI_BLUE, "green": ou.AI_GREEN,
                "yellow": ou.AI_YELLOW, "orange": ou.AI_ORANGE,
            }
            _engine = st.session_state.dados.get("consultoria_engine", "red")
            _motor_nome = _engine_map.get(_engine, ou.AI_RED)

            st.markdown(
                f"**Gerado por {_motor_nome}**\n\n"
                f"**1. Input do estudante:** Série **{st.session_state.dados.get('serie','-')}**, diagnóstico **{diag_show}**.\n\n"
                f"**2. Barreiras ativas:** detectei **{n_barreiras}** barreiras e cruzei isso com BNCC + DUA.\n\n"
                f"**3. Ponto crítico exemplo:** priorizei adaptações para reduzir impacto de **{exemplo_barreira}**."
            )

        with st.expander("🛡️ Calibragem e segurança pedagógica"):
            st.markdown(
                "- **Farmacologia:** não sugere dose/medicação; apenas sinaliza pontos de atenção.\n"
                "- **Dados sensíveis:** evite inserir PII desnecessária.\n"
                "- **Normativa:** sugestões buscam aderência à LBI/DUA e adaptações razoáveis."
            )

        st.markdown("#### 📝 Revisão do Plano")
        texto_visual = re.sub(r"\[.*?\]", "", st.session_state.dados.get("ia_sugestao", ""))
        st.markdown(texto_visual)

        st.divider()
        st.markdown("**⚠️ Responsabilidade do Educador:** a IA pode errar. Valide e ajuste antes de aplicar.")

        if st.session_state.dados.get("status_validacao_pei") == "revisao":
            c_ok, c_ajuste = st.columns(2)
            if c_ok.button("✅ Aprovar Plano", type="primary", use_container_width=True):
                st.session_state.dados["status_validacao_pei"] = "aprovado"
                ou.track_ai_feedback("pei", "validated", content_type="relatorio_pei")
                st.success("Plano aprovado ✅")
                st.rerun()
            if c_ajuste.button("❌ Solicitar Ajuste", use_container_width=True):
                st.session_state.dados["status_validacao_pei"] = "ajustando"
                st.rerun()

        elif st.session_state.dados.get("status_validacao_pei") == "aprovado":
            st.success("Plano Validado ✅")
            novo_texto = st.text_area(
                "Edição Final Manual (opcional)",
                value=st.session_state.dados.get("ia_sugestao", ""),
                height=320,
            )
            st.session_state.dados["ia_sugestao"] = novo_texto

            c1, c2 = st.columns(2)
            with c1:
                if st.button("🔁 Regerar do Zero", use_container_width=True):
                    st.session_state.dados["ia_sugestao"] = ""
                    st.session_state.dados["status_validacao_pei"] = "rascunho"
                    st.rerun()
            with c2:
                if st.button("🧹 Voltar para Revisão", use_container_width=True):
                    st.session_state.dados["status_validacao_pei"] = "revisao"
                    st.rerun()

    # 3) Ajustando: caixa de feedback + regerar
    elif st.session_state.dados.get("status_validacao_pei") == "ajustando":
        st.warning("Descreva o ajuste desejado:")
        feedback = st.text_area("Seu feedback:", placeholder="Ex: Foque mais na alfabetização…")
        if st.button("Regerar com Ajustes", type="primary", use_container_width=True):
            ou.track_ai_feedback("pei", "refazer", content_type="relatorio_pei", feedback_text=feedback or "")
            engine = st.session_state.dados.get("consultoria_engine", "red")
            with st.spinner(ou.get_loading_message(engine)):
                res, err = consultar_gpt_pedagogico(
                    api_key,
                    st.session_state.dados,
                    st.session_state.get("pdf_text", ""),
                    modo_pratico=False,
                    feedback_usuario=feedback,
                    engine=engine,
                )
            if res:
                st.session_state.dados["ia_sugestao"] = res
                st.session_state.dados["status_validacao_pei"] = "revisao"
                st.rerun()
            else:
                st.error(err or "Erro ao regerar.")

        if st.button("Cancelar", use_container_width=True):
            st.session_state.dados["status_validacao_pei"] = "revisao"
            st.rerun()


# ==============================================================================
# 19. ABA DASHBOARD & DOCS (Dashboard + Metas + Exportações + Sincronização 'rico')
# ==============================================================================
with tab9:
    render_progresso()
    st.markdown("### <i class='ri-file-pdf-line'></i> Dashboard e Exportação", unsafe_allow_html=True)

    # --------------------------------------------------------------------------
    # 0) GARANTIR CSS DO DASH
    # --------------------------------------------------------------------------
    def _ensure_dashboard_css():
        css = """
        <style>
            .dash-hero { background: linear-gradient(135deg, #0F52BA 0%, #062B61 100%); border-radius: 16px; padding: 25px; color: white; margin-bottom: 20px; display: flex; justify-content: space-between; align-items: center; box-shadow: 0 4px 12px rgba(15, 82, 186, 0.15); }
            .apple-avatar { width: 60px; height: 60px; border-radius: 50%; background: rgba(255,255,255,0.15); border: 2px solid rgba(255,255,255,0.4); color: white; font-weight: 800; font-size: 1.6rem; display: flex; align-items: center; justify-content: center; }
            .metric-card { background: white; border-radius: 16px; padding: 15px; border: 1px solid #E2E8F0; display: flex; flex-direction: column; align-items: center; justify-content: center; height: 140px; box-shadow: 0 2px 5px rgba(0,0,0,0.02); }
            .css-donut { --p: 0; --fill: #e5e7eb; width: 80px; height: 80px; border-radius: 50%; background: conic-gradient(var(--fill) var(--p), #F3F4F6 0); position: relative; display: flex; align-items: center; justify-content: center; margin-bottom: 10px; }
            .css-donut:after { content: ""; position: absolute; width: 60px; height: 60px; border-radius: 50%; background: white; }
            .d-val { position: relative; z-index: 10; font-weight: 800; font-size: 1.2rem; color: #2D3748; }
            .d-lbl { font-size: 0.75rem; font-weight: 700; color: #718096; text-transform: uppercase; letter-spacing: 0.5px; text-align:center; }
            .comp-icon-box { width: 50px; height: 50px; border-radius: 50%; background: rgba(255,255,255,0.2); display: flex; align-items: center; justify-content: center; margin-bottom: 10px; }
            .soft-card { border-radius: 12px; padding: 20px; min-height: 220px; height: 100%; display: flex; flex-direction: column; box-shadow: 0 2px 5px rgba(0,0,0,0.02); border: 1px solid rgba(0,0,0,0.05); border-left: 5px solid; position: relative; overflow: hidden; }
            .sc-orange { background-color: #FFF5F5; border-left-color: #DD6B20; }
            .sc-blue { background-color: #EBF8FF; border-left-color: #3182CE; }
            .sc-yellow { background-color: #FFFFF0; border-left-color: #D69E2E; }
            .sc-cyan { background-color: #E6FFFA; border-left-color: #0BC5EA; }
            .sc-green { background-color: #F0FFF4; border-left-color: #38A169; }
            .sc-head { display: flex; align-items: center; gap: 8px; font-weight: 800; font-size: 0.95rem; margin-bottom: 15px; color: #2D3748; }
            .sc-body { font-size: 0.85rem; color: #4A5568; line-height: 1.5; flex-grow: 1; }
            .bg-icon { position: absolute; bottom: -10px; right: -10px; font-size: 5rem; opacity: 0.08; pointer-events: none; }
            .meta-row { display: flex; align-items: center; gap: 10px; margin-bottom: 8px; font-size: 0.85rem; border-bottom: 1px solid rgba(0,0,0,0.05); padding-bottom: 5px; }
            .dna-bar-container { margin-bottom: 15px; }
            .dna-bar-flex { display: flex; justify-content: space-between; font-size: 0.8rem; margin-bottom: 3px; font-weight: 600; color: #4A5568; }
            .dna-bar-bg { width: 100%; height: 8px; background-color: #E2E8F0; border-radius: 4px; overflow: hidden; }
            .dna-bar-fill { height: 100%; border-radius: 4px; transition: width 1s ease; }
            .rede-chip { display: inline-flex; align-items: center; gap: 5px; background: white; padding: 5px 12px; border-radius: 20px; font-size: 0.85rem; font-weight: 600; color: #2D3748; box-shadow: 0 1px 2px rgba(0,0,0,0.05); border: 1px solid #E2E8F0; margin: 0 5px 5px 0; }
            .pulse-alert { animation: pulse 2s infinite; color: #E53E3E; font-weight: bold; }
            @keyframes pulse { 0% { opacity: 1; } 50% { opacity: 0.5; } 100% { opacity: 1; } }
        </style>
        """
        st.markdown(css, unsafe_allow_html=True)

    _ensure_dashboard_css()

    # --------------------------------------------------------------------------
    # 1) HELPERS (fallbacks)
    # --------------------------------------------------------------------------
    d = st.session_state.dados

    def _safe(fn_name, default=None):
        return globals().get(fn_name, default)

    calcular_idade_fn = _safe("calcular_idade", lambda x: "")
    get_hiperfoco_emoji_fn = _safe("get_hiperfoco_emoji", lambda x: "🚀")
    calcular_complexidade_pei_fn = _safe("calcular_complexidade_pei", lambda _d: ("ATENÇÃO", "#FFFFF0", "#D69E2E"))
    extrair_metas_estruturadas_fn = _safe("extrair_metas_estruturadas", lambda _t: {"Curto": "Definir...", "Medio": "Definir...", "Longo": "Definir..."})
    inferir_componentes_impactados_fn = _safe("inferir_componentes_impactados", lambda _d: [])
    get_pro_icon_fn = _safe("get_pro_icon", lambda _p: "👨‍⚕️")

    # --------------------------------------------------------------------------
    # 2) GUARD
    # --------------------------------------------------------------------------
    if not d.get("nome"):
        st.info("Preencha o estudante na aba **Estudante** para visualizar o dashboard.")
        st.stop()

    # --------------------------------------------------------------------------
    # 3) HERO
    # --------------------------------------------------------------------------
    init_avatar = d.get("nome", "?")[0].upper() if d.get("nome") else "?"
    idade_str = calcular_idade_fn(d.get("nasc"))
    serie_txt = d.get("serie") or "-"
    turma_txt = d.get("turma") or "-"
    matricula_txt = d.get("matricula") or d.get("ra") or "-"
    student_id = st.session_state.get("selected_student_id")
    vinculo_txt = "Vinculado ao Supabase ✅" if student_id else "Rascunho (não sincronizado)"

    st.markdown(
        f"""
        <div class="dash-hero">
            <div style="display:flex; align-items:center; gap:20px;">
                <div class="apple-avatar">{init_avatar}</div>
                <div style="color:white;">
                    <h1 style="margin:0; line-height:1.1;">{d.get("nome","")}</h1>
                    <p style="margin:6px 0 0 0; opacity:.9;">
                        {serie_txt} • Turma {turma_txt} • Matrícula/RA: {matricula_txt}
                    </p>
                    <p style="margin:6px 0 0 0; opacity:.8; font-size:.85rem;">{vinculo_txt}</p>
                </div>
            </div>
            <div style="text-align:right;">
                <div style="font-size:0.8rem; opacity:.85;">IDADE</div>
                <div style="font-size:1.2rem; font-weight:800;">{idade_str}</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )

    # --------------------------------------------------------------------------
    # 4) KPIs
    # --------------------------------------------------------------------------
    c_kpi1, c_kpi2, c_kpi3, c_kpi4 = st.columns(4)

    with c_kpi1:
        n_pot = len(d.get("potencias", []) or [])
        color_p = "#38A169" if n_pot > 0 else "#CBD5E0"
        st.markdown(
            f"""<div class="metric-card">
                <div class="css-donut" style="--p: {min(n_pot*10,100)}%; --fill: {color_p};">
                    <div class="d-val">{n_pot}</div>
                </div>
                <div class="d-lbl">Potencialidades</div>
            </div>""",
            unsafe_allow_html=True
        )

    with c_kpi2:
        barreiras = d.get("barreiras_selecionadas", {}) or {}
        n_bar = sum(len(v) for v in barreiras.values()) if isinstance(barreiras, dict) else 0
        color_b = "#E53E3E" if n_bar > 5 else "#DD6B20"
        st.markdown(
            f"""<div class="metric-card">
                <div class="css-donut" style="--p: {min(n_bar*5,100)}%; --fill: {color_b};">
                    <div class="d-val">{n_bar}</div>
                </div>
                <div class="d-lbl">Barreiras</div>
            </div>""",
            unsafe_allow_html=True
        )

    with c_kpi3:
        hf = d.get("hiperfoco") or "-"
        hf_emoji = get_hiperfoco_emoji_fn(hf)
        st.markdown(
            f"""<div class="metric-card">
                <div style="font-size:2.5rem;">{hf_emoji}</div>
                <div style="font-weight:800; font-size:1.1rem; color:#2D3748; margin:10px 0;">{hf}</div>
                <div class="d-lbl">Hiperfoco</div>
            </div>""",
            unsafe_allow_html=True
        )

    with c_kpi4:
        txt_comp, bg_c, txt_c = calcular_complexidade_pei_fn(d)
        st.markdown(
            f"""<div class="metric-card" style="background-color:{bg_c}; border-color:{txt_c};">
                <div class="comp-icon-box">
                    <i class="ri-error-warning-line" style="color:{txt_c}; font-size: 2rem;"></i>
                </div>
                <div style="font-weight:800; font-size:1.1rem; color:{txt_c}; margin:5px 0;">{txt_comp}</div>
                <div class="d-lbl" style="color:{txt_c};">Nível de Atenção (Execução)</div>
            </div>""",
            unsafe_allow_html=True
        )

    # --------------------------------------------------------------------------
    # 5) CARDS PRINCIPAIS (2 colunas)
    # --------------------------------------------------------------------------
    st.write("")
    c_r1, c_r2 = st.columns(2)

    with c_r1:
        lista_meds = d.get("lista_medicamentos", []) or []
        if len(lista_meds) > 0:
            nomes_meds = ", ".join([m.get("nome","").strip() for m in lista_meds if m.get("nome")])
            alerta_escola = any(bool(m.get("escola")) for m in lista_meds)

            icon_alerta = '<i class="ri-alarm-warning-fill pulse-alert" style="font-size:1.2rem; margin-left:10px;"></i>' if alerta_escola else ""
            msg_escola = '<div style="margin-top:5px; color:#C53030; font-weight:bold; font-size:0.8rem;">🚨 ATENÇÃO: ADMINISTRAÇÃO NA ESCOLA NECESSÁRIA</div>' if alerta_escola else ""

            st.markdown(
                f"""<div class="soft-card sc-orange">
                    <div class="sc-head"><i class="ri-medicine-bottle-fill" style="color:#DD6B20;"></i> Atenção Farmacológica {icon_alerta}</div>
                    <div class="sc-body"><b>Uso Contínuo:</b> {nomes_meds if nomes_meds else "Medicação cadastrada."} {msg_escola}</div>
                    <div class="bg-icon">💊</div>
                </div>""",
                unsafe_allow_html=True
            )
        else:
            st.markdown(
                """<div class="soft-card sc-green">
                    <div class="sc-head"><i class="ri-checkbox-circle-fill" style="color:#38A169;"></i> Medicação</div>
                    <div class="sc-body">Nenhuma medicação informada.</div>
                    <div class="bg-icon">✅</div>
                </div>""",
                unsafe_allow_html=True
            )

        st.write("")

        metas = extrair_metas_estruturadas_fn(d.get("ia_sugestao", ""))
        html_metas = (
            f"""<div class="meta-row"><span style="font-size:1.2rem;">🏁</span> <b>Curto:</b> {metas.get('Curto','Definir...')}</div>
                <div class="meta-row"><span style="font-size:1.2rem;">🧗</span> <b>Médio:</b> {metas.get('Medio','Definir...')}</div>
                <div class="meta-row"><span style="font-size:1.2rem;">🏔️</span> <b>Longo:</b> {metas.get('Longo','Definir...')}</div>"""
        )
        st.markdown(
            f"""<div class="soft-card sc-yellow">
                <div class="sc-head"><i class="ri-flag-2-fill" style="color:#D69E2E;"></i> Cronograma de Metas</div>
                <div class="sc-body">{html_metas}</div>
                <div class="bg-icon">🏁</div>
            </div>""",
            unsafe_allow_html=True
        )

    with c_r2:
        comps_inferidos = inferir_componentes_impactados_fn(d) or []
        if comps_inferidos:
            html_comps = "".join([f'<span class="rede-chip" style="border-color:#FC8181; color:#C53030;">{c}</span> ' for c in comps_inferidos])
            st.markdown(
                f"""<div class="soft-card sc-orange" style="border-left-color: #FC8181; background-color: #FFF5F5;">
                    <div class="sc-head"><i class="ri-radar-fill" style="color:#C53030;"></i> Radar Curricular (Automático)</div>
                    <div class="sc-body" style="margin-bottom:10px;">Componentes que exigem maior flexibilização (baseado nas barreiras):</div>
                    <div>{html_comps}</div>
                    <div class="bg-icon">🎯</div>
                </div>""",
                unsafe_allow_html=True
            )
        else:
            st.markdown(
                """<div class="soft-card sc-blue">
                    <div class="sc-head"><i class="ri-radar-line" style="color:#3182CE;"></i> Radar Curricular</div>
                    <div class="sc-body">Nenhum componente específico marcado como crítico.</div>
                    <div class="bg-icon">🎯</div>
                </div>""",
                unsafe_allow_html=True
            )

        st.write("")

        rede = d.get("rede_apoio", []) or []
        rede_html = "".join([f'<span class="rede-chip">{get_pro_icon_fn(p)} {p}</span> ' for p in rede]) if rede else "<span style='opacity:0.6;'>Sem rede.</span>"
        st.markdown(
            f"""<div class="soft-card sc-cyan">
                <div class="sc-head"><i class="ri-team-fill" style="color:#0BC5EA;"></i> Rede de Apoio</div>
                <div class="sc-body">{rede_html}</div>
                <div class="bg-icon">🤝</div>
            </div>""",
            unsafe_allow_html=True
        )

    # --------------------------------------------------------------------------
    # 6) DNA de Suporte
    # --------------------------------------------------------------------------
    st.write("")
    st.markdown("##### 🧬 DNA de Suporte")
    dna_c1, dna_c2 = st.columns(2)

    LISTAS_BARREIRAS_LOCAL = globals().get("LISTAS_BARREIRAS", {}) or {}
    areas = list(LISTAS_BARREIRAS_LOCAL.keys()) if isinstance(LISTAS_BARREIRAS_LOCAL, dict) else []

    for i, area in enumerate(areas):
        qtd = len((d.get("barreiras_selecionadas", {}) or {}).get(area, []) or [])
        val = min(qtd * 20, 100)
        target = dna_c1 if i < 3 else dna_c2

        color = "#3182CE"
        if val > 40: color = "#DD6B20"
        if val > 70: color = "#E53E3E"

        target.markdown(
            f"""<div class="dna-bar-container">
                <div class="dna-bar-flex"><span>{area}</span><span>{qtd} barreiras</span></div>
                <div class="dna-bar-bg"><div class="dna-bar-fill" style="width:{val}%; background:{color};"></div></div>
            </div>""",
            unsafe_allow_html=True
        )

    # --------------------------------------------------------------------------
    # 7) EXPORTAÇÃO + SINCRONIZAÇÃO (BLOCO COMPLETO CORRIGIDO)
    # --------------------------------------------------------------------------
    st.divider()
    st.markdown("#### 📤 Exportação e Sincronização")

    # ======================================================================
    # EXPORTAÇÃO E SINCRONIZAÇÃO (SEMPRE DISPONÍVEL)
    # ======================================================================
    col_docs, col_backup, col_sys = st.columns(3) 

    # ---------------- COLUNA 1: DOCS ----------------
    with col_docs:
        st.caption("📄 Documentos")
        
        if not d.get("nome"):
            st.info("Preencha o nome do estudante para gerar documentos.")
        else:
            pdf_bytes = None
            try:
                pdf_bytes = gerar_pdf_final(d)
            except Exception as e:
                st.error(f"Erro ao gerar PDF: {e}")

            if pdf_bytes:
                st.download_button(
                    "📄 Baixar PDF Oficial",
                    pdf_bytes,
                    f"PEI_{d.get('nome','Estudante')}.pdf",
                    "application/pdf",
                    use_container_width=True,
                    type="primary"
                )
            else:
                st.warning("Não foi possível gerar o PDF.")

            try:
                docx = gerar_docx_final(d)
                st.download_button(
                    "📝 Baixar Word Editável",
                    docx,
                    f"PEI_{d.get('nome','Estudante')}.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            except Exception as e:
                st.warning(f"Word indisponível: {str(e)}")
            
            if not d.get("ia_sugestao"):
                st.caption("💡 Gere o Plano na aba **Consultoria IA** para incluir o planejamento pedagógico detalhado no documento.")

        # ---------------- COLUNA 2: BACKUP LOCAL ----------------
        with col_backup:
            st.caption("💾 Backup (JSON)")
            st.markdown(
                "<div style='font-size:.8rem; color:#64748B; margin-bottom:8px;'>"
                "Salva um arquivo no seu computador para garantir que nada se perca."
                "</div>",
                unsafe_allow_html=True
            )
            st.download_button(
                "Salvar Arquivo .JSON",
                json.dumps(d, default=str, ensure_ascii=False),
                f"PEI_{d.get('nome','Estudante')}.json",
                "application/json",
                use_container_width=True
            )

        # ---------------- COLUNA 3: NUVEM (SUPABASE COMPLETO) ----------------
        with col_sys:
            st.caption("🌐 Nuvem (Supabase)")
            st.markdown(
                "<div style='font-size:.8rem; color:#64748B; margin-bottom:8px;'>"
                "Salva cadastro + conteúdo completo (JSON) na nuvem."
                "</div>",
                unsafe_allow_html=True
            )

            # Helper interno de verificação
            def _cloud_ready_check():
                try:
                    url = str(ou.get_setting("SUPABASE_URL", "")).strip()
                    key = str(
                        ou.get_setting("SUPABASE_SERVICE_KEY", "")
                        or ou.get_setting("SUPABASE_ANON_KEY", "")
                    ).strip()
                    return bool(url and key)
                except:
                    return False

            if st.button("🔗 Sincronizar Tudo", type="primary", use_container_width=True, key="btn_sync_final_fix"):
                if not _cloud_ready_check():
                    st.error("⚠️ Configure os Secrets do Supabase.")
                else:
                    try:
                        with st.spinner("Sincronizando dados..."):
                            # 1. Tratar datas
                            nasc_iso = d.get("nasc").isoformat() if hasattr(d.get("nasc"), "isoformat") else None
                            
                            # 2. Dados Básicos
                            student_payload = {
                                "name": d.get("nome"),
                                "birth_date": nasc_iso,
                                "grade": d.get("serie"),
                                "class_group": d.get("turma") or None,
                                "diagnosis": d.get("diagnostico") or None,
                                "workspace_id": st.session_state.get("workspace_id"),
                            }
                            
                            # 3. Identificar ou Criar
                            sid = st.session_state.get("selected_student_id")
                            
                            if not sid:
                                created = db_create_student(student_payload)
                                if created and isinstance(created, dict):
                                    sid = created.get("id")
                                    st.session_state["selected_student_id"] = sid
                                    st.session_state["students_cache_invalid"] = True  # outras páginas recarregam a lista
                            else:
                                db_update_student(sid, student_payload)

                            # 4. SALVAR CONTEÚDO COMPLETO
                            if sid:
                                # Certifique-se de ter colado a função 'db_update_pei_content' lá em cima nas funções!
                                db_update_pei_content(sid, d)
                                
                                st.session_state["ultimo_backup_json"] = json.dumps(d, default=str, ensure_ascii=False)
                                st.session_state["sync_sucesso"] = True
                                
                                st.toast("Salvo na nuvem com sucesso!", icon="☁️")
                                st.rerun()
                            else:
                                st.error("Erro: ID do estudante não encontrado.")

                    except Exception as e:
                        st.error(f"Erro na sincronização: {e}")

            # Feedback e Download Pós-Sync
            if st.session_state.get("sync_sucesso"):
                st.success("✅ Tudo salvo!")
                timestamp = datetime.now().strftime("%d-%m_%Hh%M")
                nome_clean = (d.get('nome') or 'Estudante').replace(' ', '_')
                
                st.download_button(
                    label="📂 BAIXAR CÓPIA AGORA",
                    data=st.session_state["ultimo_backup_json"],
                    file_name=f"PEI_{nome_clean}_{timestamp}.json",
                    mime="application/json",
                    type="secondary",
                    use_container_width=True,
                    key="btn_post_sync_download_fix"
                )

# ==============================================================================
# ABA RETRÁTIL — O QUE ESTÁ REGISTRADO PARA O ESTUDANTE (PARTE DE BAIXO)
# ==============================================================================
dados = st.session_state.get("dados", {})
nome_pei = (dados.get("nome") or "").strip()
if nome_pei:
    tem_relatorio_pei = bool((dados.get("ia_sugestao") or "").strip())
    tem_jornada_pei = bool((dados.get("ia_mapa_texto") or "").strip())
    ou.render_resumo_anexos_estudante(
        nome_estudante=nome_pei,
        tem_relatorio_pei=tem_relatorio_pei,
        tem_jornada=tem_jornada_pei,
        n_ciclos_pae=0,
        pagina="PEI",
    )

# ==============================================================================
# RODAPÉ COM ASSINATURA
# ==============================================================================
ou.render_footer_assinatura()
