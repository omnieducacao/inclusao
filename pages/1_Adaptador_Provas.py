import streamlit as st
from openai import OpenAI
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from streamlit_cropper import st_cropper
import re
import requests
import json
import base64
import os

# --- 1. CONFIGURAÇÃO ---
st.set_page_config(page_title="Adaptador 360º | V18.0", page_icon="🧩", layout="wide")

# --- 2. BANCO DE DADOS ---
ARQUIVO_DB = "banco_alunos.json"

def carregar_banco():
    if os.path.exists(ARQUIVO_DB):
        try:
            with open(ARQUIVO_DB, "r", encoding="utf-8") as f:
                return json.load(f)
        except: return []
    return []

if 'banco_estudantes' not in st.session_state or not st.session_state.banco_estudantes:
    st.session_state.banco_estudantes = carregar_banco()

# --- 3. ESTILO VISUAL ---
st.markdown("""
    <style>
    html, body, [class*="css"] { font-family: 'Nunito', sans-serif; color: #2D3748; }
    .header-clean { background: white; padding: 25px; border-radius: 16px; border: 1px solid #EDF2F7; margin-bottom: 20px; display: flex; gap: 20px; align-items: center; }
    .student-header { background-color: #EBF8FF; border: 1px solid #BEE3F8; border-radius: 12px; padding: 15px 25px; margin-bottom: 25px; display: flex; justify-content: space-between; align-items: center; }
    .student-label { font-size: 0.85rem; color: #718096; font-weight: 700; text-transform: uppercase; }
    .student-value { font-size: 1.1rem; color: #2C5282; font-weight: 800; }
    
    .analise-box { background-color: #F0FFF4; border: 1px solid #C6F6D5; border-radius: 8px; padding: 20px; margin-bottom: 20px; color: #22543D; }
    .analise-title { font-weight: bold; font-size: 1.1rem; margin-bottom: 10px; display: flex; align-items: center; gap: 10px; }
    
    div[data-testid="column"] .stButton button[kind="primary"] { border-radius: 12px !important; height: 50px; width: 100%; background-color: #FF6B6B !important; color: white !important; font-weight: 800 !important; }
    div[data-testid="column"] .stButton button[kind="secondary"] { border-radius: 12px !important; height: 50px; width: 100%; border: 2px solid #CBD5E0 !important; color: #4A5568 !important; font-weight: bold; }
    
    .validado-box { background-color: #C6F6D5; color: #22543D; padding: 10px; border-radius: 8px; text-align: center; font-weight: bold; margin-top: 10px; border: 1px solid #276749; }
    </style>
""", unsafe_allow_html=True)

# --- 4. FUNÇÕES DE ARQUIVO E IMAGEM ---
def extrair_dados_docx(uploaded_file):
    uploaded_file.seek(0); imagens = []; texto = ""
    try:
        doc = Document(uploaded_file)
        texto = "\n".join([p.text for p in doc.paragraphs if p.text.strip() != ""])
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                img_data = rel.target_part.blob
                if len(img_data) > 1024: imagens.append(img_data)
    except: pass
    return texto, imagens

def sanitizar_imagem(image_bytes):
    try:
        img = Image.open(BytesIO(image_bytes)).convert("RGB")
        out = BytesIO()
        img.save(out, format="JPEG", quality=90)
        return out.getvalue()
    except: return None

def baixar_imagem_url(url):
    try:
        resp = requests.get(url, timeout=10)
        if resp.status_code == 200: return BytesIO(resp.content)
    except: pass
    return None

def buscar_imagem_unsplash(query, access_key):
    if not access_key: return None
    url = f"https://api.unsplash.com/search/photos?query={query}&per_page=1&client_id={access_key}"
    try:
        resp = requests.get(url, timeout=5)
        data = resp.json()
        if data.get('results'):
            return data['results'][0]['urls']['regular']
    except: pass
    return None

def construir_docx_final(texto_ia, aluno, materia, mapa_imgs, img_dalle_url, tipo_atv, sem_cabecalho=False):
    doc = Document(); style = doc.styles['Normal']; style.font.name = 'Arial'; style.font.size = Pt(12)
    
    if not sem_cabecalho:
        doc.add_heading(f'{tipo_atv.upper()} ADAPTADA - {materia.upper()}', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Estudante: {aluno['nome']}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("_"*50)
        doc.add_heading('Atividades', level=2)

    linhas = texto_ia.split('\n')
    for linha in linhas:
        # Regex para pegar [[IMG...]] ou [[GEN_IMG...]]
        tag_match = re.search(r'\[\[(IMG|GEN_IMG).*?(\d+)\]\]', linha, re.IGNORECASE)
        
        if tag_match:
            partes = re.split(r'(\[\[(?:IMG|GEN_IMG).*?\d+\]\])', linha, flags=re.IGNORECASE)
            for parte in partes:
                sub_match = re.search(r'(\d+)', parte)
                if ("IMG" in parte.upper() or "GEN_IMG" in parte.upper()) and sub_match:
                    num = int(sub_match.group(1))
                    img_bytes = mapa_imgs.get(num)
                    # Fallback para imagem única (chave 1)
                    if not img_bytes and len(mapa_imgs) == 1: img_bytes = list(mapa_imgs.values())[0]
                    
                    if img_bytes:
                        try:
                            doc.add_picture(BytesIO(img_bytes), width=Inches(4.5))
                            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except: pass
                elif parte.strip():
                    doc.add_paragraph(parte.strip())
        else:
            if linha.strip(): doc.add_paragraph(linha.strip())
            
    buffer = BytesIO(); doc.save(buffer); buffer.seek(0)
    return buffer

# --- 5. INTELIGÊNCIA ARTIFICIAL ---
def gerar_imagem_hibrida(api_key, prompt, unsplash_key=None):
    if unsplash_key:
        termo = prompt.split('.')[0] if '.' in prompt else prompt
        url = buscar_imagem_unsplash(termo, unsplash_key)
        if url: return url
    client = OpenAI(api_key=api_key)
    try:
        didactic_prompt = f"Educational illustration, clear diagram style, white background: {prompt}"
        resp = client.images.generate(model="dall-e-3", prompt=didactic_prompt, size="1024x1024", quality="standard", n=1)
        return resp.data[0].url
    except: return None

# FUNÇÃO AUXILIAR: GUARDIÃO DA IMAGEM
def garantir_tag_imagem(texto):
    """Se a IA esquecer a tag, insere na força bruta após o primeiro parágrafo."""
    if "[[IMG" not in texto.upper():
        # Tenta achar o fim do primeiro parágrafo ou frase
        match = re.search(r'(\n|\. )', texto)
        if match:
            pos = match.end()
            return texto[:pos] + "\n[[IMG_1]]\n" + texto[pos:]
        else:
            return texto + "\n[[IMG_1]]"
    return texto

# ADAPTAR DOCX (Com Split Robusto)
def adaptar_conteudo_docx(api_key, aluno, texto, materia, tema, tipo_atv, remover_resp, questoes_mapeadas, modo_profundo=False):
    client = OpenAI(api_key=api_key)
    lista_q = ", ".join([str(n) for n in questoes_mapeadas])
    
    instruction_style = "Pense passo a passo e faça uma adaptação profunda e didática." if modo_profundo else "Seja direto e objetivo."
    
    prompt = f"""
    ESPECIALISTA EM DUA. {instruction_style}
    1. ANALISE O PERFIL: {aluno.get('ia_sugestao', '')[:1000]}
    2. ADAPTE A PROVA: Use o hiperfoco ({aluno.get('hiperfoco', 'Geral')}) em 30% das questões.
    
    REGRA DE IMAGEM: O professor indicou imagens nas questões: {lista_q}.
    Nessas questões, use OBRIGATORIAMENTE: Enunciado -> [[IMG_número]] -> Alternativas.
    
    SAÍDA OBRIGATÓRIA (Use EXATAMENTE este divisor):
    [ANÁLISE PEDAGÓGICA]
    ...texto da análise...
    ---DIVISOR---
    [ATIVIDADE]
    ...texto da atividade...
    
    CONTEXTO: {materia} | {tema}. {"REMOVA GABARITO." if remover_resp else ""}
    TEXTO ORIGINAL: {texto}
    """
    try:
        resp = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": prompt}], temperature=0.7 if modo_profundo else 0.4)
        full_text = resp.choices[0].message.content
        if "---DIVISOR---" in full_text:
            parts = full_text.split("---DIVISOR---")
            return parts[0].replace("[ANÁLISE PEDAGÓGICA]", "").strip(), parts[1].replace("[ATIVIDADE]", "").strip()
        return "Análise indisponível.", full_text
    except Exception as e: return str(e), ""

# ADAPTAR IMAGEM (OCR + Split Robusto + Guardião)
def adaptar_conteudo_imagem(api_key, aluno, imagem_bytes, materia, tema, tipo_atv, livro_professor, modo_profundo=False):
    client = OpenAI(api_key=api_key)
    if not imagem_bytes: return "Erro: Imagem vazia", ""
    b64 = base64.b64encode(imagem_bytes).decode('utf-8')
    
    instrucao_livro = "ATENÇÃO: IMAGEM COM RESPOSTAS. Remova todo gabarito/respostas." if livro_professor else ""
    instruction_style = "Faça uma análise crítica e profunda do conteúdo." if modo_profundo else "Transcreva e adapte objetivamente."
    
    prompt = f"""
    ATUAR COMO: Especialista em Acessibilidade e OCR. {instruction_style}
    1. Transcreva o texto da imagem. {instrucao_livro}
    2. Adapte para o aluno (PEI: {aluno.get('ia_sugestao', '')[:800]}).
    3. Hiperfoco ({aluno.get('hiperfoco')}): Conecte levemente.
    4. IMPORTANTE: Insira a tag [[IMG_1]] logo após o enunciado principal. Não a esqueça.
    
    SAÍDA OBRIGATÓRIA (Respeite o divisor):
    [ANÁLISE PEDAGÓGICA]
    ...resumo das adaptações...
    ---DIVISOR---
    [ATIVIDADE]
    ...conteúdo adaptado...
    """
    msgs = [{"role": "user", "content": [{"type": "text", "text": prompt}, {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64}"}}]}]
    try:
        resp = client.chat.completions.create(model="gpt-4o-mini", messages=msgs, temperature=0.7 if modo_profundo else 0.4)
        full_text = resp.choices[0].message.content
        
        analise = "Análise indisponível."
        atividade = full_text
        
        if "---DIVISOR---" in full_text:
            parts = full_text.split("---DIVISOR---")
            analise = parts[0].replace("[ANÁLISE PEDAGÓGICA]", "").strip()
            atividade = parts[1].replace("[ATIVIDADE]", "").strip()
        
        # GUARDIÃO DA IMAGEM: Se a IA esqueceu a tag, insere na marra
        atividade = garantir_tag_imagem(atividade)
            
        return analise, atividade
    except Exception as e: return str(e), ""

# CRIAR PROFISSIONAL (Com Controle e Validação)
def criar_profissional(api_key, aluno, materia, objeto, qtd, tipo_q, pct_img, modo_profundo=False):
    client = OpenAI(api_key=api_key)
    hiperfoco = aluno.get('hiperfoco', 'Geral')
    
    qtd_imgs = int(qtd * (pct_img / 100))
    instrucao_img = f"Incluir {qtd_imgs} imagens (use [[GEN_IMG: termo]])." if qtd_imgs > 0 else "Sem imagens."
    instruction_style = "Pense como uma banca examinadora rigorosa. Crie itens perfeitos." if modo_profundo else "Crie itens padrão BNCC."
    
    prompt = f"""
    VOCÊ É UMA BANCA EXAMINADORA. {instruction_style}
    Crie prova de {materia} ({objeto}). QTD: {qtd} ({tipo_q}).
    
    DIRETRIZES:
    1. **Contexto:** Use situações-problema reais.
    2. **Hiperfoco ({hiperfoco}):** Integre à lógica de 30% das questões.
    3. **Distratores:** Inteligentes e plausíveis.
    4. **Imagens:** {instrucao_img} (NUNCA repita a mesma descrição de imagem).
    5. **Divisão:** Separe CLARAMENTE a análise.
    
    SAÍDA OBRIGATÓRIA:
    [ANÁLISE PEDAGÓGICA]
    ...análise...
    ---DIVISOR---
    [ATIVIDADE]
    ...questões...
    """
    try:
        resp = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": prompt}], temperature=0.8 if modo_profundo else 0.7)
        full_text = resp.choices[0].message.content
        if "---DIVISOR---" in full_text:
            parts = full_text.split("---DIVISOR---")
            return parts[0].replace("[ANÁLISE PEDAGÓGICA]", "").strip(), parts[1].replace("[ATIVIDADE]", "").strip()
        return "Análise indisponível.", full_text
    except Exception as e: return str(e), ""

# OUTRAS (Mantidas)
def gerar_roteiro_aula(api_key, aluno, assunto):
    client = OpenAI(api_key=api_key)
    prompt = f"Roteiro de aula {assunto} para {aluno['nome']}. PEI: {aluno.get('ia_sugestao','')[:500]}."
    try:
        resp = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": prompt}], temperature=0.7)
        return resp.choices[0].message.content
    except Exception as e: return str(e)

def gerar_quebra_gelo_profundo(api_key, aluno, assunto, tema_extra=""):
    client = OpenAI(api_key=api_key)
    tema = tema_extra if tema_extra else aluno.get('hiperfoco', 'Geral')
    prompt = f"3 Tópicos profundos conectando {tema} e {assunto} para {aluno['nome']}."
    try:
        resp = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": prompt}], temperature=0.8)
        return resp.choices[0].message.content
    except Exception as e: return str(e)

def gerar_dinamica_inclusiva(api_key, aluno, assunto):
    client = OpenAI(api_key=api_key)
    prompt = f"Dinâmica de grupo sobre {assunto} inclusiva para {aluno['nome']} (PEI: {aluno.get('ia_sugestao','')[:500]})."
    try:
        resp = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": prompt}], temperature=0.7)
        return resp.choices[0].message.content
    except Exception as e: return str(e)

def sugerir_imagem_pei(api_key, aluno):
    client = OpenAI(api_key=api_key)
    prompt = f"Prompt DALL-E para recurso visual: {aluno.get('ia_sugestao','')[:500]}."
    try:
        resp = client.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": prompt}], temperature=0.7)
        return resp.choices[0].message.content
    except: return "Educational illustration"

# --- 6. INTERFACE ---
with st.sidebar:
    if 'OPENAI_API_KEY' in st.secrets: api_key = st.secrets['OPENAI_API_KEY']; st.success("✅ OpenAI OK")
    else: api_key = st.text_input("Chave OpenAI:", type="password")
    st.markdown("---")
    if 'UNSPLASH_ACCESS_KEY' in st.secrets: unsplash_key = st.secrets['UNSPLASH_ACCESS_KEY']; st.success("✅ Unsplash OK")
    else: unsplash_key = st.text_input("Chave Unsplash (Opcional):", type="password")
    
    st.markdown("---")
    if st.button("🧹 Limpar Tudo e Reiniciar", type="secondary"):
        for key in list(st.session_state.keys()):
            if key not in ['banco_estudantes', 'OPENAI_API_KEY', 'UNSPLASH_ACCESS_KEY']: del st.session_state[key]
        st.rerun()

st.markdown("""<div class="header-clean"><div style="font-size:3rem;">🧩</div><div><p style="margin:0;color:#004E92;font-size:1.5rem;font-weight:800;">Adaptador V18.0: Validação Pro</p></div></div>""", unsafe_allow_html=True)

if not st.session_state.banco_estudantes:
    st.warning("⚠️ Cadastre um aluno no PEI 360º primeiro.")
    st.stop()

lista = [a['nome'] for a in st.session_state.banco_estudantes]
nome_aluno = st.selectbox("📂 Selecione o Estudante:", lista)
aluno = next(a for a in st.session_state.banco_estudantes if a['nome'] == nome_aluno)

st.markdown(f"""
    <div class="student-header">
        <div class="student-info-item"><div class="student-label">Nome</div><div class="student-value">{aluno.get('nome')}</div></div>
        <div class="student-info-item"><div class="student-label">Série</div><div class="student-value">{aluno.get('serie', '-')}</div></div>
        <div class="student-info-item"><div class="student-label">Hiperfoco</div><div class="student-value">{aluno.get('hiperfoco', '-')}</div></div>
    </div>
""", unsafe_allow_html=True)

tabs = st.tabs(["📄 Adaptar Prova", "✂️ Adaptar Atividade", "✨ Criar do Zero", "🎨 Estúdio Visual", "📝 Roteiro de Aula", "🗣️ Papo de Mestre", "🤝 Dinâmica Inclusiva"])

# 1. ADAPTAR PROVA
with tabs[0]:
    c1, c2, c3 = st.columns(3)
    materia_d = c1.selectbox("Matéria", ["Matemática", "Português", "Ciências", "História", "Geografia", "Artes", "Ed. Física", "Inglês"], key="dm")
    tema_d = c2.text_input("Tema", placeholder="Ex: Frações", key="dt")
    tipo_d = c3.selectbox("Tipo", ["Prova", "Tarefa"], key="dtp")
    arquivo_d = st.file_uploader("Upload DOCX", type=["docx"], key="fd")
    
    if 'docx_imgs' not in st.session_state: st.session_state.docx_imgs = []
    if 'docx_txt' not in st.session_state: st.session_state.docx_txt = None
    
    if arquivo_d and arquivo_d.file_id != st.session_state.get('last_d'):
        st.session_state.last_d = arquivo_d.file_id
        txt, imgs = extrair_dados_docx(arquivo_d)
        st.session_state.docx_txt = txt; st.session_state.docx_imgs = imgs
        st.success(f"{len(imgs)} imagens encontradas.")

    map_d = {}; qs_d = []
    if st.session_state.docx_imgs:
        st.write("### Mapeamento")
        cols = st.columns(3)
        for i, img in enumerate(st.session_state.docx_imgs):
            with cols[i % 3]:
                st.image(img, width=80)
                q = st.number_input(f"Questão:", 0, 50, key=f"dq_{i}")
                if q > 0: map_d[int(q)] = img; qs_d.append(int(q))

    if st.button("🚀 ADAPTAR PROVA", type="primary", key="btn_d"):
        if not st.session_state.docx_txt: st.warning("Envie arquivo."); st.stop()
        with st.spinner("Analisando e Adaptando..."):
            rac, txt = adaptar_conteudo_docx(api_key, aluno, st.session_state.docx_txt, materia_d, tema_d, tipo_d, True, qs_d)
            st.session_state['res_docx'] = {'rac': rac, 'txt': txt, 'map': map_d, 'valid': False}
            st.rerun()

    if 'res_docx' in st.session_state:
        res = st.session_state['res_docx']
        
        # BARRA DE AÇÃO (VALIDAR / REFAZER)
        if res.get('valid'):
            st.markdown("<div class='validado-box'>✅ ATIVIDADE VALIDADA E PRONTA!</div>", unsafe_allow_html=True)
        else:
            col_v, col_r = st.columns([1, 1])
            if col_v.button("✅ Validar Atividade", key="val_d"):
                st.session_state['res_docx']['valid'] = True
                st.rerun()
            if col_r.button("🧠 Refazer (Modo Profundo)", key="redo_d"):
                with st.spinner("Refazendo com análise aprofundada..."):
                    rac, txt = adaptar_conteudo_docx(api_key, aluno, st.session_state.docx_txt, materia_d, tema_d, tipo_d, True, qs_d, modo_profundo=True)
                    st.session_state['res_docx'] = {'rac': rac, 'txt': txt, 'map': map_d, 'valid': False}
                    st.rerun()

        st.markdown(f"<div class='analise-box'><div class='analise-title'>🧠 Análise Pedagógica</div>{res['rac']}</div>", unsafe_allow_html=True)
        with st.container(border=True):
            partes = re.split(r'(\[\[IMG.*?\d+\]\])', res['txt'], flags=re.IGNORECASE)
            for p in partes:
                tag = re.search(r'(\d+)', p)
                if "IMG" in p.upper() and tag:
                    i = int(tag.group(1))
                    im = res['map'].get(i)
                    if im: st.image(im, width=300)
                elif p.strip(): st.markdown(p.strip())
        docx = construir_docx_final(res['txt'], aluno, materia_d, res['map'], None, tipo_d)
        st.download_button("📥 BAIXAR DOCX (Só Atividade)", docx, "Prova_Adaptada.docx", "primary")

# 2. ADAPTAR ATIVIDADE
with tabs[1]:
    st.info("Adaptar de Imagem/Foto.")
    c1, c2, c3 = st.columns(3)
    discip = ["Matemática", "Português", "Ciências", "História", "Geografia", "Artes", "Ed. Física", "Inglês", "Filosofia", "Sociologia"]
    materia_i = c1.selectbox("Matéria", discip, key="im")
    tema_i = c2.text_input("Tema", key="it")
    tipo_i = c3.selectbox("Tipo", ["Atividade", "Tarefa"], key="itp")
    arquivo_i = st.file_uploader("Upload", type=["png","jpg","jpeg"], key="fi")
    livro_prof = st.checkbox("📖 É Livro do Professor? (Limpar Gabarito)", value=False)
    
    if 'img_raw' not in st.session_state: st.session_state.img_raw = None
    if arquivo_i and arquivo_i.file_id != st.session_state.get('last_i'):
        st.session_state.last_i = arquivo_i.file_id
        st.session_state.img_raw = sanitizar_imagem(arquivo_i.getvalue())

    cropped_res = None
    if st.session_state.img_raw:
        st.markdown("### ✂️ Recorte")
        img_pil = Image.open(BytesIO(st.session_state.img_raw))
        img_pil.thumbnail((800, 800))
        cropped_res = st_cropper(img_pil, realtime_update=True, box_color='#FF0000', aspect_ratio=None, key="crop_i")
        if cropped_res: st.image(cropped_res, width=200, caption="Prévia do Recorte")

    if st.button("🚀 ADAPTAR RECORTE", type="primary", key="btn_i"):
        if not st.session_state.img_raw: st.warning("Envie imagem."); st.stop()
        with st.spinner("Analisando Perfil e BNCC..."):
            buf_c = BytesIO()
            cropped_res.convert('RGB').save(buf_c, format="JPEG", quality=90)
            img_bytes = buf_c.getvalue()
            
            rac, txt = adaptar_conteudo_imagem(api_key, aluno, st.session_state.img_raw, materia_i, tema_i, tipo_i, livro_prof)
            st.session_state['res_img'] = {'rac': rac, 'txt': txt, 'map': {1: img_bytes}, 'valid': False}
            st.rerun()

    if 'res_img' in st.session_state:
        res = st.session_state['res_img']
        
        if res.get('valid'):
            st.markdown("<div class='validado-box'>✅ ATIVIDADE VALIDADA E PRONTA!</div>", unsafe_allow_html=True)
        else:
            col_v, col_r = st.columns([1, 1])
            if col_v.button("✅ Validar Atividade", key="val_i"):
                st.session_state['res_img']['valid'] = True
                st.rerun()
            if col_r.button("🧠 Refazer (Modo Profundo)", key="redo_i"):
                with st.spinner("Refazendo com análise profunda..."):
                    img_bytes = res['map'][1]
                    rac, txt = adaptar_conteudo_imagem(api_key, aluno, st.session_state.img_raw, materia_i, tema_i, tipo_i, livro_prof, modo_profundo=True)
                    st.session_state['res_img'] = {'rac': rac, 'txt': txt, 'map': {1: img_bytes}, 'valid': False}
                    st.rerun()

        st.markdown(f"<div class='analise-box'><div class='analise-title'>🧠 Análise Pedagógica</div>{res['rac']}</div>", unsafe_allow_html=True)
        
        with st.container(border=True):
            partes = re.split(r'(\[\[IMG.*?\]\])', res['txt'], flags=re.IGNORECASE)
            for p in partes:
                if "IMG" in p.upper():
                    im = res['map'].get(1)
                    if im: st.image(im, width=300)
                elif p.strip(): st.markdown(p.strip())
        docx = construir_docx_final(res['txt'], aluno, materia_i, res['map'], None, tipo_i)
        st.download_button("📥 BAIXAR DOCX (Só Atividade)", docx, "Atividade.docx", "primary")

# 3. CRIAR DO ZERO
with tabs[2]:
    st.info("Criação Profissional.")
    cc1, cc2 = st.columns(2)
    mat_c = cc1.selectbox("Componente", discip, key="cm")
    obj_c = cc2.text_input("Assunto", key="co")
    
    cc3, cc4 = st.columns(2)
    qtd_c = cc3.slider("Qtd", 1, 10, 5, key="cq")
    tipo_quest = cc4.selectbox("Tipo", ["Objetiva", "Discursiva", "Mista"], key="ctq")
    
    col_img_opt, col_img_pct = st.columns([1, 2])
    usar_img = col_img_opt.checkbox("📸 Incluir Imagens?", value=True)
    pct_img = col_img_pct.slider("Porcentagem com Imagem", 0, 100, 30, disabled=not usar_img)
    
    if st.button("✨ CRIAR ATIVIDADE", type="primary", key="btn_c"):
        with st.spinner("Elaborando e Analisando..."):
            pct_final = pct_img if usar_img else 0
            rac, txt = criar_profissional(api_key, aluno, mat_c, obj_c, qtd_c, tipo_quest, pct_final)
            
            novo_map = {}; count = 0
            tags = re.findall(r'\[\[GEN_IMG: (.*?)\]\]', txt)
            for p in tags:
                count += 1
                url = gerar_imagem_hibrida(api_key, p, unsplash_key)
                if url:
                    io = baixar_imagem_url(url)
                    if io: novo_map[count] = io.getvalue()
            txt_fin = txt
            for i in range(1, count + 1): 
                txt_fin = re.sub(r'\[\[GEN_IMG: .*?\]\]', f"[[IMG_G{i}]]", txt_fin, count=1)
            st.session_state['res_create'] = {'rac': rac, 'txt': txt_fin, 'map': novo_map, 'valid': False}
            st.rerun()

    if 'res_create' in st.session_state:
        res = st.session_state['res_create']
        
        if res.get('valid'):
            st.markdown("<div class='validado-box'>✅ ATIVIDADE VALIDADA E PRONTA!</div>", unsafe_allow_html=True)
        else:
            col_v, col_r = st.columns([1, 1])
            if col_v.button("✅ Validar Atividade", key="val_c"):
                st.session_state['res_create']['valid'] = True
                st.rerun()
            if col_r.button("🧠 Refazer (Modo Profundo)", key="redo_c"):
                with st.spinner("Refazendo com novos parâmetros..."):
                    pct_final = pct_img if usar_img else 0
                    rac, txt = criar_profissional(api_key, aluno, mat_c, obj_c, qtd_c, tipo_quest, pct_final, modo_profundo=True)
                    st.session_state['res_create']['rac'] = rac
                    st.session_state['res_create']['txt'] = txt
                    st.session_state['res_create']['valid'] = False
                    st.rerun()

        st.markdown(f"<div class='analise-box'><div class='analise-title'>🧠 Análise Pedagógica</div>{res['rac']}</div>", unsafe_allow_html=True)
        
        with st.container(border=True):
            partes = re.split(r'(\[\[IMG_G\d+\]\])', res['txt'])
            for p in partes:
                tag = re.search(r'\[\[IMG_G(\d+)\]\]', p)
                if tag:
                    i = int(tag.group(1))
                    im = res['map'].get(i)
                    if im: st.image(im, width=300)
                elif p.strip(): st.markdown(p.strip())
        
        docx = construir_docx_final(res['txt'], aluno, mat_c, {}, None, "Criada", sem_cabecalho=True)
        st.download_button("📥 BAIXAR DOCX (Sem Cabeçalho)", docx, "Criada.docx", "primary")

# 4. 5. 6. 7. (MANTIDOS DA V17)
with tabs[3]:
    if st.button("✨ MÁGICA DO PEI (Gerar Auto)", type="primary"):
        with st.spinner("Desenhando..."):
            desc_auto = sugerir_imagem_pei(api_key, aluno)
            url = gerar_imagem_hibrida(api_key, desc_auto, unsplash_key)
            if url: st.image(url, caption=f"Sugestão: {desc_auto}")
    st.markdown("---")
    desc = st.text_area("Descrição Manual:", key="vd")
    if st.button("🎨 Gerar"):
        url = gerar_imagem_hibrida(api_key, f"{desc} style {aluno.get('hiperfoco')}", unsplash_key)
        if url: st.image(url)

with tabs[4]:
    ass = st.text_input("Assunto:", key="rot_ass")
    if st.button("📝 CRIAR ROTEIRO", type="primary"):
        with st.spinner("Planejando..."): st.markdown(gerar_roteiro_aula(api_key, aluno, ass))

with tabs[5]:
    c1, c2 = st.columns(2)
    ass_q = c1.text_input("Assunto:", key="qg_ass")
    tema_q = c2.text_input("Tema:", value=aluno.get('hiperfoco'), key="qg_tema")
    if st.button("🗣️ GERAR CONVERSA", type="primary"):
        with st.spinner("Pensando..."): st.markdown(gerar_quebra_gelo_profundo(api_key, aluno, ass_q, tema_q))

with tabs[6]:
    ass_d = st.text_input("Tema:", key="din_ass")
    if st.button("🤝 CRIAR DINÂMICA", type="primary"):
        with st.spinner("Adaptando..."): st.markdown(gerar_dinamica_inclusiva(api_key, aluno, ass_d))
