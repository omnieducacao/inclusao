import streamlit as st
from openai import OpenAI
from datetime import date
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from pypdf import PdfReader
from fpdf import FPDF
import base64
import os
import re
import json  # <--- NOVA IMPORTAÇÃO ESSENCIAL

# --- 1. CONFIGURAÇÃO INICIAL ---
def get_favicon():
    if os.path.exists("iconeaba.png"): return "iconeaba.png"
    return "📘"

st.set_page_config(
    page_title="PEI 360º",
    page_icon=get_favicon(),
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- 1.1 SISTEMA DE BANCO DE DADOS (A NOVIDADE) ---
ARQUIVO_DB = "banco_alunos.json"

def carregar_banco():
    """Lê o arquivo físico e traz os alunos de volta"""
    if os.path.exists(ARQUIVO_DB):
        try:
            with open(ARQUIVO_DB, "r", encoding="utf-8") as f:
                return json.load(f)
        except: return []
    return []

def salvar_no_banco(novo_aluno):
    """Salva o aluno no arquivo físico para não perder"""
    banco = carregar_banco()
    # Remove versão antiga se já existir (atualização)
    banco = [a for a in banco if a['nome'] != novo_aluno['nome']]
    banco.append(novo_aluno)
    
    with open(ARQUIVO_DB, "w", encoding="utf-8") as f:
        json.dump(banco, f, ensure_ascii=False, indent=4)
    
    # Atualiza a memória imediata também
    st.session_state.banco_estudantes = banco

# Inicializa conectando com o arquivo
if 'banco_estudantes' not in st.session_state:
    st.session_state.banco_estudantes = carregar_banco()

# --- 2. UTILITÁRIOS ---
def finding_logo():
    possiveis = ["360.png", "360.jpg", "logo.png", "logo.jpg", "iconeaba.png"]
    for nome in possiveis:
        if os.path.exists(nome): return nome
    return None

def get_base64_image(image_path):
    if not image_path: return ""
    with open(image_path, "rb") as img_file:
        return base64.b64encode(img_file.read()).decode()

def ler_pdf(arquivo):
    if arquivo is None: return ""
    try:
        reader = PdfReader(arquivo)
        texto = ""
        for i, page in enumerate(reader.pages):
            if i >= 6: break 
            texto += page.extract_text() + "\n"
        return texto
    except Exception as e: return f"Erro ao ler PDF: {e}"

def limpar_texto_pdf(texto):
    if not texto: return ""
    texto = texto.replace('**', '').replace('__', '')
    texto = texto.replace('### ', '').replace('## ', '').replace('# ', '')
    texto = texto.replace('* ', '• ')
    texto = re.sub(r'[^\x00-\xff]', '', texto) 
    return texto

def calcular_idade(data_nasc):
    if not data_nasc: return ""
    hoje = date.today()
    return str(hoje.year - data_nasc.year - ((hoje.month, hoje.day) < (data_nasc.month, data_nasc.day)))

# --- 3. CSS (SEU DESIGN ORIGINAL PRESERVADO) ---
st.markdown("""
    <link href="https://cdn.jsdelivr.net/npm/remixicon@4.1.0/fonts/remixicon.css" rel="stylesheet">
    <link href="https://fonts.googleapis.com/css2?family=Nunito:wght@400;600;700;800&display=swap" rel="stylesheet">
    
    <style>
    html, body, [class*="css"] { font-family: 'Nunito', sans-serif; color: #2D3748; }
    
    :root { 
        --brand-blue: #004E92; 
        --brand-coral: #FF6B6B; 
        --card-radius: 16px;
    }

    div[data-baseweb="tab-highlight"] { background-color: transparent !important; }

    /* CARD COMUM */
    .unified-card {
        background-color: white;
        padding: 25px;
        border-radius: var(--card-radius);
        border: 1px solid #EDF2F7;
        box-shadow: 0 4px 6px rgba(0,0,0,0.03);
        margin-bottom: 20px;
    }
    
    .interactive-card:hover {
        transform: translateY(-3px);
        border-color: var(--brand-blue);
        box-shadow: 0 8px 15px rgba(0,78,146,0.08);
    }

    /* HEADER LIMPO */
    .header-clean {
        background-color: white;
        padding: 35px 40px;
        border-radius: var(--card-radius);
        border: 1px solid #EDF2F7;
        box-shadow: 0 4px 12px rgba(0,0,0,0.04);
        margin-bottom: 30px;
        display: flex;
        align-items: center;
        gap: 30px;
    }

    /* ABAS */
    .stTabs [data-baseweb="tab-list"] { gap: 10px; padding-bottom: 10px; }
    .stTabs [data-baseweb="tab"] {
        height: 45px;
        border-radius: 25px;
        padding: 0 25px;
        background-color: white;
        border: 1px solid #E2E8F0;
        font-weight: 700;
        color: #718096;
    }
    .stTabs [aria-selected="true"] {
        background-color: var(--brand-coral) !important;
        color: white !important;
        border-color: var(--brand-coral) !important;
        box-shadow: 0 4px 10px rgba(255, 107, 107, 0.2);
    }

    /* ICONES */
    .icon-box {
        width: 45px; height: 45px;
        background: #EBF8FF;
        border-radius: 12px;
        display: flex; align-items: center; justify-content: center;
        margin-bottom: 15px;
        color: var(--brand-blue);
        font-size: 22px;
    }

    /* INPUTS & HELP */
    .stTextInput input, .stTextArea textarea, .stSelectbox div[data-baseweb="select"] {
        border-radius: 12px !important;
        border-color: #E2E8F0 !important;
    }
    .stTooltipIcon { color: var(--brand-blue) !important; }

    div[data-testid="column"] .stButton button {
        border-radius: 12px !important;
        font-weight: 800 !important;
        text-transform: uppercase;
        height: 50px !important;
        letter-spacing: 0.5px;
    }
    </style>
""", unsafe_allow_html=True)

# --- 4. IA (PROMPT TURBO - ATUALIZADO AQUI) ---
def consultar_gpt(api_key, dados, contexto_pdf=""):
    if not api_key: return None, "⚠️ Configure a Chave API OpenAI na barra lateral."
    
    try:
        client = OpenAI(api_key=api_key)
        contexto_seguro = contexto_pdf[:5000] if contexto_pdf else "Sem laudo anexado."
        
        is_ahsd = "altas habilidades" in dados['diagnostico'].lower() or "superdotação" in dados['diagnostico'].lower()
        foco = "ENRIQUECIMENTO E APROFUNDAMENTO" if is_ahsd else "FLEXIBILIZAÇÃO E SUPORTE"

        prompt_sistema = """
        Você é um Neuropsicopedagogo Sênior e Especialista em Legislação Educacional (LBI).
        Sua função é gerar o PEI (Plano de Ensino Individualizado) mais completo e seguro possível.
        """

        prompt_usuario = f"""
        ESTUDANTE: {dados['nome']} | Série: {dados['serie']}
        DIAGNÓSTICO: {dados['diagnostico']} ({foco})
        MEDICAÇÃO: {dados['medicacao']}
        
        POTENCIALIDADES (Crucial para engajamento):
        - Hiperfoco: {dados['hiperfoco']}
        - Pontos Fortes: {', '.join(dados['potencias'])}
        
        CONTEXTO: {dados['historico']} | {dados['familia']}
        REDE DE APOIO: {', '.join(dados['rede_apoio'])}
        
        BARREIRAS MAPEADAS:
        - Sensorial: {', '.join(dados['b_sensorial'])}
        - Cognitiva: {', '.join(dados['b_cognitiva'])}
        - Social: {', '.join(dados['b_social'])}
        
        ESTRATÉGIAS:
        - Acesso: {', '.join(dados['estrategias_acesso'])}
        - Ensino: {', '.join(dados['estrategias_ensino'])}
        - Avaliação: {', '.join(dados['estrategias_avaliacao'])}
        
        LAUDO PDF: {contexto_seguro}
        
        GERE O RELATÓRIO TÉCNICO ESTRUTURADO:
        1. ANÁLISE DE PERFIL: Integre diagnóstico, medicação e histórico.
        2. ANÁLISE BNCC: Cite 1 Habilidade Essencial da {dados['serie']} e como adaptá-la.
        3. PLANO DE AÇÃO: Detalhe a aplicação das estratégias.
        4. PARECER FINAL: Conclusão fundamentada.
        
        5. >>> DIRETRIZES TÉCNICAS PARA O ADAPTADOR DE ATIVIDADES <<< (ESSENCIAL):
           Crie um "Manual de Instruções" para a IA que vai adaptar as provas deste aluno. Especifique:
           - GATILHOS DE INTERESSE: Como usar o hiperfoco ({dados['hiperfoco']}) para criar enunciados de questões (dê exemplos).
           - REGRAS VISUAIS: Tamanho de fonte, tipo de imagem permitida (ex: realista ou desenho), espaçamento entre linhas.
           - REGRAS DE LINGUAGEM: Complexidade do vocabulário, tamanho máximo de frase.
           - FORMATO DE RESPOSTA: Se deve priorizar múltipla escolha, ligar colunas ou resposta oral.
        """
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "system", "content": prompt_sistema}, {"role": "user", "content": prompt_usuario}],
            temperature=0.7
        )
        return response.choices[0].message.content, None
    except Exception as e: return None, f"Erro OpenAI: {str(e)}."

# --- 5. PDF ---
class PDF_V3(FPDF):
    def header(self):
        self.set_draw_color(0, 78, 146)
        self.set_line_width(0.4)
        self.rect(5, 5, 200, 287)
        
        logo = finding_logo()
        if logo: 
            self.image(logo, 12, 12, 22)
            x_offset = 40
        else: x_offset = 12
        
        self.set_xy(x_offset, 15)
        self.set_font('Arial', 'B', 14)
        self.set_text_color(0, 78, 146)
        self.cell(0, 8, 'PLANO DE ENSINO INDIVIDUALIZADO', 0, 1, 'L')
        
        self.set_xy(x_offset, 22)
        self.set_font('Arial', 'I', 9)
        self.set_text_color(100)
        self.cell(0, 5, 'Documento Oficial de Planejamento Pedagógico', 0, 1, 'L')
        self.ln(15)

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.set_text_color(128)
        self.cell(0, 10, f'Gerado via PEI 360º | Página {self.page_no()}', 0, 0, 'C')

    def section_title(self, label):
        self.ln(5)
        self.set_fill_color(240, 248, 255)
        self.set_text_color(0, 78, 146)
        self.set_font('Arial', 'B', 11)
        self.cell(0, 8, f"  {label}", 0, 1, 'L', fill=True)
        self.ln(3)

def gerar_pdf(dados, tem_anexo):
    pdf = PDF_V3()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)
    
    # Identificação
    pdf.section_title("1. IDENTIFICAÇÃO E CONTEXTO")
    pdf.set_font("Arial", size=10); pdf.set_text_color(0)
    
    nasc = dados['nasc'].strftime('%d/%m/%Y') if dados['nasc'] else "-"
    diag_display = dados['diagnostico'] if dados['diagnostico'] else ("Em análise (Vide laudo anexo)" if tem_anexo else "Não informado")
    med_display = dados['medicacao'] if dados['medicacao'] else "Não informado"

    txt_ident = (
        f"Nome: {dados['nome']}\n"
        f"Nascimento: {nasc}\n"
        f"Série: {dados['serie']} | Turma: {dados['turma']}\n"
        f"Diagnóstico: {diag_display}\n"
        f"Medicação: {med_display}"
    )
    pdf.multi_cell(0, 6, limpar_texto_pdf(txt_ident))
    
    # Rede de Apoio
    if dados['rede_apoio'] or dados['orientacoes_especialistas']:
        pdf.ln(3)
        pdf.set_font("Arial", 'B', 10)
        pdf.cell(0, 6, "Suporte Multidisciplinar:", 0, 1)
        pdf.set_font("Arial", size=10)
        
        prof = ', '.join(dados['rede_apoio']) if dados['rede_apoio'] else "-"
        ori = dados['orientacoes_especialistas'] if dados['orientacoes_especialistas'] else "-"
        pdf.multi_cell(0, 6, limpar_texto_pdf(f"Profissionais: {prof}.\nOrientações: {ori}"))

    # Relatório IA
    if dados['ia_sugestao']:
        pdf.ln(5)
        txt_ia = limpar_texto_pdf(dados['ia_sugestao'])
        pdf.multi_cell(0, 6, txt_ia)
        
    # Assinaturas
    pdf.ln(20)
    y = pdf.get_y()
    if y > 250: pdf.add_page(); y = 40
    pdf.line(20, y, 90, y); pdf.line(120, y, 190, y)
    pdf.set_font("Arial", 'I', 8)
    pdf.text(35, y+5, "Coordenação / Direção"); pdf.text(135, y+5, "Família / Responsável")
    
    return pdf.output(dest='S').encode('latin-1', 'replace')

def gerar_docx(dados):
    doc = Document()
    style = doc.styles['Normal']; style.font.name = 'Arial'; style.font.size = Pt(11)
    
    doc.add_heading('PLANO DE ENSINO INDIVIDUALIZADO', 0)
    doc.add_paragraph(f"Estudante: {dados['nome']}")
    doc.add_paragraph(f"Série: {dados['serie']} | Turma: {dados['turma']}")
    doc.add_paragraph(f"Diagnóstico: {dados['diagnostico']}")
    doc.add_paragraph(f"Medicação: {dados['medicacao']}")
    
    if dados['ia_sugestao']:
        doc.add_heading('Parecer Pedagógico', level=1)
        doc.add_paragraph(dados['ia_sugestao'])
        
    buffer = BytesIO(); doc.save(buffer); buffer.seek(0)
    return buffer

# --- 6. ESTADO ---
if 'dados' not in st.session_state:
    st.session_state.dados = {
        'nome': '', 'nasc': None, 'serie': None, 'turma': '', 
        'diagnostico': '', 'medicacao': '', 
        'historico': '', 'familia': '', 'hiperfoco': '', 'potencias': [],
        'rede_apoio': [], 'orientacoes_especialistas': '',
        'b_sensorial': [], 'sup_sensorial': '🟡 Monitorado',
        'b_cognitiva': [], 'sup_cognitiva': '🟡 Monitorado',
        'b_social': [], 'sup_social': '🟡 Monitorado',
        'estrategias_acesso': [], 'estrategias_ensino': [], 'estrategias_avaliacao': [],
        'ia_sugestao': ''
    }
if 'pdf_text' not in st.session_state: st.session_state.pdf_text = ""

# --- 7. SIDEBAR ---
with st.sidebar:
    logo = finding_logo()
    if logo: st.image(logo, width=120)
    
    if 'OPENAI_API_KEY' in st.secrets:
        api_key = st.secrets['OPENAI_API_KEY']
        st.success("✅ OpenAI Ativa")
    else:
        api_key = st.text_input("Chave OpenAI (sk-...):", type="password")
        
    st.markdown("---")
    st.markdown("### 📂 Meus Alunos")
    
    # LISTA DE ALUNOS SALVOS (MENU DE ACESSO RÁPIDO)
    banco = carregar_banco() # Lê do arquivo
    if banco:
        for aluno in banco:
            # Botão para carregar aluno na tela
            if st.button(f"👤 {aluno['nome']}", key=f"btn_{aluno['nome']}"):
                st.session_state.dados = aluno # Carrega os dados do aluno clicado
                st.success(f"Perfil de {aluno['nome']} carregado!")
    else:
        st.caption("Nenhum aluno salvo ainda.")
        
    st.markdown("---")
    st.markdown("<div style='font-size:0.8rem; color:#A0AEC0;'>PEI 360º v5.0<br>DB + Prompt Turbo</div>", unsafe_allow_html=True)

# --- 8. LAYOUT ---

# CABEÇALHO LIMPO (Proporção Ajustada)
logo_path = finding_logo()
b64_logo = get_base64_image(logo_path)
mime = "image/png" if logo_path and logo_path.endswith("png") else "image/jpeg"
img_html = f'<img src="data:{mime};base64,{b64_logo}" style="height: 80px;">' if logo_path else "" # Aumentei levemente a logo

st.markdown(f"""
    <div class="header-clean">
        {img_html}
        <div>
            <p style="margin: 0; color: #004E92; font-size: 1.3rem; font-weight: 800;">Ecossistema de Inteligência Pedagógica e Inclusiva</p>
        </div>
    </div>
""", unsafe_allow_html=True)

# ABAS
abas = ["Início", "Estudante", "Rede de Apoio", "Mapeamento", "Plano de Ação", "Consultoria IA", "Documento", "💾 Salvar & Integrar"]
tab0, tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs(abas)

# TAB 0: INÍCIO
with tab0:
    st.markdown("### <i class='ri-dashboard-line'></i> Visão Geral", unsafe_allow_html=True)
    c1, c2 = st.columns(2)
    with c1:
        st.markdown("""
        <div class="unified-card interactive-card">
            <div class="icon-box"><i class="ri-book-read-line"></i></div>
            <h4>O que é o PEI?</h4>
            <p>Instrumento oficial de acessibilidade curricular. Garante o acesso ao conhecimento, conforme a LBI.</p>
        </div>""", unsafe_allow_html=True)
    with c2:
        st.markdown("""
        <div class="unified-card interactive-card">
            <div class="icon-box"><i class="ri-scales-3-line"></i></div>
            <h4>Legalidade</h4>
            <p>Conforme Decreto 12.686/2025. O PEI independe de laudo médico fechado, focando nas barreiras.</p>
        </div>""", unsafe_allow_html=True)

    c3, c4 = st.columns(2)
    with c3:
        st.markdown("""
        <div class="unified-card interactive-card">
            <div class="icon-box"><i class="ri-brain-line"></i></div>
            <h4>Neurociência</h4>
            <p>Mapeamos Funções Executivas e Perfil Sensorial para estratégias assertivas.</p>
        </div>""", unsafe_allow_html=True)
    with c4:
        st.markdown("""
        <div class="unified-card interactive-card">
            <div class="icon-box"><i class="ri-compass-3-line"></i></div>
            <h4>BNCC</h4>
            <p>Garantia das Aprendizagens Essenciais através da flexibilização curricular.</p>
        </div>""", unsafe_allow_html=True)

# TAB 1: ESTUDANTE
with tab1:
    st.markdown("### <i class='ri-user-smile-line'></i> Dossiê do Estudante", unsafe_allow_html=True)
    
    c1, c2, c3, c4 = st.columns([3, 2, 2, 1])
    st.session_state.dados['nome'] = c1.text_input("Nome Completo", st.session_state.dados['nome'])
    st.session_state.dados['nasc'] = c2.date_input("Nascimento", st.session_state.dados['nasc'])
    st.session_state.dados['serie'] = c3.selectbox("Série/Ano", ["Infantil", "1º Ano", "2º Ano", "3º Ano", "4º Ano", "5º Ano", "Fund. II", "Ensino Médio"], placeholder="Selecione...")
    st.session_state.dados['turma'] = c4.text_input("Turma", st.session_state.dados['turma'])

    st.markdown("---")
    st.markdown("##### 1. Contexto Escolar e Familiar")
    
    ch, cf = st.columns(2)
    with ch:
        st.session_state.dados['historico'] = st.text_area(
            "Histórico Escolar", 
            st.session_state.dados['historico'], 
            height=100, 
            help="Descreva brevemente a trajetória escolar, retenções, mudanças de escola e relação com a aprendizagem."
        )
    with cf:
        st.session_state.dados['familia'] = st.text_area(
            "Contexto Familiar", 
            st.session_state.dados['familia'], 
            height=100,
            help="Rotina em casa, quem são os cuidadores principais e quais as expectativas da família."
        )

    st.markdown("##### 2. Saúde e Diagnóstico")
    col_d, col_m = st.columns(2)
    with col_d:
        st.session_state.dados['diagnostico'] = st.text_input(
            "Diagnóstico Clínico", 
            st.session_state.dados['diagnostico'], 
            placeholder="Ex: TEA, TDAH. (Se vazio, a IA buscará no PDF)",
            help="Se o aluno não tiver laudo fechado, insira a hipótese diagnóstica ou 'Em investigação'."
        )
    with col_m:
        st.session_state.dados['medicacao'] = st.text_input(
            "Medicação em uso", 
            st.session_state.dados['medicacao'], 
            placeholder="Ex: Ritalina, Risperidona...",
            help="Informação crucial para que a escola entenda possíveis efeitos colaterais (sono, sede, agitação)."
        )
    
    with st.expander("📎 Anexar Laudo (PDF)"):
        up = st.file_uploader("Arquivo PDF", type="pdf")
        if up:
            st.session_state.pdf_text = ler_pdf(up)
            st.success("PDF Anexado e pronto para análise.")

# TAB 2: REDE DE APOIO
with tab2:
    st.markdown("### <i class='ri-team-line'></i> Rede de Apoio", unsafe_allow_html=True)
    st.caption("Registre os profissionais que atendem o aluno fora da escola. A inclusão efetiva depende dessa parceria.")
    
    c_rede1, c_rede2 = st.columns(2)
    st.session_state.dados['rede_apoio'] = c_rede1.multiselect(
        "Profissionais:", 
        ["Psicólogo", "Fonoaudiólogo", "Terapeuta Ocupacional", "Neuropediatra", "Psicopedagogo", "Professor Particular"],
        placeholder="Selecione..."
    )
    st.session_state.dados['orientacoes_especialistas'] = st.text_area(
        "Orientações Técnicas (Resumo)", 
        placeholder="Ex: A fonoaudióloga recomendou pistas visuais...", 
        height=150,
        help="O que a escola precisa fazer para seguir a linha de trabalho dos terapeutas?"
    )

# TAB 3: MAPEAMENTO
with tab3:
    st.markdown("### <i class='ri-map-pin-user-line'></i> Mapeamento Integral", unsafe_allow_html=True)
    
    with st.container(border=True):
        st.markdown("#### <i class='ri-lightbulb-flash-line' style='color: var(--brand-blue);'></i> Potencialidades e Hiperfoco", unsafe_allow_html=True)
        cp1, cp2 = st.columns(2)
        with cp1:
            st.session_state.dados['hiperfoco'] = st.text_input(
                "Hiperfoco (Interesses intensos)", 
                placeholder="Ex: Dinossauros, Minecraft...",
                help="Use o interesse do aluno para criar pontes de aprendizagem."
            )
        with cp2:
            st.session_state.dados['potencias'] = st.multiselect(
                "Pontos Fortes", 
                ["Memória Visual", "Lógica Matemática", "Criatividade", "Oralidade", "Tecnologia", "Artes", "Música", "Liderança"], 
                placeholder="Selecione...",
                key="potencias_v38"
            )

    st.markdown("#### Barreiras e Suporte")
    c_bar1, c_bar2, c_bar3 = st.columns(3)
    with c_bar1:
        with st.container():
            st.markdown("##### Sensorial")
            st.session_state.dados['b_sensorial'] = st.multiselect("Barreiras:", ["Hipersensibilidade Auditiva", "Hipersensibilidade Visual", "Busca Sensorial", "Baixo Tônus"], key="b1", placeholder="Selecione...")
            st.session_state.dados['sup_sensorial'] = st.select_slider("Suporte", ["Autônomo", "Monitorado", "Substancial", "Muito Substancial"], value="Monitorado", key="s1")
    with c_bar2:
        with st.container():
            st.markdown("##### Cognitivo")
            st.session_state.dados['b_cognitiva'] = st.multiselect("Barreiras:", ["Atenção", "Memória", "Rigidez Mental", "Processamento Lento"], key="b2", placeholder="Selecione...")
            st.session_state.dados['sup_cognitiva'] = st.select_slider("Suporte", ["Autônomo", "Monitorado", "Substancial", "Muito Substancial"], value="Monitorado", key="s2")
    with c_bar3:
        with st.container():
            st.markdown("##### Social")
            st.session_state.dados['b_social'] = st.multiselect("Barreiras:", ["Interação", "Frustração", "Regras", "Isolamento"], key="b3", placeholder="Selecione...")
            st.session_state.dados['sup_social'] = st.select_slider("Suporte", ["Autônomo", "Monitorado", "Substancial", "Muito Substancial"], value="Monitorado", key="s3")

# TAB 4: PLANO DE AÇÃO
with tab4:
    st.markdown("### <i class='ri-tools-line'></i> Estratégias Pedagógicas", unsafe_allow_html=True)
    
    c_acesso, c_ensino = st.columns(2)
    with c_acesso:
        st.markdown("#### 1. Acesso ao Currículo")
        st.session_state.dados['estrategias_acesso'] = st.multiselect(
            "Recursos de Acessibilidade:", 
            ["Tempo Estendido (+25%)", "Apoio à Leitura e Escrita", "Material Ampliado", "Sala com Redução de Estímulos", "Tecnologia Assistiva", "Pausas Sensoriais"],
            placeholder="Selecione...",
            key="acesso_v38"
        )
    with c_ensino:
        st.markdown("#### 2. Metodologia de Ensino")
        st.session_state.dados['estrategias_ensino'] = st.multiselect(
            "Estratégias Didáticas:", 
            ["Fragmentação de Tarefas", "Pistas Visuais", "Enriquecimento Curricular (AH/SD)", "Antecipação de Rotina", "Projetos Práticos"],
            placeholder="Selecione...",
            key="ensino_v38"
        )
    
    st.write("")
    st.markdown("#### 3. Avaliação")
    st.session_state.dados['estrategias_avaliacao'] = st.multiselect(
        "Formato Avaliativo:", 
        ["Prova Adaptada", "Consulta Permitida", "Avaliação Oral", "Trabalho Prático", "Enunciados Curtos"],
        placeholder="Selecione...",
        key="aval_v38"
    )

# TAB 5: IA (TEXTO AMIGÁVEL)
with tab5:
    st.markdown("### <i class='ri-robot-2-line'></i> Consultoria Pedagógica Inteligente", unsafe_allow_html=True)
    
    col_btn, col_txt = st.columns([1, 2])
    with col_btn:
        st.markdown("""
        <div style="font-size: 0.95rem; color: #4A5568; line-height: 1.6; margin-bottom: 20px;">
            Olá! Sou seu assistente especializado. Fui calibrado para cruzar todas as informações que você inseriu 
            (histórico, barreiras, laudo) com a <b>Legislação Vigente</b>, a <b>BNCC</b> e os princípios da <b>Neurociência</b>.
            <br><br>
            Meu objetivo é gerar um plano seguro, fundamentado e personalizado.
        </div>
        """, unsafe_allow_html=True)
        
        if st.button("GERAR PLANO AGORA", type="primary"):
            if not st.session_state.dados['nome']: st.error("Preencha o Nome do estudante.")
            else:
                with st.spinner("Analisando dados e cruzando com a BNCC..."):
                    res, err = consultar_gpt(api_key, st.session_state.dados, st.session_state.pdf_text)
                    if err: st.error(err)
                    else: st.session_state.dados['ia_sugestao'] = res; st.success("Plano Gerado!")
    with col_txt:
        if st.session_state.dados['ia_sugestao']:
            st.text_area("Parecer Técnico:", st.session_state.dados['ia_sugestao'], height=550)
        else:
            st.markdown("<div style='padding:50px; text-align:center; color:#CBD5E0; border:2px dashed #E2E8F0; border-radius:12px;'>O Parecer Técnico aparecerá aqui.</div>", unsafe_allow_html=True)

# TAB 6: DOCUMENTO
with tab6:
    st.markdown("### <i class='ri-file-pdf-line'></i> Exportação", unsafe_allow_html=True)
    if st.session_state.dados['ia_sugestao']:
        c_pdf, c_word = st.columns(2)
        tem_anexo = len(st.session_state.pdf_text) > 0
        
        with c_pdf:
            pdf_bytes = gerar_pdf(st.session_state.dados, tem_anexo)
            st.download_button("📥 Baixar PDF", pdf_bytes, f"PEI_{st.session_state.dados['nome']}.pdf", "application/pdf", type="primary")
        with c_word:
            docx_bytes = gerar_docx(st.session_state.dados)
            st.download_button("📥 Baixar Word", docx_bytes, f"PEI_{st.session_state.dados['nome']}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.warning("Primeiro, gere o plano na aba de Consultoria IA.")

# TAB 7: SALVAR (MODIFICADA PARA USAR O BANCO)
with tab7:
    st.markdown("### <i class='ri-save-3-line'></i> Integração e Banco de Dados", unsafe_allow_html=True)
    st.markdown("""
    <div class="unified-card">
        <h4>Salvar para Adaptação de Provas</h4>
        <p>Ao salvar aqui, os dados deste estudante (incluindo as diretrizes técnicas geradas pela IA) 
        ficarão salvos no banco de dados do sistema e estarão disponíveis no módulo de <b>Adaptação de Atividades</b>,
        mesmo se você fechar o navegador.</p>
    </div>
    """, unsafe_allow_html=True)
    
    if st.button("💾 SALVAR ALUNO NO SISTEMA", type="primary"):
        if st.session_state.dados['nome']:
            # Cria o perfil rico para o adaptador
            novo_aluno = st.session_state.dados.copy()
            # Converte data para string para poder salvar no JSON
            if novo_aluno['nasc']:
                novo_aluno['nasc'] = str(novo_aluno['nasc'])
            novo_aluno['idade_calculada'] = calcular_idade(st.session_state.dados.get('nasc'))
            
            # CHAMA A FUNÇÃO DE SALVAR NO ARQUIVO
            salvar_no_banco(novo_aluno)
            st.success(f"Sucesso! O estudante **{novo_aluno['nome']}** foi salvo no banco de dados permanente.")
        else:
            st.warning("Preencha pelo menos o nome do estudante.")

st.markdown("---")
st.markdown("<div style='text-align: center; color: #A0AEC0; font-size: 0.8rem;'>PEI 360º v5.0 | Powered by OpenAI</div>", unsafe_allow_html=True)
