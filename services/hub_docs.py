"""
Serviço de geração de documentos para o Hub de Inclusão.
Funções puras de DOCX/PDF reutilizáveis.
"""
from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from fpdf import FPDF

try:
    from docx.oxml.ns import qn
except ImportError:
    qn = lambda x: x


def _adicionar_paragrafo_formatado(
    doc, texto: str, usar_caixa_alta: bool = False,
    usar_opendyslexic: bool = False, espacamento: float = 1.5
):
    """Adiciona parágrafo formatado com detecção de títulos e listas."""
    try:
        texto_formatado = texto.upper() if usar_caixa_alta else texto
        if re.match(r"^#{1,3}\s+", texto_formatado):
            nivel = len(re.match(r"^(#+)", texto_formatado).group(1))
            texto_limpo = re.sub(r"^#+\s+", "", texto_formatado)
            heading = doc.add_heading(texto_limpo, level=min(nivel, 3))
            if heading.runs:
                heading.runs[0].font.size = Pt(16 - nivel)
                if usar_opendyslexic:
                    heading.runs[0].font.name = "OpenDyslexic"
        elif re.match(r"^\d+[\.\)]\s+", texto_formatado):
            p = doc.add_paragraph(texto_formatado, style="List Number")
            if p.runs:
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.name = "OpenDyslexic" if usar_opendyslexic else "Arial"
                p.paragraph_format.line_spacing = espacamento
        elif re.match(r"^[-•*]\s+", texto_formatado):
            texto_limpo = re.sub(r"^[-•*]\s+", "", texto_formatado)
            p = doc.add_paragraph(texto_limpo, style="List Bullet")
            if p.runs:
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.name = "OpenDyslexic" if usar_opendyslexic else "Arial"
                p.paragraph_format.line_spacing = espacamento
        elif texto_formatado.isupper() and len(texto_formatado) < 100:
            p = doc.add_paragraph(texto_formatado)
            if p.runs:
                p.runs[0].font.size = Pt(15)
                p.runs[0].bold = True
                p.runs[0].font.name = "OpenDyslexic" if usar_opendyslexic else "Arial"
                p.paragraph_format.line_spacing = espacamento
        else:
            p = doc.add_paragraph(texto_formatado)
            if p.runs:
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.name = "OpenDyslexic" if usar_opendyslexic else "Arial"
                p.paragraph_format.line_spacing = espacamento
    except Exception:
        doc.add_paragraph(texto_formatado if usar_caixa_alta else texto)


def construir_docx_final(
    texto_ia: str,
    aluno: dict,
    materia: str,
    mapa_imgs: dict,
    tipo_atv: str,
    sem_cabecalho: bool = False,
    checklist_adaptacao=None,
) -> BytesIO:
    """Constrói documento DOCX final com formatação baseada no checklist."""
    doc = Document()
    usar_caixa_alta = False
    usar_opendyslexic = False
    espacamento_linhas = 1.5
    if checklist_adaptacao:
        if checklist_adaptacao.get("paragrafos_curtos") or not checklist_adaptacao.get("compreende_instrucoes_complexas"):
            usar_caixa_alta = True
            usar_opendyslexic = True
            espacamento_linhas = 1.8

    style = doc.styles["Normal"]
    if usar_opendyslexic:
        try:
            style.font.name = "OpenDyslexic"
            try:
                rFonts = style.element.rPr.rFonts
                if rFonts is not None:
                    rFonts.set(qn("w:ascii"), "OpenDyslexic")
                    rFonts.set(qn("w:hAnsi"), "OpenDyslexic")
            except Exception:
                pass
        except Exception:
            style.font.name = "Arial"
    else:
        style.font.name = "Arial"
    style.font.size = Pt(14)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = espacamento_linhas

    if not sem_cabecalho:
        titulo = doc.add_heading(f"{tipo_atv.upper()} ADAPTADA - {materia.upper()}", 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if titulo.runs:
            titulo.runs[0].font.size = Pt(16)
            titulo.runs[0].bold = True
        p_estudante = doc.add_paragraph(f"Estudante: {aluno['nome']}")
        p_estudante.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p_estudante.runs:
            p_estudante.runs[0].font.size = Pt(11)
        linha_sep = doc.add_paragraph("_" * 50)
        linha_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        secao = doc.add_heading("Atividades", level=2)
        if secao.runs:
            secao.runs[0].font.size = Pt(14)

    linhas = texto_ia.split("\n")
    for linha in linhas:
        linha_limpa = linha.strip()
        if not linha_limpa:
            doc.add_paragraph()
            continue
        tag_match = re.search(r"\[\[(IMG|GEN_IMG).*?(\d+)\]\]", linha, re.IGNORECASE)
        if tag_match:
            partes = re.split(r"(\[\[(?:IMG|GEN_IMG).*?\d+\]\])", linha, flags=re.IGNORECASE)
            for parte in partes:
                sub_match = re.search(r"(\d+)", parte)
                if ("IMG" in parte.upper() or "GEN_IMG" in parte.upper()) and sub_match:
                    num = int(sub_match.group(1))
                    img_bytes = mapa_imgs.get(num)
                    if not img_bytes and len(mapa_imgs) == 1:
                        img_bytes = list(mapa_imgs.values())[0]
                    if img_bytes:
                        try:
                            p = doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            r = p.add_run()
                            r.add_picture(BytesIO(img_bytes), width=Inches(4.5))
                            doc.add_paragraph()
                        except Exception:
                            pass
                elif parte.strip():
                    _adicionar_paragrafo_formatado(doc, parte.strip(), usar_caixa_alta, usar_opendyslexic, espacamento_linhas)
        else:
            _adicionar_paragrafo_formatado(doc, linha_limpa, usar_caixa_alta, usar_opendyslexic, espacamento_linhas)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def criar_docx_simples(texto: str, titulo: str = "Documento") -> BytesIO:
    """Cria um DOCX simples a partir de texto."""
    doc = Document()
    doc.add_heading(titulo, 0)
    for para in texto.split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def criar_pdf_generico(texto: str) -> bytes:
    """Cria um PDF simples a partir de texto."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    texto_safe = texto.encode("latin-1", "replace").decode("latin-1")
    pdf.multi_cell(0, 10, texto_safe)
    return pdf.output(dest="S").encode("latin-1")
